"""Generate DARPA-Ready Implementation Plans .docx"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

doc = Document()

# Styles
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
    return h

def add_file_table(rows):
    """rows = list of (file, change) tuples"""
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = table.rows[0].cells
    hdr[0].text = 'File'
    hdr[1].text = 'Change'
    for f, c in rows:
        row = table.add_row().cells
        row[0].text = f
        row[1].text = c
    return table

# ============================================================
# TITLE PAGE
# ============================================================
doc.add_paragraph()
doc.add_paragraph()
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = t.add_run('PIN Fusion Engine\nImplementation Plans')
run.font.size = Pt(28)
run.bold = True
run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run('DARPA-Ready Roadmap: From "It Runs" to "It Detects Threats"')
r.font.size = Pt(14)
r.font.color.rgb = RGBColor(0x64, 0x64, 0x64)

doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = meta.add_run('March 2026  |  Covers 5 implementation phases\nEstimated: 5–10 weeks total engineering')
r.font.size = Pt(11)
r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

doc.add_page_break()

# ============================================================
# EXECUTIVE SUMMARY
# ============================================================
add_heading('Executive Summary')

doc.add_paragraph(
    'The PIN Fusion Engine is operational. USGS Instantaneous Values (IV) data flows from 19 priority '
    'states through a detection pipeline that computes z-score deviations, runs binomial statistical '
    'tests for spatial coordination, clusters anomalies via BFS on HUC-8 adjacency graphs, and '
    'delivers alerts through the Dashboard webhook to email recipients via Resend.'
)
doc.add_paragraph(
    'What follows are five implementation plans that transform this working pipeline into a '
    'DARPA-deployable threat detection system. Each plan is self-contained with file-level '
    'implementation detail, test specifications, and verification criteria.'
)

doc.add_paragraph()
summary_table = doc.add_table(rows=1, cols=4)
summary_table.style = 'Light Grid Accent 1'
hdr = summary_table.rows[0].cells
hdr[0].text = '#'
hdr[1].text = 'Workstream'
hdr[2].text = 'Effort'
hdr[3].text = 'DARPA Impact'

data = [
    ('1', 'Confounder Rejection (NWS Weather + Seasonal)', '1–2 weeks', 'Low false-positive rate'),
    ('2', 'Red-Team Simulation & Validation Suite', '1–2 weeks', 'Proof of detection on attack patterns'),
    ('3', 'Alert Enrichment & Operational Hardening', '3–5 days', 'Actionable operator alerts'),
    ('4', 'ATTAINS/ECHO Connectors (Historical Depth)', '1–2 weeks', 'Real-time + regulatory context'),
    ('5', 'Documentation, API Specs & Transition Package', '1–2 weeks', 'Award-ready deliverable'),
]
for row_data in data:
    row = summary_table.add_row().cells
    for i, val in enumerate(row_data):
        row[i].text = val

doc.add_page_break()

# ============================================================
# PLAN 1: CONFOUNDER REJECTION
# ============================================================
add_heading('Plan 1: Confounder Rejection — NWS Weather, Seasonal, Infrastructure')

add_heading('Problem Statement', level=2)
doc.add_paragraph(
    'The detection pipeline fires anomalies whenever basin parameters deviate from baseline. '
    'Storms, seasonal turnover, snowmelt, and infrastructure events (dam releases, treatment plant maintenance) '
    'all cause legitimate deviations that are not pollution or attack signals. Without confounder rejection, '
    'every major storm produces a flood of false-positive alerts that erodes operator trust and makes the system '
    'unusable for real threat detection. DARPA will evaluate false-positive rate as a primary metric.'
)

add_heading('Design Decisions', level=2)
p = doc.add_paragraph()
p.add_run('1. NWS Weather API is the primary exclusion source. ').bold = True
p.add_run(
    'Free, real-time, covers all 19 states. The API returns active alerts by point (lat/lng) with '
    'event types (Flood Watch, Severe Thunderstorm, Flash Flood, etc.) that map directly to water quality '
    'confounders. We check each triggered basin\'s centroid against active NWS alerts.'
)
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('2. Seasonal baseline adjustment via monthly correction factors. ').bold = True
p.add_run(
    'Rather than building a full seasonal model (which requires a year of data), we apply published '
    'monthly correction factors for DO, temperature, and pH. These are well-characterized in limnology '
    'literature. Correction factors widen the baseline stddev during transitional months (spring/fall) '
    'to reduce false triggers from normal seasonal change.'
)
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('3. Maintenance windows from a JSON schedule file. ').bold = True
p.add_run(
    'Station-level maintenance data is not available via API. We support a manual JSON schedule '
    '(site → time range → reason) that operators can maintain. This covers known USGS station '
    'maintenance, dam releases, and treatment plant outages.'
)
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('4. Exclusion results annotate but don\'t suppress by default. ').bold = True
p.add_run(
    'Each exclusion filter returns matched=true/false with a reason. The correlation engine '
    'uses these to adjust confidence (weather match → confidence *= 0.3) and annotate the narrative, '
    'but still emits the anomaly. Operators see "WEATHER-ASSOCIATED" in the alert rather than missing '
    'a real event hidden behind a storm.'
)

add_heading('Implementation Sequence', level=2)

add_heading('Step 1: NWS Weather Exclusion Filter', level=3)
doc.add_paragraph('File: fusion-engine/src/exclusions/nwsWeather.ts')
bullets = [
    'Implement evaluate(): for each trigger, look up basin centroid via getCentroid(huc8)',
    'Fetch active NWS alerts: GET https://api.weather.gov/alerts/active?point={lat},{lng}&status=actual',
    'Cache responses for 15 minutes (NWS rate limit: ~100 req/min, we have up to 340 basins per tick)',
    'Map NWS event types to water quality impact: Flood Warning/Watch → HIGH, '
    'Flash Flood → HIGH, Severe Thunderstorm → MODERATE, Winter Storm → MODERATE, '
    'Heat Advisory → LOW (affects DO), Hurricane → HIGH',
    'Return ExclusionResult with matched=true, reason="Active NWS alert: {event} ({severity})", '
    'metadata includes alert headline, onset, expires, event type',
    'Handle 503/rate limits gracefully — return matched=false (fail-open, not fail-closed)',
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

add_heading('Step 2: Seasonal Baseline Adjustment', level=3)
doc.add_paragraph('File: fusion-engine/src/exclusions/seasonalAdjustment.ts (new)')
bullets = [
    'New ExclusionFilter implementation: SeasonalAdjustmentFilter',
    'Monthly stddev multipliers per parameter (published limnology values): '
    'DO peaks in winter (cold water holds more O2), dips in summer — widen stddev by 1.5x in Jun-Aug; '
    'Temperature naturally varies more in spring/fall — widen by 1.3x in Mar-May, Sep-Nov; '
    'pH is relatively stable — no adjustment; Turbidity spikes after rain — handled by NWS filter',
    'evaluate() checks current month, returns ExclusionResult with metadata.adjustmentFactor',
    'The correlation engine multiplies baseline stddev by this factor before z-score calculation, '
    'effectively raising the trigger threshold during transitional periods',
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

add_heading('Step 3: Maintenance Window Filter', level=3)
doc.add_paragraph('File: fusion-engine/src/exclusions/maintenance.ts')
bullets = [
    'Load schedule from data/maintenance-windows.json on init()',
    'Schema: { siteId?, huc8?, from: ISO, to: ISO, reason: string }',
    'evaluate() checks each trigger\'s sourceId and huc8 against active windows',
    'Return matched=true with reason="Scheduled maintenance: {reason}" if within window',
    'Support wildcard huc8 patterns (e.g., "031002*" for all Tampa Bay basins)',
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

add_heading('Step 4: Wire Exclusions into Correlation Engine', level=3)
doc.add_paragraph('File: fusion-engine/src/detection/correlationEngine.ts')
bullets = [
    'After detectAnomalies() produces candidates, run all exclusion filters on each candidate\'s triggers',
    'Attach ExclusionResult[] to CoordinatedAnomaly.exclusions',
    'Adjust confidence: if any exclusion matched, multiply confidence by 0.3 (weather) or 0.5 (maintenance)',
    'Annotate narrative: append "Note: {N} triggers coincide with active weather alerts" or similar',
    'Do NOT filter out the anomaly — let it through with reduced confidence and annotation',
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

add_heading('Step 5: Pass Exclusion Filters Through Tick Loop', level=3)
doc.add_paragraph('File: fusion-engine/src/index.ts')
bullets = [
    'Initialize exclusion filters in main() (already scaffolded)',
    'Pass filters array to detectAnomalies() or run separately after detection',
    'Update detectAnomalies() signature to accept ExclusionFilter[] parameter',
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

add_heading('File Summary', level=2)
add_file_table([
    ('src/exclusions/nwsWeather.ts', 'Implement NWS API fetch, centroid lookup, alert mapping, 15-min cache'),
    ('src/exclusions/seasonalAdjustment.ts', 'NEW — monthly stddev multipliers per parameter'),
    ('src/exclusions/maintenance.ts', 'Implement JSON schedule loader, time window matching'),
    ('src/exclusions/index.ts', 'Export SeasonalAdjustmentFilter'),
    ('src/detection/correlationEngine.ts', 'Accept exclusion results, adjust confidence, annotate narrative'),
    ('src/index.ts', 'Wire filters into tick loop'),
    ('data/maintenance-windows.json', 'NEW — empty schedule template'),
    ('test/exclusions/nwsWeather.test.ts', 'NEW — mock NWS API, verify matching/non-matching alerts'),
    ('test/exclusions/seasonalAdjustment.test.ts', 'NEW — verify monthly multipliers'),
    ('test/exclusions/maintenance.test.ts', 'NEW — verify window matching, wildcards'),
])

add_heading('Verification', level=2)
bullets = [
    'npm run typecheck — zero errors',
    'npm test — all existing + new exclusion tests pass',
    'Integration: inject storm-correlated triggers, verify anomaly emits with reduced confidence and weather annotation',
    'Integration: inject triggers during maintenance window, verify maintenance annotation',
    'Negative: inject triggers with no active NWS alerts, verify confidence unchanged',
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

doc.add_page_break()

# ============================================================
# PLAN 2: RED-TEAM SIMULATION & VALIDATION
# ============================================================
add_heading('Plan 2: Red-Team Simulation & Validation Suite')

add_heading('Problem Statement', level=2)
doc.add_paragraph(
    'The detection pipeline has been tested with synthetic data in unit tests, but has never been validated '
    'against realistic attack scenarios or known historical incidents. "The math works" is not "it catches threats." '
    'DARPA will require evidence of detection performance on documented events: fish kills, confirmed spills, '
    'impairment listings with dated onset. The replay harness and simulation suite provide this evidence and '
    'produce the precision/recall metrics needed for the METRICS.md credibility document.'
)

add_heading('Design Decisions', level=2)
p = doc.add_paragraph()
p.add_run('1. Replay harness operates on archived USGS IV data. ').bold = True
p.add_run(
    'USGS provides historical IV data via the same API with startDT/endDT parameters. '
    'The replay harness fetches historical windows, feeds them through the full pipeline '
    '(buffer → snapshot → baselines → detection), and collects anomalies. This validates '
    'detection against real data, not synthetic.'
)
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('2. Synthetic attack generator for red-team scenarios. ').bold = True
p.add_run(
    'A ScenarioGenerator class produces RawReading[] arrays that simulate coordinated contamination events: '
    'gradual onset (ramp over 2–4 hours), multi-parameter signatures (DO drops + pH drops + conductance spikes), '
    'spatial spread (N adjacent basins with time delay). Parameterized by attack type, magnitude, basin set, timing.'
)
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('3. Scoring framework: precision, recall, F1, detection latency. ').bold = True
p.add_run(
    'Each scenario defines expected outcomes (which basins, which parameters, minimum severity). '
    'The scorer compares pipeline output to expected outcomes and produces: precision (anomalies that matched), '
    'recall (expected events that were detected), F1 score, and detection latency (time from onset to first alert).'
)
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('4. Maryland validation dataset as first proof point. ').bold = True
p.add_run(
    'Maryland has the best data coverage (Chesapeake Bay monitoring, detailed ATTAINS records). '
    'Target incidents: documented fish kills (MD DNR records), Deer Creek PCB site, Back River WWTP '
    'consent decree events. These are public-record events with known dates and locations.'
)

add_heading('Implementation Sequence', level=2)

add_heading('Step 1: Replay Harness — Historical Data Fetch', level=3)
doc.add_paragraph('File: fusion-engine/src/replay/replayHarness.ts')
bullets = [
    'Implement runReplay(config): fetch historical USGS IV data using startDT/endDT URL parameters',
    'Chunk replay period into tick-sized windows (default 5 min, configurable)',
    'For each window: fetch readings for that time slice, feed through buffer → snapshot → detect',
    'Use a fresh TimeWindowBuffer and InMemoryBaselineStore per replay (isolated from live state)',
    'Collect all CoordinatedAnomaly[] with timestamps, return as ReplayResult',
    'Add progress callback for long replays: onTick(tickNumber, totalTicks, anomaliesFound)',
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

add_heading('Step 2: Replay Data Source — Historical USGS IV Fetcher', level=3)
doc.add_paragraph('File: fusion-engine/src/replay/historicalFetcher.ts (new)')
bullets = [
    'fetchHistoricalIv(states, from, to): same URL as live connector but with startDT/endDT instead of period',
    'Chunk large date ranges into 7-day blocks (USGS API limit)',
    'Reuse parseIvResponse() from usgsIv.ts for response parsing',
    'Cache fetched data to data/replay-cache/ to avoid re-fetching (keyed by state+dateRange)',
    'Support single-state fetch for targeted validation',
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

add_heading('Step 3: Synthetic Attack Generator', level=3)
doc.add_paragraph('File: fusion-engine/src/replay/scenarioGenerator.ts (new)')
bullets = [
    'ScenarioGenerator class with factory methods for common attack types:',
    '  - pointSource(basin, params, magnitude): single basin, sharp onset',
    '  - coordinatedContamination(basins, params, magnitude, delayMinutes): multi-basin with time delay',
    '  - gradualOnset(basins, params, peakMagnitude, rampHours): slow ramp to peak',
    '  - upstreamPropagation(startBasin, adjacencyLookup, params, magnitude): follows flow direction',
    'Each scenario returns: { readings: RawReading[], expected: ExpectedOutcome }',
    'ExpectedOutcome: { basins: string[], parameters: string[], minSeverity: Severity, onsetTime: string }',
    'Readings inject baseline-breaking values atop normal background noise',
    'Background noise generated from per-parameter normal distributions matching real USGS ranges',
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

add_heading('Step 4: Scoring Framework', level=3)
doc.add_paragraph('File: fusion-engine/src/replay/scorer.ts (new)')
bullets = [
    'scoreReplay(anomalies, expectedOutcomes): compare pipeline output to expected events',
    'Metrics computed:',
    '  - Precision: anomalies that match an expected outcome / total anomalies',
    '  - Recall: expected outcomes that were detected / total expected outcomes',
    '  - F1: harmonic mean of precision and recall',
    '  - Detection latency: median time from expected onset to first matching anomaly',
    '  - False positive rate: anomalies not matching any expected outcome / total anomalies',
    'Matching logic: anomaly matches expected if ≥50% of expected basins appear and ≥1 expected parameter triggers',
    'Output: ScoreReport with per-scenario breakdown and aggregate metrics',
    'Write report to output/validation-report-{date}.json',
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

add_heading('Step 5: Validation Scenarios', level=3)
doc.add_paragraph('File: fusion-engine/test/validation/scenarios.ts (new)')
bullets = [
    'Define 5+ named scenarios with real basin codes:',
    '  1. "tampa-bay-coordinated": 4 FL basins (03100201–04), DO+pH drop, CRITICAL',
    '  2. "chesapeake-gradual": 3 MD basins, gradual DO decline over 3 hours',
    '  3. "single-basin-false-positive": 1 basin spike (should NOT trigger coordinated anomaly)',
    '  4. "weather-confounded": 4 basins during simulated storm (should trigger but with weather annotation)',
    '  5. "upstream-propagation": sequential onset in 5 connected basins',
    'Each scenario self-contained: baselines + readings + expected outcome',
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

add_heading('Step 6: CLI Integration', level=3)
doc.add_paragraph('File: fusion-engine/src/index.ts')
bullets = [
    'Wire --replay mode to use new replay harness',
    'Add --scenario flag: npx tsx src/index.ts --scenario tampa-bay-coordinated',
    'Add --validate flag: run all scenarios, print score report',
    'Output: summary table of scenario name | detected | severity | latency | pass/fail',
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

add_heading('File Summary', level=2)
add_file_table([
    ('src/replay/replayHarness.ts', 'Full implementation: historical fetch → tick simulation → collect anomalies'),
    ('src/replay/historicalFetcher.ts', 'NEW — USGS IV historical data fetch with date chunking and caching'),
    ('src/replay/scenarioGenerator.ts', 'NEW — synthetic attack pattern generator (point source, coordinated, gradual, propagation)'),
    ('src/replay/scorer.ts', 'NEW — precision/recall/F1/latency scoring framework'),
    ('src/replay/types.ts', 'Add ExpectedOutcome, ScenarioConfig, ScoreReport types'),
    ('src/replay/index.ts', 'Export new modules'),
    ('test/validation/scenarios.ts', 'NEW — 5+ named validation scenarios with real basin codes'),
    ('test/validation/scoring.test.ts', 'NEW — unit tests for scoring logic'),
    ('test/replay/replayHarness.test.ts', 'NEW — integration test with mock historical data'),
])

add_heading('Verification', level=2)
bullets = [
    'npm test — all new tests pass',
    'npx tsx src/index.ts --validate — runs all scenarios, prints score report',
    'F1 score ≥ 0.8 on synthetic scenarios (baseline target)',
    'Detection latency ≤ 2 ticks (10 minutes) on coordinated attack scenarios',
    'Single-basin scenario correctly produces 0 anomalies (no false positive)',
    'Score report JSON written to output/ directory',
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

doc.add_page_break()

# ============================================================
# PLAN 3: ALERT ENRICHMENT & OPERATIONAL HARDENING
# ============================================================
add_heading('Plan 3: Alert Enrichment & Operational Hardening')

add_heading('Problem Statement', level=2)
doc.add_paragraph(
    'Alerts currently fire with a text narrative, severity, affected basins, and confidence score. '
    'An operator receiving this alert cannot immediately act: they need map context, parameter-level detail, '
    'escalation guidance, and links to the specific basins on the dashboard. The email template in '
    'lib/alerts/channels/email.ts already supports EnrichedAlert data (affected HUCs table, parameter '
    'deviations, coordination analysis, threat classification) — but the fusion-ingest endpoint does not '
    'produce this enrichment. This plan bridges that gap.'
)

add_heading('Design Decisions', level=2)
p = doc.add_paragraph()
p.add_run('1. Enrich at the Dashboard ingest endpoint, not in the fusion engine. ').bold = True
p.add_run(
    'The fusion engine should remain a pure detection system. Enrichment (map URLs, HUC names, '
    'classification labels) uses Dashboard-side data (HUC name lookups, dashboard URLs, classification config). '
    'The fusion-ingest endpoint transforms CoordinatedAnomaly into EnrichedAlert before dispatch.'
)
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('2. Map URL generation using basin centroids. ').bold = True
p.add_run(
    'Generate a dashboard map URL that centers on the affected basins with appropriate zoom. '
    'Format: /dashboard?lat={centroid_lat}&lng={centroid_lng}&zoom=10&highlight={huc8_list}'
)
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('3. Threat classification from fusion confidence + exclusion state. ').bold = True
p.add_run(
    'Map fusion engine output to AttackClassification: confidence ≥ 0.8 + no weather exclusions = likely_attack; '
    'confidence ≥ 0.5 = possible_attack; weather exclusions matched = likely_benign; '
    'confidence < 0.5 = insufficient_data. Include reasoning rules explaining the classification.'
)
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('4. Escalation rules based on severity × classification. ').bold = True
p.add_run(
    'CRITICAL + likely_attack → immediate escalation (additional recipients, SMS if available). '
    'HIGH + possible_attack → standard alert. MODERATE + likely_benign → digest-only (batched). '
    'LOW → log only, no alert.'
)

add_heading('Implementation Sequence', level=2)

add_heading('Step 1: Build EnrichedAlert from CoordinatedAnomaly', level=3)
doc.add_paragraph('File: Dashboard/app/api/alerts/fusion-ingest/route.ts')
bullets = [
    'After parsing the anomaly, build EnrichedAlert object:',
    '  affectedHucs: map basin codes to { huc8, name, state } using HUC lookup table',
    '  parameterDeviations: extract from triggers — paramCd, value, baseline { mean, stddev }, zScore',
    '  coordinationContext: { coordinationScore: confidence, clusterSize: basins.length, memberHucs, temporalSpread }',
    '  classification: compute from confidence + exclusion data (see decision #3)',
    '  mapUrl: generate from centroid of affected basins',
    '  relatedEvents: query recent alert log for events in same basins (24h window)',
    '  sourceHealth: current connector health status (optional, from engine state)',
    'Pass enrichment to dispatchAlerts() alongside the AlertEvent',
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

add_heading('Step 2: HUC Name Lookup Table', level=3)
doc.add_paragraph('File: Dashboard/lib/hucNames.ts (new or extend existing)')
bullets = [
    'Load HUC-8 names from a static JSON file (USGS WBD provides these)',
    'Export getHucName(huc8): string function',
    'Include state mapping: getHucState(huc8): string',
    'Data file: data/huc8-names.json (~2,500 entries, one-time generation from WBD)',
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

add_heading('Step 3: Update dispatchAlerts to Accept Enrichment', level=3)
doc.add_paragraph('File: Dashboard/lib/alerts/engine.ts')
bullets = [
    'Update dispatchAlerts signature: dispatchAlerts(events, enrichment?)',
    'Pass enrichment through to sendAlertEmail(event, enrichment)',
    'Email template already handles EnrichedAlert — no template changes needed',
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

add_heading('Step 4: Escalation Rules', level=3)
doc.add_paragraph('File: Dashboard/lib/alerts/escalation.ts (new)')
bullets = [
    'Define escalation matrix: severity × classification → action',
    'Actions: "immediate" (send now, add escalation recipients), "standard" (normal dispatch), '
    '"digest" (batch for hourly digest), "log" (record only, no send)',
    'Export getEscalationAction(severity, classification): EscalationAction',
    'Wire into fusion-ingest: apply escalation before dispatch',
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

add_heading('File Summary', level=2)
add_file_table([
    ('Dashboard/app/api/alerts/fusion-ingest/route.ts', 'Build EnrichedAlert, classification, map URL, escalation'),
    ('Dashboard/lib/hucNames.ts', 'NEW — HUC-8 name/state lookup from static JSON'),
    ('Dashboard/lib/alerts/engine.ts', 'Accept optional EnrichedAlert in dispatchAlerts'),
    ('Dashboard/lib/alerts/escalation.ts', 'NEW — severity × classification escalation matrix'),
    ('Dashboard/data/huc8-names.json', 'NEW — HUC-8 watershed names (~2,500 entries)'),
])

add_heading('Verification', level=2)
bullets = [
    'npm run build — Dashboard builds cleanly',
    'POST to fusion-ingest with test anomaly — verify email contains enrichment sections',
    'Verify map URL opens correct location with highlighted basins',
    'Verify classification badge appears in email (likely_attack / possible_attack / likely_benign)',
    'Verify escalation: CRITICAL+likely_attack sends to escalation recipients',
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

doc.add_page_break()

# ============================================================
# PLAN 4: ATTAINS/ECHO CONNECTORS
# ============================================================
add_heading('Plan 4: ATTAINS & ECHO Connectors — Historical Regulatory Depth')

add_heading('Problem Statement', level=2)
doc.add_paragraph(
    'USGS IV provides the real-time pulse. But the PIN value proposition is real-time detection layered '
    'on historical regulatory context. ATTAINS holds 430M+ assessment records mapping water body impairments. '
    'ECHO holds compliance and enforcement data (permit violations, inspections, enforcement actions). '
    'Together, they answer: "Is this anomaly in a basin with a history of problems?" and "Are there '
    'upstream facilities with compliance issues?" These are daily-cadence data sources, not real-time, '
    'but they transform raw anomalies into regulatory intelligence.'
)

add_heading('Design Decisions', level=2)
p = doc.add_paragraph()
p.add_run('1. ATTAINS data as context enrichment, not detection input. ').bold = True
p.add_run(
    'ATTAINS data doesn\'t flow through the z-score pipeline. Instead, when an anomaly fires, '
    'we look up the affected basins in ATTAINS to find: active impairments (303(d) listings), '
    'assessment status, cause/parameter matches. This context goes into the anomaly narrative '
    'and alert enrichment.'
)
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('2. ECHO data for upstream facility correlation. ').bold = True
p.add_run(
    'When an anomaly fires, query ECHO for facilities in the affected and upstream basins. '
    'Flag facilities with recent violations, active enforcement, or significant noncompliance (SNC). '
    'This answers "who might be causing this?" and is extremely valuable for regulatory users.'
)
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('3. Daily cache refresh, not per-tick fetch. ').bold = True
p.add_run(
    'Both APIs are slow (ATTAINS: 5-30s per request, ECHO: similar). Cache results in-memory '
    'with JSON persistence, refresh daily via the existing cron pattern. Detection pipeline '
    'reads from cache synchronously during tick.'
)

add_heading('Implementation Sequence', level=2)

add_heading('Step 1: ATTAINS Connector', level=3)
doc.add_paragraph('File: fusion-engine/src/ingest/epaAttains.ts')
bullets = [
    'Implement poll(): fetch assessment data for priority states from ATTAINS Public API',
    'Endpoint: https://attains.epa.gov/attains-public/api/assessments?reportingCycle=2022&state={ST}',
    'Parse response: extract impairments per HUC-12, roll up to HUC-8',
    'Map to RawReading format: source=epa-attains, parameter=impairment count or cause codes, '
    'value=number of active impairments',
    'Cache results in-memory Map<huc8, AttainsRecord> with 24h TTL',
    'Export getAttainsContext(huc8): look up cached impairment data for enrichment',
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

add_heading('Step 2: ECHO Connector', level=3)
doc.add_paragraph('File: fusion-engine/src/ingest/epaEcho.ts')
bullets = [
    'Implement poll(): fetch facility compliance data from ECHO Web Services',
    'Endpoint: https://echodata.epa.gov/echo/dfr_rest_services.get_facilities?p_huc={huc8}&output=JSON',
    'Parse response: extract facility ID, name, permit type, SNC status, last inspection, violation count',
    'Store as Map<huc8, EchoFacility[]> in memory with 24h TTL',
    'Export getUpstreamFacilities(huc8, adjacencyLookup): check affected basin + upstream neighbors',
    'Flag facilities with: SNC status active, violations in last 12 months, active enforcement orders',
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

add_heading('Step 3: Context Enrichment in Correlation Engine', level=3)
doc.add_paragraph('File: fusion-engine/src/detection/correlationEngine.ts')
bullets = [
    'After building CoordinatedAnomaly, enrich narrative with ATTAINS/ECHO context:',
    '  "Basin 03100201 has 3 active impairments (nutrients, dissolved oxygen, fecal coliform)"',
    '  "Upstream facility NPDES-MD0021234 (Back River WWTP) has 2 recent violations"',
    'Add regulatoryContext field to CoordinatedAnomaly metadata',
    'Boost confidence by 0.1 if affected basins have matching ATTAINS impairments (known problem area)',
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

add_heading('File Summary', level=2)
add_file_table([
    ('fusion-engine/src/ingest/epaAttains.ts', 'Implement ATTAINS fetch, HUC-12→8 rollup, in-memory cache'),
    ('fusion-engine/src/ingest/epaEcho.ts', 'Implement ECHO facility fetch, SNC/violation flagging'),
    ('fusion-engine/src/detection/correlationEngine.ts', 'Add regulatory context to narrative and metadata'),
    ('fusion-engine/src/types.ts', 'Add AttainsRecord, EchoFacility types'),
    ('test/ingest/epaAttains.test.ts', 'NEW — mock ATTAINS API, verify rollup and caching'),
    ('test/ingest/epaEcho.test.ts', 'NEW — mock ECHO API, verify facility flagging'),
])

doc.add_page_break()

# ============================================================
# PLAN 5: DOCUMENTATION & TRANSITION
# ============================================================
add_heading('Plan 5: Documentation, API Specs & Transition Package')

add_heading('Problem Statement', level=2)
doc.add_paragraph(
    'A working system without documentation is a demo, not a deliverable. DARPA transitions require: '
    'architecture documentation, API specifications, deployment guides, performance benchmarks, '
    'and a handover plan. This package also serves as the "leave-behind" for any stakeholder meeting — '
    'EPA, state agencies, DoD water infrastructure teams — where you need to explain what the system does, '
    'how it works, and how to deploy it without you in the room.'
)

add_heading('Deliverables', level=2)

add_heading('1. Architecture Document', level=3)
bullets = [
    'System overview diagram: data sources → fusion engine → detection → dashboard → alerts',
    'Component breakdown: ingest layer (4 connectors), basin layer (HUC-8 grouping, time windows), '
    'detection layer (z-scores, binomial test, BFS clustering, correlation engine), '
    'exclusion layer (NWS, seasonal, maintenance), output layer (console, JSONL, webhook)',
    'Data flow diagram showing a single tick from poll to alert delivery',
    'Deployment architecture: fusion engine (Node.js process) + Dashboard (Next.js on Vercel) + Resend (email)',
    'Security model: CRON_SECRET Bearer token, same-origin browser auth, no PII in alerts',
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

add_heading('2. API Specification', level=3)
bullets = [
    'Fusion Engine CLI: --once, --continuous, --replay, --scenario, --validate flags',
    'Dashboard endpoints: POST /api/alerts/fusion-ingest (webhook receiver), GET /api/alerts/history, '
    'CRUD /api/alerts/recipients, /rules, /suppress',
    'Request/response schemas in OpenAPI 3.0 format',
    'Authentication requirements per endpoint',
    'Rate limits and error codes',
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

add_heading('3. Deployment & Operations Guide', level=3)
bullets = [
    'Environment variables reference (CRON_SECRET, DASHBOARD_WEBHOOK_URL, FUSION_ALERTS_ENABLED, etc.)',
    'Baseline warmup documentation: engine requires ~50 min (10 ticks) before meaningful detection',
    'Monitoring: engine state output, tick logs, anomaly JSONL files',
    'Troubleshooting: common failure modes (USGS API rate limits, NWS API outages, webhook failures)',
    'Scaling considerations: state count, tick interval, concurrent fetches',
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

add_heading('4. Validation Report (METRICS.md)', level=3)
bullets = [
    'Generated from replay/validation suite output (Plan 2)',
    'Precision, recall, F1 on each validation scenario',
    'Detection latency distribution',
    'False positive analysis with confounder breakdown',
    'Comparison table: baseline system vs. with exclusion filters',
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

add_heading('5. DARPA Transition Plan', level=3)
bullets = [
    'Current TRL assessment (TRL 5-6: system validated in relevant environment)',
    'Path to TRL 7: operational demonstration at 2-3 state agencies',
    'Integration points for DoD water infrastructure monitoring (military bases, forward operating bases)',
    'Data rights and licensing: all data sources are public (USGS, EPA, NWS)',
    'Open-source vs. proprietary considerations',
    'Sustainment plan: cron-based operation, minimal human oversight',
    'Team and capabilities summary',
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

add_heading('File Deliverables', level=2)
add_file_table([
    ('docs/architecture.md', 'System architecture with diagrams (Mermaid)'),
    ('docs/api-spec.yaml', 'OpenAPI 3.0 specification'),
    ('docs/deployment-guide.md', 'Environment setup, warmup, monitoring, troubleshooting'),
    ('docs/validation-report.md', 'METRICS.md — precision/recall/F1 from validation suite'),
    ('docs/darpa-transition.md', 'TRL assessment, path to deployment, sustainment'),
    ('fusion-engine/README.md', 'Quickstart, architecture overview, CLI reference'),
])

doc.add_page_break()

# ============================================================
# IMPLEMENTATION TIMELINE
# ============================================================
add_heading('Implementation Timeline')

doc.add_paragraph(
    'Plans are ordered by impact and dependency. Plan 1 (confounder rejection) should ship before '
    'Plan 2 (validation) so the validation suite measures performance with exclusion filters active.'
)

timeline = doc.add_table(rows=1, cols=4)
timeline.style = 'Light Grid Accent 1'
hdr = timeline.rows[0].cells
hdr[0].text = 'Week'
hdr[1].text = 'Plan'
hdr[2].text = 'Key Deliverable'
hdr[3].text = 'Gate'

rows = [
    ('1–2', 'Plan 1: Confounder Rejection', 'NWS filter, seasonal adjustment, maintenance windows', 'False positives annotated in output'),
    ('3–4', 'Plan 2: Red-Team Simulation', 'Replay harness, scenario generator, scorer', 'F1 ≥ 0.8 on 5 scenarios'),
    ('5', 'Plan 3: Alert Enrichment', 'EnrichedAlert, map URLs, classification, escalation', 'Enriched email with all sections'),
    ('6–7', 'Plan 4: ATTAINS/ECHO', 'Assessment context, facility correlation', 'Regulatory context in narratives'),
    ('8–9', 'Plan 5: Documentation', 'Architecture, API spec, validation report, transition plan', 'Complete deliverable package'),
]
for r in rows:
    row = timeline.add_row().cells
    for i, val in enumerate(r):
        row[i].text = val

doc.add_paragraph()
doc.add_paragraph(
    'Total estimated timeline: 8–9 weeks. Plans 1 and 2 are the critical path — they establish '
    'credibility (low false positives + proven detection). Plans 3-5 are polish and packaging that '
    'make the system presentable for DARPA award and stakeholder demos.',
)

# Save
output_path = os.path.expanduser('~/Dashboard/docs/PIN-Fusion-Engine-Implementation-Plans.docx')
os.makedirs(os.path.dirname(output_path), exist_ok=True)
doc.save(output_path)
print(f'Saved to {output_path}')
