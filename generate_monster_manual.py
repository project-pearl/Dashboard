#!/usr/bin/env python3
"""
Voxels Encyclopedia Generator
==============================
Generates a DOCX encyclopedia with:
- BESTIARY: 500 procedurally-generated creatures with Pillow-rendered
  vector-style art, visual descriptions, stats tables, and lore
- ARMORY: 30 items with icon art, limited-use effects, stat bonuses,
  and cross-references to creature drop sources

Output: C:\\Users\\Doug\\Downloads\\monster_manual.docx

Dependencies: python-docx, Pillow
"""

import hashlib
import io
import math
import os
import random
from datetime import datetime

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw

# ---------------------------------------------------------------------------
# CREATURE DATA TABLES
# ---------------------------------------------------------------------------

ADJECTIVES = [
    "Crimson", "Frozen", "Fiery", "Shadow", "Venomous", "Ancient", "Cursed",
    "Spectral", "Iron", "Golden", "Obsidian", "Emerald", "Storm", "Blood",
    "Bone", "Crystal", "Void", "Ashen", "Molten", "Thorn", "Plague", "Lunar",
    "Solar", "Abyssal", "Ethereal", "Feral", "Savage", "Dire", "Eldritch",
    "Infernal", "Glacial", "Thunder", "Silent", "Howling", "Radiant", "Blighted",
    "Corrupted", "Arcane", "Primal", "Dread", "Pale", "Dark", "Scarlet",
    "Azure", "Amber", "Ivory", "Onyx", "Jade", "Ruby", "Sapphire",
]

NOUNS = [
    "Dragon", "Wolf", "Spider", "Slime", "Ogre", "Troll", "Serpent",
    "Griffin", "Phoenix", "Wyvern", "Skeleton", "Zombie", "Vampire",
    "Werewolf", "Goblin", "Orc", "Knight", "Mage", "Bear", "Lion",
    "Hound", "Imp", "Wraith", "Elemental", "Hydra", "Golem", "Shade",
    "Specter", "Dwarf", "Gnome", "Harpy", "Bat", "Eagle", "Raven",
    "Fox", "Boar", "Warg", "Sprite", "Rat", "Wasp", "Bug",
    "Blight", "Echo", "Wight", "Stalker", "Brute", "Fiend", "Sentinel",
    "Colossus", "Revenant",
]

TITLES = [
    "of the Abyss", "of the Frozen Wastes", "of the Deep", "the Destroyer",
    "the Eternal", "of the Crimson Tide", "the Devourer", "of the Void",
    "the Relentless", "of the Storm", "the Undying", "of the Flame",
    "the Silent", "of the Shadows", "the Corrupted", "of the Ancient Ruins",
    "the Forsaken", "of the Iron Gate", "the Merciless", "of the Blood Moon",
    "the Hollow", "of the Shattered Realm", "the Awakened", "of the Last Dawn",
    "the Withered", "of the Bone Fields", "the Ravenous", "of the Crystal Cavern",
    "the Cursed", "of the Emerald Forest",
]

SIZE_CATEGORIES = ["Tiny", "Small", "Medium", "Large", "Huge", "Colossal"]
SIZE_HEIGHTS = {
    "Tiny": (0.1, 0.3), "Small": (0.4, 0.9), "Medium": (1.0, 2.0),
    "Large": (2.1, 4.0), "Huge": (4.1, 8.0), "Colossal": (8.1, 15.0),
}

ABILITIES = [
    "Fire Breath", "Ice Storm", "Poison Cloud", "Lightning Strike",
    "Shadow Step", "Regeneration", "Stone Skin", "Teleport",
    "Mind Control", "Earthquake", "Plague Touch", "Wind Slash",
    "Blood Drain", "Soul Steal", "Acid Spray", "Thorn Barrage",
    "Crystal Shield", "Void Blast", "Lava Burst", "Frost Nova",
    "Thunder Clap", "Death Gaze", "Summon Minions", "Berserk Rage",
    "Petrifying Gaze", "Venom Bite", "Flame Aura", "Ice Armor",
    "Shadow Cloak", "Holy Smite",
]

HABITATS = [
    "Volcanic caverns", "Frozen tundra", "Dense jungle", "Underground ruins",
    "Swamplands", "Mountain peaks", "Desert wastelands", "Deep ocean",
    "Ancient forests", "Cursed graveyards", "Crystal caves", "Shadow realm",
    "Floating islands", "Corrupted plains", "Abyssal rift", "Enchanted grove",
]

BEHAVIORS = [
    "Highly aggressive, attacks on sight",
    "Territorial, guards its lair fiercely",
    "Nocturnal hunter, stalks prey silently",
    "Pack hunter, always found in groups",
    "Ambush predator, uses terrain for cover",
    "Solitary wanderer, avoids confrontation unless provoked",
    "Intelligent strategist, sets traps for adventurers",
    "Frenzied berserker, becomes more dangerous when wounded",
    "Patient stalker, follows targets for days",
    "Hive-minded, coordinates attacks with others of its kind",
]

DROPS = [
    "Dragon Scale", "Shadow Essence", "Frozen Heart", "Venom Sac",
    "Ancient Bone", "Crystal Shard", "Demon Horn", "Phoenix Feather",
    "Troll Blood", "Ogre Tooth", "Spectral Dust", "Iron Fragment",
    "Golden Coin", "Enchanted Gem", "Cursed Relic", "Beast Hide",
    "Elemental Core", "Dark Rune", "Luminous Pearl", "Bone Charm",
    "Stormweave Thread", "Abyssal Fang", "Moonfire Elixir", "Warden's Sigil",
    "Bloodforge Ingot", "Wraithbound Lantern", "Thornheart Seed",
    "Celestial Prism", "Venomsteel Shard", "Echoing Skull",
]

# ---------------------------------------------------------------------------
# ITEM ENCYCLOPEDIA
# ---------------------------------------------------------------------------

ITEM_ENCYCLOPEDIA = {
    "Dragon Scale": {
        "display_name": "Drakescale Aegis",
        "item_type": "Armor",
        "rarity": "Epic",
        "uses": 4,
        "effect": (
            "Wraps the wearer in overlapping dragon scales that deflect physical blows and absorb fire damage "
            "for 30 seconds. Each hit absorbed cracks one scale; after 4 absorbed hits the aegis shatters, "
            "releasing a shockwave that staggers nearby enemies. While active, the wearer cannot sprint."
        ),
        "stat_bonuses": {"DEF": 12, "HP": 30},
        "flavor_text": "Forged from shed scales of an elder wyrm, still warm to the touch centuries later.",
        "icon_shape": "shield",
        "icon_primary": "#CC2222",
        "icon_secondary": "#FFD700",
    },
    "Shadow Essence": {
        "display_name": "Vial of Condensed Nightfall",
        "item_type": "Consumable",
        "rarity": "Rare",
        "uses": 3,
        "effect": (
            "Drink to become incorporeal for 8 seconds — you can pass through walls and enemies but cannot "
            "attack or interact. Exiting inside a solid object forces you to the nearest open space and deals "
            "15% max HP damage. Leaves a shadow trail that reveals your path to observant foes."
        ),
        "stat_bonuses": {"SPD": 8},
        "flavor_text": "Bottled from the space between midnight and the witching hour.",
        "icon_shape": "potion",
        "icon_primary": "#333344",
        "icon_secondary": "#7744AA",
    },
    "Frozen Heart": {
        "display_name": "Frostbound Core",
        "item_type": "Relic",
        "rarity": "Legendary",
        "uses": 1,
        "effect": (
            "Detonates in a 15m blast that flash-freezes everything — allies, enemies, terrain — for 12 seconds. "
            "The user is encased last, gaining 3 seconds of action. The core disintegrates on use. "
            "Frozen creatures take triple damage from the next strike but are immune to all other damage while frozen."
        ),
        "stat_bonuses": {"ATK": 20, "DEF": 10},
        "flavor_text": "Torn from the chest of a frost titan. It still beats, once per century.",
        "icon_shape": "orb",
        "icon_primary": "#88CCFF",
        "icon_secondary": "#FFFFFF",
    },
    "Venom Sac": {
        "display_name": "Virulent Phial",
        "item_type": "Consumable",
        "rarity": "Uncommon",
        "uses": 6,
        "effect": (
            "Apply to any weapon to coat it in neurotoxin for 60 seconds. Poisoned strikes deal 3 damage/sec "
            "for 10 seconds (stacks up to 3 times). Handling the phial without gloves has a 10% chance to "
            "poison the user for 5 seconds each application."
        ),
        "stat_bonuses": {"ATK": 5},
        "flavor_text": "The sac still twitches reflexively when squeezed. Try not to think about it.",
        "icon_shape": "potion",
        "icon_primary": "#44CC44",
        "icon_secondary": "#228822",
    },
    "Ancient Bone": {
        "display_name": "Primordial Femur",
        "item_type": "Weapon",
        "rarity": "Rare",
        "uses": 5,
        "effect": (
            "A massive bone club that deals 2x damage to undead and skeletal creatures. Critical hits have a "
            "25% chance to summon a bone servant (10 HP, lasts 30 sec). The femur loses durability rapidly "
            "against armored targets, reducing damage by 10% per armored foe struck."
        ),
        "stat_bonuses": {"ATK": 14, "SPD": -3},
        "flavor_text": "Carbon-dated to before the first civilizations. Whatever it belonged to was enormous.",
        "icon_shape": "sword",
        "icon_primary": "#AA8855",
        "icon_secondary": "#DDCCBB",
    },
    "Crystal Shard": {
        "display_name": "Prismatic Sliver",
        "item_type": "Material",
        "rarity": "Common",
        "uses": 8,
        "effect": (
            "Can be affixed to any equipment to add a random elemental resistance (+15%) for 5 minutes. "
            "The element is determined when attached and cannot be changed. Shatters after the duration, "
            "dealing 2 damage to the wearer as fragments scatter."
        ),
        "stat_bonuses": {"DEF": 3},
        "flavor_text": "Refracts light into colors that shouldn't exist.",
        "icon_shape": "gem",
        "icon_primary": "#AAEEFF",
        "icon_secondary": "#FF88CC",
    },
    "Demon Horn": {
        "display_name": "Hellspire Gorget",
        "item_type": "Armor",
        "rarity": "Epic",
        "uses": 3,
        "effect": (
            "A neck guard carved from demon horn that grants immunity to fear and charm effects. When struck "
            "by a critical hit, the gorget retaliates with a hellfire burst dealing 25 damage in a 3m radius. "
            "Prolonged wear (>5 minutes) causes whispered temptations that invert the minimap for 10 seconds."
        ),
        "stat_bonuses": {"DEF": 15, "HP": 20},
        "flavor_text": "The horn remembers its owner and resents its new purpose.",
        "icon_shape": "shield",
        "icon_primary": "#DD3300",
        "icon_secondary": "#FFD700",
    },
    "Phoenix Feather": {
        "display_name": "Ashborn Quill",
        "item_type": "Talisman",
        "rarity": "Legendary",
        "uses": 2,
        "effect": (
            "Automatically activates upon lethal damage, fully restoring HP and granting 5 seconds of "
            "invulnerability. The revival erupts in a phoenix flame that deals 40 fire damage to all enemies "
            "within 8m. After activation, all healing received is halved for 3 minutes. Cannot trigger while "
            "the debuff is active."
        ),
        "stat_bonuses": {"HP": 50},
        "flavor_text": "Write your name with it and death forgets you — but only twice.",
        "icon_shape": "scroll",
        "icon_primary": "#FF4400",
        "icon_secondary": "#FFD700",
    },
    "Troll Blood": {
        "display_name": "Phlogiston Tincture",
        "item_type": "Consumable",
        "rarity": "Uncommon",
        "uses": 7,
        "effect": (
            "Drink to regenerate 8 HP/sec for 15 seconds. Side effect: skin turns green and emits a foul "
            "stench, reducing stealth by 50% for 1 minute. Fire damage taken during regeneration is doubled. "
            "Tastes exactly as bad as you'd expect."
        ),
        "stat_bonuses": {"HP": 10},
        "flavor_text": "Harvested under duress. The troll grew it back, and a grudge.",
        "icon_shape": "potion",
        "icon_primary": "#66AA44",
        "icon_secondary": "#AA4422",
    },
    "Ogre Tooth": {
        "display_name": "Bonecrusher Fang",
        "item_type": "Weapon",
        "rarity": "Uncommon",
        "uses": 6,
        "effect": (
            "A brutal dagger fashioned from an ogre's molar. Deals bonus damage equal to 5% of the target's "
            "current HP on each strike. Against targets below 25% HP, attacks have a 15% chance to instantly "
            "shatter bones, stunning for 2 seconds. Unwieldy — attack speed reduced by 20%."
        ),
        "stat_bonuses": {"ATK": 8, "SPD": -2},
        "flavor_text": "The ogre didn't notice it was missing. That's the scary part.",
        "icon_shape": "sword",
        "icon_primary": "#DDCCBB",
        "icon_secondary": "#886644",
    },
    "Spectral Dust": {
        "display_name": "Ghostveil Powder",
        "item_type": "Consumable",
        "rarity": "Rare",
        "uses": 4,
        "effect": (
            "Throw to create a 10m cloud of spectral particles lasting 20 seconds. Allies inside gain "
            "50% transparency (reduced enemy accuracy). Enemies inside are slowed 30% and cannot use "
            "teleportation abilities. The user becomes fully visible for 5 seconds after throwing as the "
            "powder marks the thrower's aura."
        ),
        "stat_bonuses": {"SPD": 5, "DEF": 5},
        "flavor_text": "Ground from the chains that once bound a thousand ghosts.",
        "icon_shape": "potion",
        "icon_primary": "#AABBFF",
        "icon_secondary": "#FFFFFF",
    },
    "Iron Fragment": {
        "display_name": "Forgeborn Shard",
        "item_type": "Material",
        "rarity": "Common",
        "uses": 9,
        "effect": (
            "Embed in armor to gain +5 DEF for 10 minutes. Stacks up to 3 times on the same piece. "
            "Each shard adds 0.5 kg of weight, reducing dodge chance by 2% per shard. Fragments can be "
            "recovered after combat if the armor isn't destroyed."
        ),
        "stat_bonuses": {"DEF": 5},
        "flavor_text": "Supposedly from the first sword ever forged. There are a lot of supposedly first swords.",
        "icon_shape": "ingot",
        "icon_primary": "#888899",
        "icon_secondary": "#CCCCDD",
    },
    "Golden Coin": {
        "display_name": "Sovereign's Tribute",
        "item_type": "Talisman",
        "rarity": "Uncommon",
        "uses": 8,
        "effect": (
            "Flip the coin before a battle: heads grants +20% ATK for the encounter, tails grants +20% DEF. "
            "The coin always shows the opposite of what you wanted. If used to bribe an NPC, success rate "
            "is doubled but the coin returns to your pocket — the NPC won't remember the deal."
        ),
        "stat_bonuses": {"ATK": 4, "DEF": 4},
        "flavor_text": "Minted in a kingdom that exists only in probability.",
        "icon_shape": "gem",
        "icon_primary": "#FFD700",
        "icon_secondary": "#CC9900",
    },
    "Enchanted Gem": {
        "display_name": "Arcanist's Keystone",
        "item_type": "Relic",
        "rarity": "Epic",
        "uses": 3,
        "effect": (
            "Socket into any ability slot to amplify the next ability cast by 75%. The amplified ability's "
            "cooldown is tripled. If the amplified ability kills a target, the keystone absorbs the soul, "
            "reducing the cooldown penalty by half. The gem dims with each use and cannot be recharged."
        ),
        "stat_bonuses": {"ATK": 10, "SPD": 5},
        "flavor_text": "The gem contains a library. The library contains a universe. The universe is running out of pages.",
        "icon_shape": "gem",
        "icon_primary": "#8855CC",
        "icon_secondary": "#FFD700",
    },
    "Cursed Relic": {
        "display_name": "Doomheld Idol",
        "item_type": "Relic",
        "rarity": "Legendary",
        "uses": 1,
        "effect": (
            "Instantly kills any single non-boss creature, but permanently reduces the user's max HP by 10%. "
            "The idol whispers the victim's last thought. If used against a boss, it instead deals 50% of the "
            "boss's current HP as damage but the user is stunned for 5 seconds. The idol crumbles after use, "
            "leaving only a faint sense of regret."
        ),
        "stat_bonuses": {"ATK": 25},
        "flavor_text": "It wants to be used. That should concern you.",
        "icon_shape": "skull",
        "icon_primary": "#6622AA",
        "icon_secondary": "#110011",
    },
    "Beast Hide": {
        "display_name": "Wildleather Wrap",
        "item_type": "Armor",
        "rarity": "Common",
        "uses": 9,
        "effect": (
            "Wear as a cloak to gain +5 DEF and reduce damage from beast-type creatures by 20%. While worn, "
            "beast-type creatures are less likely to initiate combat (aggro range reduced 30%). Cannot be worn "
            "with heavy armor. Loses effectiveness when wet."
        ),
        "stat_bonuses": {"DEF": 5, "SPD": 2},
        "flavor_text": "Smells like wet dog, courage, and poor life decisions.",
        "icon_shape": "shield",
        "icon_primary": "#AA8855",
        "icon_secondary": "#665533",
    },
    "Elemental Core": {
        "display_name": "Convergence Sphere",
        "item_type": "Relic",
        "rarity": "Epic",
        "uses": 2,
        "effect": (
            "Activate to merge two different elemental damage types on your weapon for 45 seconds, creating "
            "a hybrid element (fire+ice = steam, lightning+earth = magnetism, etc.). Hybrid elements bypass "
            "single-element resistances. The sphere overheats after use, dealing 10 damage to the holder "
            "and requiring 60 seconds to cool before reuse."
        ),
        "stat_bonuses": {"ATK": 12, "DEF": 8},
        "flavor_text": "Four elements in perfect, furious disagreement.",
        "icon_shape": "orb",
        "icon_primary": "#FF4400",
        "icon_secondary": "#88CCFF",
    },
    "Dark Rune": {
        "display_name": "Glyphstone of the Abyss",
        "item_type": "Talisman",
        "rarity": "Rare",
        "uses": 4,
        "effect": (
            "Place on the ground to create a 5m dark zone for 30 seconds. Enemies inside take 3 shadow "
            "damage/sec and have their abilities silenced. Allies inside (including the user) lose all "
            "light-based buffs and cannot receive healing from outside the zone. Picking up the rune early "
            "ends the zone but consumes the use."
        ),
        "stat_bonuses": {"ATK": 8},
        "flavor_text": "The inscription changes each time you look away. It's always worse.",
        "icon_shape": "scroll",
        "icon_primary": "#333344",
        "icon_secondary": "#6622AA",
    },
    "Luminous Pearl": {
        "display_name": "Tidecaller's Eye",
        "item_type": "Talisman",
        "rarity": "Rare",
        "uses": 5,
        "effect": (
            "Activate to illuminate a 20m radius for 5 minutes, revealing hidden enemies, traps, and "
            "secret passages. Water-based creatures within range are charmed for 10 seconds (one-time per "
            "creature). Using the pearl underwater triples its radius but halves its duration. Undead "
            "within the light take 2 damage/sec."
        ),
        "stat_bonuses": {"DEF": 6, "SPD": 3},
        "flavor_text": "Plucked from an oyster the size of a galleon. The oyster wants it back.",
        "icon_shape": "orb",
        "icon_primary": "#EEEEFF",
        "icon_secondary": "#88BBDD",
    },
    "Bone Charm": {
        "display_name": "Ossuary Pendant",
        "item_type": "Talisman",
        "rarity": "Uncommon",
        "uses": 7,
        "effect": (
            "Wear to gain +15% damage against undead creatures. When an undead is killed while wearing the "
            "pendant, gain a bone stack (max 5). Each stack grants +2 ATK. At 5 stacks, the pendant emits "
            "a death pulse dealing 20 damage to all undead in 8m and resets stacks. Living enemies can "
            "sense the pendant and receive +10% accuracy against the wearer."
        ),
        "stat_bonuses": {"ATK": 6, "DEF": 3},
        "flavor_text": "Carved from the knuckle of a lich who died of embarrassment.",
        "icon_shape": "skull",
        "icon_primary": "#DDCCBB",
        "icon_secondary": "#665544",
    },
    "Stormweave Thread": {
        "display_name": "Tempest Filament",
        "item_type": "Material",
        "rarity": "Uncommon",
        "uses": 8,
        "effect": (
            "Sew into any cloth armor to add lightning resistance (+20%) and a 10% chance to discharge "
            "a static shock (8 damage) when struck in melee. Thread degrades in heavy rain, losing its "
            "properties until dried. Can be woven into bowstrings for +5% arrow speed."
        ),
        "stat_bonuses": {"DEF": 4, "SPD": 3},
        "flavor_text": "Spun from clouds by spiders that live in thunderheads.",
        "icon_shape": "scroll",
        "icon_primary": "#5566AA",
        "icon_secondary": "#FFDD44",
    },
    "Abyssal Fang": {
        "display_name": "Voidbite Edge",
        "item_type": "Weapon",
        "rarity": "Epic",
        "uses": 3,
        "effect": (
            "A serrated blade that cuts through magical barriers and shields. Each successful hit drains "
            "5 MP (or equivalent) from the target and transfers it to the wielder. Against targets with "
            "no mana, the blade instead siphons 3 HP/hit. The void energy makes the blade vibrate "
            "uncomfortably — accuracy reduced by 10% on rapid successive swings."
        ),
        "stat_bonuses": {"ATK": 18, "SPD": -1},
        "flavor_text": "Pulled from the jaw of something that eats light.",
        "icon_shape": "sword",
        "icon_primary": "#1A0033",
        "icon_secondary": "#7744AA",
    },
    "Moonfire Elixir": {
        "display_name": "Lunaris Draught",
        "item_type": "Consumable",
        "rarity": "Rare",
        "uses": 4,
        "effect": (
            "Drink to gain moonfire vision for 3 minutes: see through illusions, detect shapechangers, "
            "and deal 25% bonus damage to lycanthropes. At night, also grants +10% SPD and silent footsteps. "
            "During daytime the draught is half as effective and causes mild photosensitivity (-5% accuracy "
            "in bright light)."
        ),
        "stat_bonuses": {"ATK": 6, "SPD": 5},
        "flavor_text": "Brewed only during lunar eclipses. Tastes like silver and regret.",
        "icon_shape": "potion",
        "icon_primary": "#CCCCEE",
        "icon_secondary": "#FFD700",
    },
    "Warden's Sigil": {
        "display_name": "Aegis Glyph",
        "item_type": "Talisman",
        "rarity": "Rare",
        "uses": 5,
        "effect": (
            "Press against a doorway or passage to seal it with a warding barrier (500 HP) for 2 minutes. "
            "Allies can pass through freely; enemies must break it. The sigil alerts the placer when the "
            "barrier is struck. Can also be placed on the ground to create a 3m safe zone where healing "
            "is boosted by 25%. Only one sigil can be active at a time."
        ),
        "stat_bonuses": {"DEF": 10},
        "flavor_text": "Bears the mark of a warden-order that guards things even gods forgot.",
        "icon_shape": "scroll",
        "icon_primary": "#FFD700",
        "icon_secondary": "#3388DD",
    },
    "Bloodforge Ingot": {
        "display_name": "Crimson Steel Bar",
        "item_type": "Material",
        "rarity": "Uncommon",
        "uses": 7,
        "effect": (
            "Use at a forge to upgrade any weapon by +1 tier, adding lifesteal (2% of damage dealt healed). "
            "The upgrade is permanent but the weapon gains a bloodthirst — if no enemy is struck within "
            "60 seconds, the weapon deals 1 damage/sec to its wielder until sheathed. Incompatible with "
            "holy-enchanted weapons."
        ),
        "stat_bonuses": {"ATK": 7},
        "flavor_text": "Smelted at temperatures that required burning something alive. Don't ask what.",
        "icon_shape": "ingot",
        "icon_primary": "#991111",
        "icon_secondary": "#CC4444",
    },
    "Wraithbound Lantern": {
        "display_name": "Lantern of Lost Souls",
        "item_type": "Relic",
        "rarity": "Legendary",
        "uses": 2,
        "effect": (
            "Open the lantern shutter to release a wraith swarm that targets the nearest 5 enemies, "
            "dealing 30 spirit damage each and terrifying them (flee for 6 seconds). Each enemy killed by "
            "the wraiths adds a new soul to the lantern, increasing its damage by 5 for future uses. "
            "If no enemies are killed, one of the lantern's souls escapes permanently. The wraiths cannot "
            "distinguish friend from foe if the user is below 20% HP."
        ),
        "stat_bonuses": {"ATK": 15, "DEF": 5},
        "flavor_text": "The lantern flickers in rhythm with a heartbeat. Not yours.",
        "icon_shape": "orb",
        "icon_primary": "#667799",
        "icon_secondary": "#AABBFF",
    },
    "Thornheart Seed": {
        "display_name": "Briarcrown Kernel",
        "item_type": "Consumable",
        "rarity": "Common",
        "uses": 10,
        "effect": (
            "Plant in any soil to grow a thorn barrier (3m wide, 2m tall) in 5 seconds. The barrier "
            "has 80 HP and deals 5 damage to anything that passes through. Lasts 2 minutes or until "
            "destroyed. Fire damage destroys it instantly. Can be eaten in desperation to heal 15 HP "
            "but causes 5 seconds of choking (cannot use abilities)."
        ),
        "stat_bonuses": {"DEF": 2},
        "flavor_text": "Grows in soil, sand, stone, or flesh. It's not picky.",
        "icon_shape": "gem",
        "icon_primary": "#558822",
        "icon_secondary": "#88CC44",
    },
    "Celestial Prism": {
        "display_name": "Astral Refractor",
        "item_type": "Relic",
        "rarity": "Epic",
        "uses": 2,
        "effect": (
            "Aim at any single-target ability (yours or an enemy's) to split it into 3 weaker copies "
            "targeting different enemies (40% damage each). Can also redirect an incoming projectile ability "
            "back at its caster at 60% power. The prism absorbs ambient light when used, plunging a 10m "
            "area into darkness for 8 seconds."
        ),
        "stat_bonuses": {"ATK": 8, "SPD": 4},
        "flavor_text": "Fell from the sky trailing seven colors that burned different songs into the grass.",
        "icon_shape": "gem",
        "icon_primary": "#AAEEFF",
        "icon_secondary": "#FF88CC",
    },
    "Venomsteel Shard": {
        "display_name": "Toxiforge Splinter",
        "item_type": "Material",
        "rarity": "Uncommon",
        "uses": 6,
        "effect": (
            "Embed in a weapon to add poison damage (4/sec for 8 sec) to attacks. The venom corrodes "
            "the weapon over time — after 20 poisoned strikes, the weapon's base damage is permanently "
            "reduced by 5% unless repaired at a forge. Glows faintly green in the presence of antidotes."
        ),
        "stat_bonuses": {"ATK": 6},
        "flavor_text": "An alloy that shouldn't exist: metal that is also alive, and angry.",
        "icon_shape": "ingot",
        "icon_primary": "#44CC44",
        "icon_secondary": "#888899",
    },
    "Echoing Skull": {
        "display_name": "Resonance Cranium",
        "item_type": "Talisman",
        "rarity": "Rare",
        "uses": 3,
        "effect": (
            "Activate to replay the last ability used by any creature within 15m (friend or foe). The "
            "replayed ability uses the skull-holder's stats, not the original caster's. Cannot echo "
            "Legendary-tier abilities. After echoing, the user hears phantom whispers for 30 seconds, "
            "making it impossible to detect audio-based traps or ambushes."
        ),
        "stat_bonuses": {"ATK": 5, "SPD": 3},
        "flavor_text": "It repeats the last thing it heard. It heard something awful.",
        "icon_shape": "skull",
        "icon_primary": "#DDCCBB",
        "icon_secondary": "#8855CC",
    },
}

RARITY_COLORS = {
    "Common": "#999999",
    "Uncommon": "#22AA44",
    "Rare": "#3388DD",
    "Epic": "#8855CC",
    "Legendary": "#FF8800",
}

# ---------------------------------------------------------------------------
# COLOR PALETTES (adjective -> primary hex, secondary hex)
# ---------------------------------------------------------------------------

ADJECTIVE_COLORS = {
    "Crimson":    ("#CC2222", "#880000"),
    "Frozen":     ("#88CCFF", "#CCE8FF"),
    "Fiery":      ("#FF4400", "#FFD700"),
    "Shadow":     ("#333344", "#111122"),
    "Venomous":   ("#44CC44", "#228822"),
    "Ancient":    ("#AA8855", "#665533"),
    "Cursed":     ("#6622AA", "#331166"),
    "Spectral":   ("#AABBFF", "#667799"),
    "Iron":       ("#888899", "#555566"),
    "Golden":     ("#FFD700", "#CC9900"),
    "Obsidian":   ("#222222", "#444455"),
    "Emerald":    ("#22AA44", "#116633"),
    "Storm":      ("#5566AA", "#334477"),
    "Blood":      ("#991111", "#660000"),
    "Bone":       ("#DDCCBB", "#AA9988"),
    "Crystal":    ("#AAEEFF", "#88BBDD"),
    "Void":       ("#110022", "#220044"),
    "Ashen":      ("#999988", "#777766"),
    "Molten":     ("#FF6600", "#CC3300"),
    "Thorn":      ("#558822", "#336611"),
    "Plague":     ("#88AA22", "#556611"),
    "Lunar":      ("#CCCCEE", "#9999BB"),
    "Solar":      ("#FFEE44", "#DDBB22"),
    "Abyssal":    ("#1A0033", "#0D001A"),
    "Ethereal":   ("#CCDDFF", "#99AADD"),
    "Feral":      ("#AA6633", "#774422"),
    "Savage":     ("#BB4422", "#882211"),
    "Dire":       ("#663333", "#441111"),
    "Eldritch":   ("#7744AA", "#553388"),
    "Infernal":   ("#DD3300", "#AA1100"),
    "Glacial":    ("#66BBEE", "#4499CC"),
    "Thunder":    ("#FFDD44", "#CCAA22"),
    "Silent":     ("#556677", "#334455"),
    "Howling":    ("#7788AA", "#556688"),
    "Radiant":    ("#FFEEAA", "#DDCC88"),
    "Blighted":   ("#667722", "#445511"),
    "Corrupted":  ("#553366", "#331144"),
    "Arcane":     ("#8855CC", "#6633AA"),
    "Primal":     ("#885522", "#663311"),
    "Dread":      ("#442233", "#221111"),
    "Pale":       ("#DDDDDD", "#BBBBBB"),
    "Dark":       ("#222233", "#111122"),
    "Scarlet":    ("#DD2244", "#AA1133"),
    "Azure":      ("#3388DD", "#2266AA"),
    "Amber":      ("#DDAA33", "#BB8822"),
    "Ivory":      ("#EEEECC", "#CCCCAA"),
    "Onyx":       ("#111111", "#333333"),
    "Jade":       ("#33AA66", "#228844"),
    "Ruby":       ("#CC1144", "#990033"),
    "Sapphire":   ("#2244BB", "#113388"),
}

# ---------------------------------------------------------------------------
# BODY ARCHETYPE MAPPING
# ---------------------------------------------------------------------------

BODY_ARCHETYPES = {
    "Ogre": "humanoid", "Troll": "humanoid", "Knight": "humanoid",
    "Goblin": "humanoid", "Orc": "humanoid", "Dwarf": "humanoid",
    "Gnome": "humanoid", "Mage": "humanoid", "Skeleton": "humanoid",
    "Zombie": "humanoid", "Vampire": "humanoid", "Werewolf": "humanoid",
    "Wight": "humanoid", "Golem": "humanoid", "Brute": "humanoid",
    "Stalker": "humanoid", "Fiend": "humanoid", "Sentinel": "humanoid",
    "Colossus": "humanoid", "Revenant": "humanoid",
    "Wolf": "quadruped", "Bear": "quadruped", "Lion": "quadruped",
    "Hound": "quadruped", "Fox": "quadruped", "Boar": "quadruped",
    "Warg": "quadruped",
    "Dragon": "winged", "Griffin": "winged", "Phoenix": "winged",
    "Wyvern": "winged", "Eagle": "winged", "Harpy": "winged",
    "Bat": "winged", "Raven": "winged",
    "Serpent": "serpentine", "Hydra": "serpentine",
    "Slime": "amorphous", "Elemental": "amorphous", "Wraith": "amorphous",
    "Shade": "amorphous", "Specter": "amorphous", "Echo": "amorphous",
    "Blight": "amorphous",
    "Spider": "arachnid",
    "Imp": "small", "Sprite": "small", "Bug": "small",
    "Wasp": "small", "Rat": "small",
}

# ---------------------------------------------------------------------------
# TEXTURE / MATERIAL MAPPING
# ---------------------------------------------------------------------------

NOUN_TEXTURES = {
    "Dragon": "scales", "Serpent": "scales", "Hydra": "scales",
    "Wyvern": "scales", "Griffin": "feathers and fur",
    "Phoenix": "blazing feathers", "Wolf": "fur", "Bear": "thick fur",
    "Lion": "tawny fur", "Hound": "short fur", "Fox": "sleek fur",
    "Boar": "bristly hide", "Warg": "matted fur",
    "Spider": "chitinous exoskeleton", "Bug": "chitin", "Wasp": "chitin",
    "Slime": "translucent gel", "Elemental": "raw elemental matter",
    "Golem": "stone", "Skeleton": "bare bone", "Zombie": "rotting flesh",
    "Vampire": "pale skin", "Werewolf": "coarse fur",
    "Wraith": "shadow-smoke", "Shade": "shadow-smoke",
    "Specter": "ethereal mist", "Echo": "semi-transparent ectoplasm",
    "Blight": "diseased organic matter",
    "Ogre": "thick warty hide", "Troll": "rubbery hide",
    "Knight": "plate armor", "Mage": "robes",
    "Goblin": "leathery green skin", "Orc": "tough green skin",
    "Dwarf": "weathered skin under armor", "Gnome": "tanned skin",
    "Imp": "leathery skin", "Sprite": "luminous skin",
    "Rat": "mangy fur", "Harpy": "feathers",
    "Bat": "leathery membrane", "Eagle": "feathers", "Raven": "black feathers",
    "Wight": "desiccated flesh", "Stalker": "dark leather",
    "Brute": "scarred hide", "Fiend": "infernal hide",
    "Sentinel": "enchanted armor", "Colossus": "stone and metal",
    "Revenant": "decayed flesh under armor",
}

ADJECTIVE_TEXTURES = {
    "Crystal": "crystalline", "Iron": "metallic", "Golden": "gilded",
    "Obsidian": "volcanic glass", "Bone": "ossite",
    "Molten": "magma-crusted", "Frozen": "frost-rimed", "Glacial": "ice-encrusted",
}

# ---------------------------------------------------------------------------
# ABILITY AURA / EFFECT DESCRIPTIONS
# ---------------------------------------------------------------------------

ABILITY_EFFECTS = {
    "Fire Breath": "wreathed in flickering flames",
    "Ice Storm": "surrounded by swirling ice crystals",
    "Poison Cloud": "emanating a sickly green miasma",
    "Lightning Strike": "crackling with arcs of electricity",
    "Shadow Step": "trailing wisps of living shadow",
    "Regeneration": "pulsing with a soft green healing glow",
    "Stone Skin": "coated in a layer of grinding stone plates",
    "Teleport": "shimmering with unstable spatial distortion",
    "Mind Control": "radiating a faint psychic halo",
    "Earthquake": "causing the ground to tremble underfoot",
    "Plague Touch": "dripping with infectious ichor",
    "Wind Slash": "surrounded by razor-thin wind currents",
    "Blood Drain": "trailing crimson mist",
    "Soul Steal": "orbited by faint ghostly wisps",
    "Acid Spray": "dripping caustic yellow-green fluid",
    "Thorn Barrage": "bristling with regenerating thorns",
    "Crystal Shield": "encased in floating crystal shards",
    "Void Blast": "distorting light around its body",
    "Lava Burst": "oozing rivulets of molten rock",
    "Frost Nova": "radiating waves of freezing cold",
    "Thunder Clap": "humming with static energy",
    "Death Gaze": "with eyes that glow an eerie pale light",
    "Summon Minions": "accompanied by lesser creatures",
    "Berserk Rage": "covered in self-inflicted battle scars",
    "Petrifying Gaze": "with stone-grey eyes that paralyze",
    "Venom Bite": "with fangs dripping venom",
    "Flame Aura": "engulfed in a constant fire aura",
    "Ice Armor": "sheathed in thick blue ice",
    "Shadow Cloak": "partially invisible in darkness",
    "Holy Smite": "glowing with divine golden light",
}

ARCHETYPE_FORMS = {
    "humanoid": "Bipedal humanoid with an upright torso, two arms, and two legs",
    "quadruped": "Four-legged beast with a powerful horizontal body and broad head",
    "winged": "Winged creature with a muscular body and large triangular wings",
    "serpentine": "Long sinuous serpentine body coiling in an S-curve",
    "amorphous": "Amorphous, shifting form with no fixed shape — ethereal and bloblike",
    "arachnid": "Eight-legged arachnid with a bulbous abdomen and fanged head",
    "small": "Compact, diminutive creature with quick darting movements",
}

NOUN_FEATURES = {
    "Dragon": "massive horns, razor claws, spiked tail, armored underbelly",
    "Griffin": "eagle head, lion body, taloned forelegs, feathered wings",
    "Phoenix": "crest of blazing plumage, long tail feathers trailing fire",
    "Wyvern": "barbed tail, leathery wings, elongated snout",
    "Hydra": "multiple serpentine heads, regenerating necks",
    "Spider": "eight multi-jointed legs, cluster of eyes, venomous chelicerae",
    "Wolf": "pointed ears, bushy tail, powerful jaws",
    "Bear": "massive paws with claws, humped shoulders, small rounded ears",
    "Lion": "flowing mane, retractable claws, muscular build",
    "Ogre": "jutting lower fangs, thick brow ridge, enormous hands",
    "Troll": "elongated arms, hunched posture, protruding nose",
    "Skeleton": "exposed rib cage, hollow eye sockets, bony claws",
    "Zombie": "shambling gait, exposed muscle tissue, vacant eyes",
    "Vampire": "elegant features, elongated canines, clawed fingers",
    "Werewolf": "lupine snout on humanoid frame, clawed hands, bristling fur",
    "Goblin": "oversized pointed ears, sharp teeth, beady eyes",
    "Orc": "tusks, broad shoulders, scarred face",
    "Knight": "full plate armor, visored helm, shield and weapon",
    "Mage": "flowing robes, glowing staff, arcane sigils on skin",
    "Golem": "carved runes, block-like limbs, featureless face",
    "Slime": "undulating surface, visible trapped objects within, pseudopods",
    "Elemental": "body composed of raw element — fire/water/earth/air swirling",
    "Wraith": "hooded skeletal figure, spectral chains, glowing eyes",
    "Shade": "featureless dark silhouette, elongated limbs, no face",
    "Specter": "translucent floating form, hollow screaming mouth",
    "Imp": "small horns, bat-like wings, barbed tail, mischievous grin",
    "Sprite": "delicate insect wings, luminous body, tiny stature",
    "Harpy": "woman's torso with bird legs and wings, talons",
    "Bat": "large membrane wings, sonar-dish ears, fanged mouth",
    "Eagle": "sharp hooked beak, powerful talons, keen eyes",
    "Raven": "glossy black plumage, intelligent eyes, sharp beak",
    "Hound": "lean muscular body, keen nose, bared teeth",
    "Fox": "bushy tail, pointed snout, alert triangular ears",
    "Boar": "curved tusks, bristly back, compact muscular body",
    "Warg": "dire-wolf proportions, scarred muzzle, glowing eyes",
    "Dwarf": "stocky build, braided beard, heavy armor",
    "Gnome": "oversized head, pointed hat, tool belt",
    "Bug": "segmented body, compound eyes, mandibles",
    "Wasp": "striped abdomen, translucent wings, stinger",
    "Rat": "long naked tail, twitching whiskers, yellow teeth",
    "Echo": "flickering duplicate of another creature, slightly out of phase",
    "Blight": "twisted plant-like growths erupting from its form",
    "Wight": "gaunt desiccated corpse in rusted armor, glowing eyes",
    "Stalker": "lean predatory build, elongated fingers, hooded face",
    "Brute": "massively muscled, small head, enormous fists",
    "Fiend": "horns, cloven hooves, barbed tail, sulfurous skin",
    "Sentinel": "towering armored figure, glowing visor, massive weapon",
    "Colossus": "enormous stone-and-metal construct, ancient runes",
    "Revenant": "skeletal warrior in battered plate, spectral fire in eye sockets",
}

# ---------------------------------------------------------------------------
# CREATURE GENERATION
# ---------------------------------------------------------------------------

def _seeded_random(seed_str):
    h = int(hashlib.md5(seed_str.encode()).hexdigest(), 16)
    return random.Random(h)


def generate_creatures(count=500):
    creatures = []
    used_names = set()
    master_rng = random.Random(42)

    attempts = 0
    while len(creatures) < count and attempts < count * 10:
        attempts += 1
        adj = master_rng.choice(ADJECTIVES)
        noun = master_rng.choice(NOUNS)
        title = master_rng.choice(TITLES)
        name = f"{adj} {noun} {title}"
        if name in used_names:
            continue
        used_names.add(name)

        rng = _seeded_random(name)
        size = rng.choice(SIZE_CATEGORIES)
        h_lo, h_hi = SIZE_HEIGHTS[size]
        height_m = round(rng.uniform(h_lo, h_hi), 1)
        level = rng.randint(1, 100)
        hp = level * rng.randint(8, 25) + rng.randint(50, 500)
        attack = rng.randint(5, 50) + level // 2
        defense = rng.randint(3, 40) + level // 3
        speed = rng.randint(1, 30)
        ability = rng.choice(ABILITIES)
        habitat = rng.choice(HABITATS)
        behavior = rng.choice(BEHAVIORS)
        drops = rng.sample(DROPS, k=rng.randint(1, 4))

        archetype = BODY_ARCHETYPES.get(noun, "humanoid")
        primary, secondary = ADJECTIVE_COLORS.get(adj, ("#888888", "#555555"))
        texture = ADJECTIVE_TEXTURES.get(adj, NOUN_TEXTURES.get(noun, "rough hide"))
        features = NOUN_FEATURES.get(noun, "unremarkable features")
        body_form = ARCHETYPE_FORMS.get(archetype, "humanoid form")
        aura_effect = ABILITY_EFFECTS.get(ability, "no visible magical aura")

        lore = _generate_lore(adj, noun, title, ability, habitat, behavior, rng)
        visual_desc = _build_visual_description(
            name, body_form, primary, secondary, size, height_m,
            features, texture, aura_effect, archetype
        )

        creatures.append({
            "name": name, "adjective": adj, "noun": noun, "title": title,
            "size": size, "height_m": height_m, "level": level, "hp": hp,
            "attack": attack, "defense": defense, "speed": speed,
            "ability": ability, "habitat": habitat, "behavior": behavior,
            "drops": drops, "archetype": archetype,
            "primary_color": primary, "secondary_color": secondary,
            "texture": texture, "features": features,
            "body_form": body_form, "aura_effect": aura_effect,
            "lore": lore, "visual_description": visual_desc,
        })

    creatures.sort(key=lambda c: c["name"])
    return creatures


def _generate_lore(adj, noun, title, ability, habitat, behavior, rng):
    templates = [
        f"Long feared by travelers, the {adj} {noun} {title} haunts the {habitat.lower()}. "
        f"It is known for its {ability.lower()} ability, which has devastated entire settlements. "
        f"{behavior}. Few who encounter this creature live to tell the tale.",
        f"Legends speak of the {adj} {noun} {title}, a terror born in the {habitat.lower()}. "
        f"Wielding the power of {ability.lower()}, it has earned its fearsome reputation. "
        f"{behavior}. Adventurers are advised to approach with extreme caution.",
        f"The {adj} {noun} {title} emerged from the depths of the {habitat.lower()} centuries ago. "
        f"Its mastery of {ability.lower()} makes it a formidable opponent. "
        f"{behavior}. Scholars debate whether it can truly be slain.",
        f"In the {habitat.lower()}, whispers tell of the {adj} {noun} {title}. "
        f"Armed with {ability.lower()}, this creature has claimed countless victims. "
        f"{behavior}. Bounties on its head remain uncollected to this day.",
    ]
    return rng.choice(templates)


def _build_visual_description(name, body_form, primary, secondary, size, height_m,
                               features, texture, aura_effect, archetype):
    return (
        f"BODY FORM: {body_form}. "
        f"SIZE: {size} category — approximately {height_m}m tall/long. "
        f"COLORING: Primary {primary}, secondary {secondary}. "
        f"TEXTURE/MATERIAL: Covered in {texture}. "
        f"DISTINGUISHING FEATURES: {features}. "
        f"AURA/EFFECTS: {aura_effect}. "
        f"ARCHETYPE: {archetype}."
    )


# ---------------------------------------------------------------------------
# PILLOW VECTOR-STYLE ART GENERATION
# ---------------------------------------------------------------------------

CANVAS_W, CANVAS_H = 400, 500  # 2x for crisp rendering, embed at smaller size
SCALE = 2  # coordinates are authored at 200x250, drawn at 2x


def _hex_to_rgb(h):
    h = h.lstrip("#")
    return tuple(int(h[i:i+2], 16) for i in (0, 2, 4))


def _rgba(hex_color, alpha=255):
    return _hex_to_rgb(hex_color) + (alpha,)


def _darken(hex_color, factor=0.6):
    r, g, b = _hex_to_rgb(hex_color)
    return (int(r * factor), int(g * factor), int(b * factor))


def _darken_hex(hex_color, factor=0.6):
    r, g, b = _darken(hex_color, factor)
    return f"#{r:02x}{g:02x}{b:02x}"


def _lighten(hex_color, factor=0.3):
    r, g, b = _hex_to_rgb(hex_color)
    return (
        min(255, int(r + (255 - r) * factor)),
        min(255, int(g + (255 - g) * factor)),
        min(255, int(b + (255 - b) * factor)),
    )


def _accent_color(primary, secondary):
    r1, g1, b1 = _hex_to_rgb(primary)
    r2, g2, b2 = _hex_to_rgb(secondary)
    return (
        min(255, (r1 + r2) // 2 + 40),
        min(255, (g1 + g2) // 2 + 20),
        min(255, (b1 + b2) // 2 + 30),
    )


def s(val):
    """Scale a coordinate by SCALE factor."""
    return int(val * SCALE)


def _ellipse_bbox(cx, cy, rx, ry):
    """Return bounding box for an ellipse centered at (cx, cy) with radii rx, ry. Coords pre-scaled."""
    return [cx - rx, cy - ry, cx + rx, cy + ry]


def _draw_ellipse(draw, cx, cy, rx, ry, fill=None, outline=None, width=1):
    """Draw an ellipse at logical coords (auto-scaled)."""
    bbox = _ellipse_bbox(s(cx), s(cy), s(rx), s(ry))
    draw.ellipse(bbox, fill=fill, outline=outline, width=width)


def _draw_rect(draw, x, y, w, h, fill=None, outline=None):
    """Draw a rectangle at logical coords (auto-scaled)."""
    draw.rectangle([s(x), s(y), s(x + w), s(y + h)], fill=fill, outline=outline)


def _draw_polygon(draw, points, fill=None, outline=None):
    """Draw polygon. Points are logical coords, auto-scaled."""
    scaled = [(s(px), s(py)) for px, py in points]
    draw.polygon(scaled, fill=fill, outline=outline)


def _draw_line(draw, x1, y1, x2, y2, fill, width=2):
    """Draw a line at logical coords (auto-scaled)."""
    draw.line([(s(x1), s(y1)), (s(x2), s(y2))], fill=fill, width=int(width * SCALE))


def _draw_thick_curve(draw, points, fill, width=3):
    """Approximate a curve with line segments. Points are logical coords."""
    scaled = [(s(px), s(py)) for px, py in points]
    for i in range(len(scaled) - 1):
        draw.line([scaled[i], scaled[i + 1]], fill=fill, width=int(width * SCALE))


def _bezier_points(p0, p1, p2, p3, steps=20):
    """Compute cubic Bezier curve points."""
    pts = []
    for i in range(steps + 1):
        t = i / steps
        u = 1 - t
        x = u**3 * p0[0] + 3 * u**2 * t * p1[0] + 3 * u * t**2 * p2[0] + t**3 * p3[0]
        y = u**3 * p0[1] + 3 * u**2 * t * p1[1] + 3 * u * t**2 * p2[1] + t**3 * p3[1]
        pts.append((x, y))
    return pts


def _quad_bezier_points(p0, p1, p2, steps=20):
    """Compute quadratic Bezier curve points."""
    pts = []
    for i in range(steps + 1):
        t = i / steps
        u = 1 - t
        x = u**2 * p0[0] + 2 * u * t * p1[0] + t**2 * p2[0]
        y = u**2 * p0[1] + 2 * u * t * p1[1] + t**2 * p2[1]
        pts.append((x, y))
    return pts


# ---------------------------------------------------------------------------
# NOUN-SPECIFIC PROPORTION MODIFIERS
# (head_scale, torso_w, torso_h, arm_len, leg_len, stance_w, hunch)
# ---------------------------------------------------------------------------

HUMANOID_PROPORTIONS = {
    # noun:       head_s, torso_w, torso_h, arm_l, leg_l, stance, hunch
    "Ogre":      (1.3,    1.4,     1.1,     1.1,   0.9,   1.3,    0.05),
    "Troll":     (1.0,    1.0,     1.3,     1.5,   1.0,   1.0,    0.15),
    "Knight":    (1.0,    1.2,     1.0,     1.0,   1.0,   1.1,    0.0),
    "Goblin":    (1.3,    0.7,     0.7,     0.8,   0.7,   0.9,    0.1),
    "Orc":       (1.1,    1.3,     1.0,     1.1,   1.0,   1.2,    0.05),
    "Dwarf":     (1.1,    1.3,     0.8,     0.8,   0.6,   1.3,    0.0),
    "Gnome":     (1.4,    0.7,     0.6,     0.7,   0.6,   0.8,    0.0),
    "Mage":      (1.0,    0.9,     1.1,     1.0,   1.0,   0.9,    0.0),
    "Skeleton":  (1.0,    0.7,     1.0,     1.0,   1.0,   0.9,    0.0),
    "Zombie":    (1.1,    1.0,     1.0,     1.1,   0.9,   1.1,    0.12),
    "Vampire":   (0.95,   0.9,     1.1,     1.0,   1.1,   0.9,    0.0),
    "Werewolf":  (1.15,   1.2,     1.0,     1.2,   1.0,   1.2,    0.1),
    "Wight":     (1.0,    0.8,     1.1,     1.1,   1.0,   0.9,    0.08),
    "Golem":     (0.9,    1.5,     1.2,     1.3,   0.9,   1.4,    0.0),
    "Brute":     (0.8,    1.5,     1.1,     1.3,   1.0,   1.4,    0.05),
    "Stalker":   (0.9,    0.8,     1.2,     1.2,   1.2,   0.8,    0.1),
    "Fiend":     (1.1,    1.1,     1.0,     1.1,   1.0,   1.1,    0.0),
    "Sentinel":  (0.9,    1.3,     1.3,     1.0,   1.1,   1.2,    0.0),
    "Colossus":  (0.8,    1.5,     1.4,     1.2,   1.1,   1.4,    0.0),
    "Revenant":  (1.0,    1.0,     1.0,     1.0,   1.0,   1.0,    0.05),
}

QUADRUPED_PROPORTIONS = {
    # noun:      body_rx, body_ry, leg_h, head_s, tail_len, ear_s
    "Wolf":     (1.0,     1.0,     1.0,   1.0,    1.2,      1.0),
    "Bear":     (1.3,     1.3,     0.8,   1.3,    0.3,      0.5),
    "Lion":     (1.1,     1.0,     1.0,   1.2,    1.4,      0.7),
    "Hound":    (0.9,     0.85,    1.1,   0.9,    0.8,      1.2),
    "Fox":      (0.8,     0.8,     0.9,   0.85,   1.5,      1.4),
    "Boar":     (1.1,     1.2,     0.7,   1.1,    0.3,      0.6),
    "Warg":     (1.2,     1.1,     1.0,   1.2,    1.0,      0.9),
}

WINGED_PROPORTIONS = {
    # noun:      body_s, wing_s, head_s, tail_l, neck_l, leg_l
    "Dragon":   (1.2,    1.3,    1.1,    1.3,    1.2,    1.0),
    "Griffin":  (1.1,    1.1,    1.0,    1.1,    0.9,    1.1),
    "Phoenix":  (0.9,    1.2,    0.9,    1.5,    1.0,    0.8),
    "Wyvern":   (1.1,    1.2,    1.0,    1.4,    1.1,    0.9),
    "Eagle":    (0.8,    1.3,    0.9,    0.8,    0.7,    0.9),
    "Harpy":    (0.9,    1.0,    1.1,    0.5,    0.8,    1.2),
    "Bat":      (0.7,    1.4,    0.8,    0.3,    0.5,    0.7),
    "Raven":    (0.7,    1.1,    0.9,    0.9,    0.7,    0.8),
}

# Size category -> overall scale factor for the creature within the canvas
SIZE_SCALES = {
    "Tiny": 0.45, "Small": 0.65, "Medium": 0.85,
    "Large": 1.0, "Huge": 1.1, "Colossal": 1.15,
}


def generate_creature_image(creature):
    """Generate a PIL Image for the creature using vector-style drawing."""
    primary = creature["primary_color"]
    secondary = creature["secondary_color"]
    accent = _accent_color(primary, secondary)
    dark = _darken(primary)
    light = _lighten(primary)
    primary_rgb = _hex_to_rgb(primary)
    secondary_rgb = _hex_to_rgb(secondary)
    archetype = creature["archetype"]
    noun = creature["noun"]
    ability = creature["ability"]
    size_cat = creature["size"]

    rng = _seeded_random(creature["name"] + "_img")
    size_scale = SIZE_SCALES.get(size_cat, 0.85)

    # Background color (light tint of secondary)
    bg = _lighten(secondary, 0.8)
    img = Image.new("RGBA", (CANVAS_W, CANVAS_H), bg + (255,))
    draw = ImageDraw.Draw(img, "RGBA")

    # Radial-ish background glow
    bg_center = _lighten(secondary, 0.9)
    for r_factor in [0.8, 0.6, 0.4]:
        rx = int(CANVAS_W * r_factor / 2)
        ry = int(CANVAS_H * r_factor / 2)
        alpha = int(30 * (1 - r_factor))
        draw.ellipse(
            [CANVAS_W // 2 - rx, CANVAS_H // 2 - ry - 40,
             CANVAS_W // 2 + rx, CANVAS_H // 2 + ry - 40],
            fill=bg_center + (alpha,)
        )

    # Ground shadow (scales with size)
    shadow_rx = int(55 * size_scale)
    _draw_ellipse(draw, 100, 235, shadow_rx, 10, fill=(0, 0, 0, 38))

    # Draw body based on archetype — now pass noun and size_scale
    if archetype == "humanoid":
        _draw_humanoid(draw, primary_rgb, secondary_rgb, dark, accent, light, rng, noun, size_scale)
    elif archetype == "quadruped":
        _draw_quadruped(draw, primary_rgb, secondary_rgb, dark, accent, light, rng, noun, size_scale)
    elif archetype == "winged":
        _draw_winged(draw, primary_rgb, secondary_rgb, dark, accent, light, rng, noun, size_scale)
    elif archetype == "serpentine":
        _draw_serpentine(draw, primary_rgb, secondary_rgb, dark, accent, light, rng, noun, size_scale)
    elif archetype == "amorphous":
        _draw_amorphous(draw, primary_rgb, secondary_rgb, dark, accent, light, rng, noun, size_scale)
    elif archetype == "arachnid":
        _draw_arachnid(draw, primary_rgb, secondary_rgb, dark, accent, light, rng, noun, size_scale)
    elif archetype == "small":
        _draw_small(draw, primary_rgb, secondary_rgb, dark, accent, light, rng, noun, size_scale)

    # Ability VFX overlay
    _draw_ability_vfx(draw, ability, accent, rng)

    # Convert to RGB for DOCX embedding
    final = Image.new("RGB", img.size, (255, 255, 255))
    final.paste(img, mask=img.split()[3])
    return final


# --- Archetype drawing functions (with noun-specific differentiation) ---

def _draw_humanoid(draw, primary, secondary, dark, accent, light, rng, noun, sz):
    props = HUMANOID_PROPORTIONS.get(noun, (1.0, 1.0, 1.0, 1.0, 1.0, 1.0, 0.0))
    head_s, tw, th, arm_l, leg_l, stance, hunch = props
    dark_dark = tuple(max(0, int(c * 0.8)) for c in dark)

    # RNG jitter for pose variation
    lean = rng.randint(-4, 4)
    arm_angle_l = rng.randint(-8, 5)
    arm_angle_r = rng.randint(-5, 8)

    cx = 100 + lean
    foot_y = 215
    leg_h = int(55 * leg_l * sz)
    torso_w = int(44 * tw * sz)
    torso_h = int(60 * th * sz)
    arm_len = int(50 * arm_l * sz)
    arm_w = max(8, int(14 * tw * 0.7 * sz))
    stance_w = int(22 * stance * sz)
    head_rx = int(18 * head_s * sz)
    head_ry = int(20 * head_s * sz)
    hunch_off = int(hunch * 40 * sz)

    torso_top = foot_y - leg_h - torso_h
    torso_bot = foot_y - leg_h
    torso_left = cx - torso_w // 2
    neck_y = torso_top - int(8 * sz)
    head_cy = neck_y - head_ry + int(4 * sz) - hunch_off

    # Legs
    ll_x = cx - stance_w // 2
    rl_x = cx + stance_w // 2
    leg_w = max(8, int(14 * sz))
    _draw_rect(draw, ll_x - leg_w // 2, torso_bot, leg_w, leg_h, fill=dark)
    _draw_rect(draw, rl_x - leg_w // 2, torso_bot, leg_w, leg_h, fill=dark)
    # Feet
    _draw_ellipse(draw, ll_x, foot_y, int(10 * sz), int(5 * sz), fill=dark_dark)
    _draw_ellipse(draw, rl_x, foot_y, int(10 * sz), int(5 * sz), fill=dark_dark)

    # Torso
    _draw_rect(draw, torso_left, torso_top, torso_w, torso_h, fill=primary)
    # Torso shading
    shade_w = max(4, torso_w - int(14 * sz))
    _draw_rect(draw, cx - shade_w // 2, torso_top + int(8 * sz), shade_w, torso_h - int(16 * sz),
               fill=dark + (50,))

    # --- Noun-specific torso decorations ---
    if noun == "Knight":
        # Chest plate cross
        _draw_rect(draw, cx - 2, torso_top + int(8*sz), 4, int(30*sz), fill=accent + (180,))
        _draw_rect(draw, cx - int(12*sz), torso_top + int(18*sz), int(24*sz), 4, fill=accent + (180,))
        # Visor slit (drawn later on head)
    elif noun == "Skeleton":
        # Rib lines
        for i in range(4):
            ry = torso_top + int((12 + i * 12) * sz)
            _draw_line(draw, cx - int(12*sz), ry, cx + int(12*sz), ry, fill=light + (160,), width=1.5)
    elif noun == "Mage":
        # Robe triangle extends below torso
        robe_w = int(torso_w * 1.3)
        _draw_polygon(draw, [(cx - robe_w//2, torso_bot + int(10*sz)),
                              (cx, torso_top + int(20*sz)),
                              (cx + robe_w//2, torso_bot + int(10*sz))],
                      fill=secondary + (100,))
    elif noun == "Golem":
        # Rune circles on torso
        for i in range(3):
            gy = torso_top + int((15 + i * 15) * sz)
            _draw_ellipse(draw, cx + rng.randint(-8, 8), gy, int(5*sz), int(5*sz),
                          fill=None, outline=accent + (200,), width=2)
    elif noun in ("Orc", "Ogre", "Brute"):
        # Battle scars
        for _ in range(rng.randint(2, 4)):
            sy = rng.randint(int(torso_top + 5*sz), int(torso_bot - 5*sz))
            sx = rng.randint(int(cx - torso_w//3), int(cx + torso_w//3))
            _draw_line(draw, sx, sy, sx + rng.randint(-8, 8), sy + rng.randint(5, 15),
                       fill=dark_dark + (150,), width=1.5)
    elif noun == "Vampire":
        # Cape / collar flare
        cape_w = int(torso_w * 0.8)
        _draw_polygon(draw, [(cx - cape_w, torso_top),
                              (cx - torso_w//2, torso_top + int(5*sz)),
                              (cx - torso_w//2, torso_top)],
                      fill=secondary + (180,))
        _draw_polygon(draw, [(cx + cape_w, torso_top),
                              (cx + torso_w//2, torso_top + int(5*sz)),
                              (cx + torso_w//2, torso_top)],
                      fill=secondary + (180,))
    elif noun == "Sentinel":
        # Glowing visor band on torso
        _draw_rect(draw, torso_left + 2, torso_top + int(5*sz), torso_w - 4, int(6*sz),
                   fill=accent + (120,))

    # Belt
    _draw_rect(draw, torso_left, torso_bot - int(6*sz), torso_w, int(6*sz), fill=accent)

    # Arms
    arm_lx = torso_left - arm_w + 2
    arm_rx = torso_left + torso_w - 2
    _draw_rect(draw, arm_lx, torso_top + int(2*sz) + arm_angle_l, arm_w, arm_len, fill=primary)
    _draw_rect(draw, arm_rx, torso_top + int(2*sz) + arm_angle_r, arm_w, arm_len, fill=primary)
    # Hands
    hand_r = max(4, int(6 * sz))
    _draw_ellipse(draw, arm_lx + arm_w//2, torso_top + int(2*sz) + arm_angle_l + arm_len,
                  hand_r, hand_r, fill=light)
    _draw_ellipse(draw, arm_rx + arm_w//2, torso_top + int(2*sz) + arm_angle_r + arm_len,
                  hand_r, hand_r, fill=light)

    # --- Noun-specific hand items ---
    if noun == "Knight":
        # Shield on left, sword on right
        shx = arm_lx - int(6*sz)
        shy = torso_top + int(15*sz) + arm_angle_l
        _draw_ellipse(draw, shx + int(8*sz), shy + int(12*sz), int(10*sz), int(14*sz), fill=accent + (200,))
        _draw_ellipse(draw, shx + int(8*sz), shy + int(12*sz), int(6*sz), int(9*sz), fill=secondary + (150,))
        # Sword
        swx = arm_rx + arm_w + int(3*sz)
        swy = torso_top - int(5*sz)
        _draw_rect(draw, swx, swy, int(3*sz), int(45*sz), fill=light + (220,))
        _draw_rect(draw, swx - int(4*sz), swy + int(8*sz), int(11*sz), int(3*sz), fill=accent)
    elif noun == "Mage":
        # Staff
        stx = arm_rx + arm_w + int(4*sz)
        _draw_line(draw, stx, torso_top - int(20*sz), stx, foot_y, fill=light + (200,), width=2)
        _draw_ellipse(draw, stx, torso_top - int(25*sz), int(6*sz), int(6*sz), fill=accent + (200,))

    # Shoulders
    shl_r = max(5, int(8 * sz * tw * 0.6))
    _draw_ellipse(draw, torso_left, torso_top + int(2*sz), shl_r, shl_r, fill=secondary)
    _draw_ellipse(draw, torso_left + torso_w, torso_top + int(2*sz), shl_r, shl_r, fill=secondary)

    # Neck
    neck_w = max(8, int(14 * sz * 0.8))
    _draw_rect(draw, cx - neck_w//2, neck_y, neck_w, int(12 * sz), fill=light)

    # Head
    _draw_ellipse(draw, cx + hunch_off, head_cy, head_rx, head_ry, fill=primary)

    # --- Noun-specific head decorations ---
    ey = head_cy - int(3 * sz * head_s)
    eye_r = max(2, int(4 * sz))
    eye_sep = int(8 * head_s * sz)

    if noun == "Skeleton":
        # Hollow sockets
        _draw_ellipse(draw, cx - eye_sep + hunch_off, ey, eye_r, eye_r, fill=(20, 20, 20))
        _draw_ellipse(draw, cx + eye_sep + hunch_off, ey, eye_r, eye_r, fill=(20, 20, 20))
        # Nose hole
        _draw_polygon(draw, [(cx + hunch_off, ey + int(6*sz)),
                              (cx - int(2*sz) + hunch_off, ey + int(10*sz)),
                              (cx + int(2*sz) + hunch_off, ey + int(10*sz))], fill=(20, 20, 20))
        # Teeth
        for tx in range(-2, 3):
            _draw_rect(draw, cx + tx * int(3*sz) + hunch_off - 1, head_cy + head_ry - int(7*sz),
                       int(2.5*sz), int(4*sz), fill=light)
    elif noun == "Knight":
        # Helmet visor slit
        _draw_rect(draw, cx - eye_sep - int(2*sz) + hunch_off, ey - int(1*sz),
                   eye_sep * 2 + int(4*sz), int(4*sz), fill=(20, 20, 20, 200))
        # Helmet plume
        _draw_polygon(draw, [(cx + hunch_off, head_cy - head_ry),
                              (cx - int(3*sz) + hunch_off, head_cy - head_ry - int(15*sz)),
                              (cx + int(3*sz) + hunch_off, head_cy - head_ry - int(12*sz))],
                      fill=accent + (200,))
    elif noun in ("Goblin",):
        # Big pointy ears
        ear_h = int(18 * sz)
        _draw_polygon(draw, [(cx - head_rx + hunch_off, head_cy - int(2*sz)),
                              (cx - head_rx - int(14*sz) + hunch_off, head_cy - ear_h),
                              (cx - head_rx + int(4*sz) + hunch_off, head_cy + int(2*sz))],
                      fill=secondary)
        _draw_polygon(draw, [(cx + head_rx + hunch_off, head_cy - int(2*sz)),
                              (cx + head_rx + int(14*sz) + hunch_off, head_cy - ear_h),
                              (cx + head_rx - int(4*sz) + hunch_off, head_cy + int(2*sz))],
                      fill=secondary)
        # Normal eyes
        _draw_ellipse(draw, cx - eye_sep + hunch_off, ey, eye_r, int(eye_r*0.7), fill=(255, 255, 0))
        _draw_ellipse(draw, cx + eye_sep + hunch_off, ey, eye_r, int(eye_r*0.7), fill=(255, 255, 0))
        _draw_ellipse(draw, cx - eye_sep + hunch_off, ey, int(eye_r*0.5), int(eye_r*0.5), fill=(20,20,20))
        _draw_ellipse(draw, cx + eye_sep + hunch_off, ey, int(eye_r*0.5), int(eye_r*0.5), fill=(20,20,20))
    elif noun in ("Orc", "Ogre"):
        # Tusks
        tusk_h = int(10 * sz)
        mouth_y = head_cy + int(head_ry * 0.5)
        _draw_polygon(draw, [(cx - int(5*sz) + hunch_off, mouth_y),
                              (cx - int(7*sz) + hunch_off, mouth_y - tusk_h),
                              (cx - int(3*sz) + hunch_off, mouth_y)], fill=(240, 230, 200))
        _draw_polygon(draw, [(cx + int(5*sz) + hunch_off, mouth_y),
                              (cx + int(7*sz) + hunch_off, mouth_y - tusk_h),
                              (cx + int(3*sz) + hunch_off, mouth_y)], fill=(240, 230, 200))
        # Brow ridge
        _draw_rect(draw, cx - head_rx + int(2*sz) + hunch_off, ey - int(5*sz),
                   head_rx * 2 - int(4*sz), int(4*sz), fill=dark + (80,))
        # Eyes
        _draw_ellipse(draw, cx - eye_sep + hunch_off, ey, eye_r, int(eye_r*0.6), fill=(255, 50, 0))
        _draw_ellipse(draw, cx + eye_sep + hunch_off, ey, eye_r, int(eye_r*0.6), fill=(255, 50, 0))
        _draw_ellipse(draw, cx - eye_sep + hunch_off, ey, int(eye_r*0.4), int(eye_r*0.4), fill=(20,20,20))
        _draw_ellipse(draw, cx + eye_sep + hunch_off, ey, int(eye_r*0.4), int(eye_r*0.4), fill=(20,20,20))
    elif noun == "Werewolf":
        # Pointed ears
        _draw_polygon(draw, [(cx - head_rx + int(3*sz) + hunch_off, head_cy - int(5*sz)),
                              (cx - int(6*sz) + hunch_off, head_cy - head_ry - int(12*sz)),
                              (cx - int(2*sz) + hunch_off, head_cy - int(2*sz))], fill=secondary)
        _draw_polygon(draw, [(cx + head_rx - int(3*sz) + hunch_off, head_cy - int(5*sz)),
                              (cx + int(6*sz) + hunch_off, head_cy - head_ry - int(12*sz)),
                              (cx + int(2*sz) + hunch_off, head_cy - int(2*sz))], fill=secondary)
        # Snout
        _draw_ellipse(draw, cx + hunch_off, head_cy + int(5*sz), int(8*sz), int(6*sz), fill=light)
        _draw_ellipse(draw, cx + hunch_off, head_cy + int(2*sz), int(3*sz), int(2*sz), fill=(30, 30, 30))
        # Eyes
        _draw_ellipse(draw, cx - eye_sep + hunch_off, ey, eye_r, int(eye_r*0.6), fill=(255, 200, 0))
        _draw_ellipse(draw, cx + eye_sep + hunch_off, ey, eye_r, int(eye_r*0.6), fill=(255, 200, 0))
        _draw_ellipse(draw, cx - eye_sep + hunch_off, ey, int(eye_r*0.4), int(eye_r*0.4), fill=(20,20,20))
        _draw_ellipse(draw, cx + eye_sep + hunch_off, ey, int(eye_r*0.4), int(eye_r*0.4), fill=(20,20,20))
    elif noun == "Vampire":
        # Slicked hair peak
        _draw_polygon(draw, [(cx - head_rx + hunch_off, head_cy - int(8*sz)),
                              (cx + hunch_off, head_cy - head_ry - int(8*sz)),
                              (cx + head_rx + hunch_off, head_cy - int(8*sz))], fill=(20, 20, 30))
        # Fangs
        mouth_y = head_cy + int(head_ry * 0.4)
        _draw_line(draw, cx - int(4*sz) + hunch_off, mouth_y,
                   cx - int(4*sz) + hunch_off, mouth_y + int(5*sz), fill=(240, 240, 240), width=1.5)
        _draw_line(draw, cx + int(4*sz) + hunch_off, mouth_y,
                   cx + int(4*sz) + hunch_off, mouth_y + int(5*sz), fill=(240, 240, 240), width=1.5)
        # Red eyes
        _draw_ellipse(draw, cx - eye_sep + hunch_off, ey, eye_r, int(eye_r*0.7), fill=(220, 0, 0))
        _draw_ellipse(draw, cx + eye_sep + hunch_off, ey, eye_r, int(eye_r*0.7), fill=(220, 0, 0))
        _draw_ellipse(draw, cx - eye_sep + hunch_off, ey, int(eye_r*0.4), int(eye_r*0.4), fill=(20,20,20))
        _draw_ellipse(draw, cx + eye_sep + hunch_off, ey, int(eye_r*0.4), int(eye_r*0.4), fill=(20,20,20))
    elif noun == "Zombie":
        # Messy, asymmetric eyes
        _draw_ellipse(draw, cx - eye_sep + hunch_off, ey - 1, eye_r, int(eye_r*0.8), fill=(200, 200, 150))
        _draw_ellipse(draw, cx + eye_sep + hunch_off + 1, ey + 1, int(eye_r*0.8), eye_r, fill=(200, 200, 150))
        _draw_ellipse(draw, cx - eye_sep + hunch_off, ey - 1, int(eye_r*0.3), int(eye_r*0.3), fill=(80,80,20))
        _draw_ellipse(draw, cx + eye_sep + hunch_off + 1, ey + 1, int(eye_r*0.3), int(eye_r*0.3), fill=(80,80,20))
        # Exposed jaw
        mouth_y = head_cy + int(head_ry * 0.5)
        _draw_line(draw, cx - int(6*sz) + hunch_off, mouth_y, cx + int(6*sz) + hunch_off, mouth_y,
                   fill=dark_dark, width=2)
    elif noun == "Fiend":
        # Horns
        horn_h = int(18 * sz)
        _draw_polygon(draw, [(cx - int(6*sz) + hunch_off, head_cy - head_ry + int(3*sz)),
                              (cx - int(14*sz) + hunch_off, head_cy - head_ry - horn_h),
                              (cx - int(3*sz) + hunch_off, head_cy - head_ry + int(6*sz))], fill=accent)
        _draw_polygon(draw, [(cx + int(6*sz) + hunch_off, head_cy - head_ry + int(3*sz)),
                              (cx + int(14*sz) + hunch_off, head_cy - head_ry - horn_h),
                              (cx + int(3*sz) + hunch_off, head_cy - head_ry + int(6*sz))], fill=accent)
        # Glowing eyes
        _draw_ellipse(draw, cx - eye_sep + hunch_off, ey, eye_r, eye_r, fill=(255, 100, 0))
        _draw_ellipse(draw, cx + eye_sep + hunch_off, ey, eye_r, eye_r, fill=(255, 100, 0))
    elif noun == "Dwarf":
        # Beard
        beard_top = head_cy + int(5*sz)
        for bi in range(5):
            bx = cx + hunch_off + (bi - 2) * int(4*sz)
            _draw_line(draw, bx, beard_top, bx + rng.randint(-2, 2), beard_top + int(15*sz),
                       fill=secondary + (200,), width=2)
        # Helmet
        _draw_ellipse(draw, cx + hunch_off, head_cy - int(5*sz), head_rx + int(3*sz), int(10*sz),
                      fill=accent + (180,))
        # Eyes
        _draw_ellipse(draw, cx - eye_sep + hunch_off, ey, eye_r, int(eye_r*0.7), fill=(255, 255, 255))
        _draw_ellipse(draw, cx + eye_sep + hunch_off, ey, eye_r, int(eye_r*0.7), fill=(255, 255, 255))
        _draw_ellipse(draw, cx - eye_sep + hunch_off, ey, int(eye_r*0.5), int(eye_r*0.5), fill=accent)
        _draw_ellipse(draw, cx + eye_sep + hunch_off, ey, int(eye_r*0.5), int(eye_r*0.5), fill=accent)
    elif noun == "Gnome":
        # Pointed hat
        _draw_polygon(draw, [(cx - head_rx + hunch_off, head_cy - int(5*sz)),
                              (cx + hunch_off, head_cy - head_ry - int(25*sz)),
                              (cx + head_rx + hunch_off, head_cy - int(5*sz))], fill=accent)
        # Big nose
        _draw_ellipse(draw, cx + hunch_off, head_cy + int(3*sz), int(5*sz), int(4*sz), fill=light)
        # Eyes
        _draw_ellipse(draw, cx - eye_sep + hunch_off, ey, eye_r + 1, eye_r + 1, fill=(255, 255, 255))
        _draw_ellipse(draw, cx + eye_sep + hunch_off, ey, eye_r + 1, eye_r + 1, fill=(255, 255, 255))
        _draw_ellipse(draw, cx - eye_sep + hunch_off, ey, int(eye_r*0.5), int(eye_r*0.5), fill=(30, 30, 30))
        _draw_ellipse(draw, cx + eye_sep + hunch_off, ey, int(eye_r*0.5), int(eye_r*0.5), fill=(30, 30, 30))
    else:
        # Default eyes (Wight, Stalker, Revenant, Colossus, Sentinel, Brute, etc.)
        eye_color = accent if noun in ("Wight", "Revenant", "Stalker", "Sentinel") else (255, 255, 255)
        _draw_ellipse(draw, cx - eye_sep + hunch_off, ey, eye_r, int(eye_r*0.7), fill=eye_color)
        _draw_ellipse(draw, cx + eye_sep + hunch_off, ey, eye_r, int(eye_r*0.7), fill=eye_color)
        if eye_color == (255, 255, 255):
            _draw_ellipse(draw, cx - eye_sep + hunch_off, ey, int(eye_r*0.5), int(eye_r*0.5), fill=accent)
            _draw_ellipse(draw, cx + eye_sep + hunch_off, ey, int(eye_r*0.5), int(eye_r*0.5), fill=accent)
        # Mouth
        mouth_y = head_cy + int(head_ry * 0.4)
        pts = _quad_bezier_points((cx - int(6*sz) + hunch_off, mouth_y),
                                   (cx + hunch_off, mouth_y + int(4*sz)),
                                   (cx + int(6*sz) + hunch_off, mouth_y))
        _draw_thick_curve(draw, pts, fill=dark, width=1.5)

    # --- Noun-specific extras ---
    if noun == "Werewolf":
        # Tail
        tx = cx + torso_w // 2 + int(5*sz)
        ty = torso_bot - int(10*sz)
        pts = _quad_bezier_points((tx, ty), (tx + int(15*sz), ty - int(20*sz)), (tx + int(10*sz), ty + int(5*sz)))
        _draw_thick_curve(draw, pts, fill=primary, width=3)
    elif noun == "Fiend":
        # Barbed tail
        tx = cx + torso_w // 2 + int(3*sz)
        ty = torso_bot - int(5*sz)
        pts = _quad_bezier_points((tx, ty), (tx + int(20*sz), ty + int(10*sz)), (tx + int(25*sz), ty - int(5*sz)))
        _draw_thick_curve(draw, pts, fill=dark, width=2.5)
        # Barb at tip
        tip = pts[-1]
        _draw_polygon(draw, [(tip[0], tip[1]), (tip[0] + 5, tip[1] - 8), (tip[0] + 8, tip[1] + 2)], fill=accent)


def _draw_quadruped(draw, primary, secondary, dark, accent, light, rng, noun, sz):
    props = QUADRUPED_PROPORTIONS.get(noun, (1.0, 1.0, 1.0, 1.0, 1.0, 1.0))
    body_rx_f, body_ry_f, leg_h_f, head_s, tail_f, ear_f = props
    dark_dark = tuple(max(0, int(c * 0.7)) for c in dark)

    by = 190
    body_rx = int(50 * body_rx_f * sz)
    body_ry = int(25 * body_ry_f * sz)
    body_cy = by - int(30 * sz)
    leg_h = int(35 * leg_h_f * sz)
    head_r = int(14 * head_s * sz)

    # RNG jitter
    body_lean = rng.randint(-3, 3)
    leg_spread = max(20, int(70 * sz * body_rx_f))

    # Body
    _draw_ellipse(draw, 100 + body_lean, body_cy, body_rx, body_ry, fill=primary)
    # Back highlight
    _draw_ellipse(draw, 100 + body_lean, body_cy - int(12*sz), int(body_rx*0.8), int(6*sz),
                  fill=secondary + (100,))

    # Legs at spread positions
    leg_positions = [100 - leg_spread//2, 100 - leg_spread//5,
                     100 + leg_spread//5, 100 + leg_spread//2]
    for i, lx in enumerate(leg_positions):
        off = rng.randint(-2, 2)
        lx += body_lean
        leg_w = max(6, int(10 * sz * body_rx_f * 0.5))
        _draw_rect(draw, lx - leg_w//2, body_cy + body_ry - int(5*sz) + off, leg_w, leg_h, fill=dark)
        _draw_ellipse(draw, lx, body_cy + body_ry - int(5*sz) + off + leg_h,
                      int(7*sz), int(4*sz), fill=dark_dark)

    # Tail
    tail_len = int(30 * tail_f * sz)
    tx_start = 100 + body_lean + body_rx - int(5*sz)
    ty_start = body_cy - int(5*sz)
    pts = _quad_bezier_points(
        (tx_start, ty_start),
        (tx_start + tail_len, ty_start - int(25*sz*tail_f)),
        (tx_start + int(tail_len*0.8), ty_start + int(10*sz)))
    tail_w = 5 if noun != "Fox" else 8
    _draw_thick_curve(draw, pts, fill=primary, width=tail_w * sz)
    if noun == "Fox":
        # Fluffy tail tip
        _draw_ellipse(draw, int(pts[-1][0]), int(pts[-1][1]), int(8*sz), int(6*sz), fill=light)
    elif noun == "Lion":
        # Tail tuft
        _draw_ellipse(draw, int(pts[-1][0]), int(pts[-1][1]), int(5*sz), int(5*sz), fill=dark)

    # Neck
    neck_w = max(8, int(16 * sz * head_s))
    neck_h = max(10, int(22 * sz))
    head_x = 100 + body_lean - body_rx + int(12*sz)
    neck_x = head_x + int(5*sz)
    _draw_rect(draw, neck_x - neck_w//2, body_cy - body_ry + int(2*sz) - int(8*sz), neck_w, neck_h, fill=primary)

    # Head
    head_y = body_cy - body_ry - int(15*sz)
    head_rx = int(18 * head_s * sz)
    head_ry_val = head_r
    _draw_ellipse(draw, head_x, head_y, head_rx, head_ry_val, fill=primary)

    # --- Noun-specific head features ---
    if noun == "Bear":
        # Round small ears
        _draw_ellipse(draw, head_x - int(10*sz), head_y - int(10*sz), int(5*sz*ear_f), int(5*sz*ear_f),
                      fill=secondary)
        _draw_ellipse(draw, head_x + int(10*sz), head_y - int(10*sz), int(5*sz*ear_f), int(5*sz*ear_f),
                      fill=secondary)
        # Snout
        _draw_ellipse(draw, head_x - int(8*sz), head_y + int(2*sz), int(10*sz), int(8*sz), fill=light)
        _draw_ellipse(draw, head_x - int(12*sz), head_y + int(1*sz), int(3*sz), int(3*sz), fill=(30, 30, 30))
    elif noun == "Lion":
        # Mane
        mane_r = int(head_rx * 1.6)
        _draw_ellipse(draw, head_x, head_y, mane_r, int(head_ry_val * 1.5), fill=secondary + (150,))
        # Re-draw head over mane
        _draw_ellipse(draw, head_x, head_y, head_rx, head_ry_val, fill=primary)
        # Snout
        _draw_ellipse(draw, head_x - int(10*sz), head_y + int(2*sz), int(8*sz), int(6*sz), fill=light)
        _draw_ellipse(draw, head_x - int(14*sz), head_y + int(1*sz), int(3*sz), int(2*sz), fill=(30, 30, 30))
    elif noun == "Boar":
        # Tusks
        _draw_polygon(draw, [(head_x - int(12*sz), head_y + int(2*sz)),
                              (head_x - int(16*sz), head_y - int(8*sz)),
                              (head_x - int(10*sz), head_y)], fill=(240, 230, 200))
        # Bristles on back
        for bx in range(3):
            _draw_line(draw, 100 + body_lean - int(15*sz) + bx * int(10*sz), body_cy - body_ry,
                       100 + body_lean - int(15*sz) + bx * int(10*sz), body_cy - body_ry - int(8*sz),
                       fill=dark + (150,), width=2)
        # Snout
        _draw_ellipse(draw, head_x - int(10*sz), head_y + int(1*sz), int(10*sz), int(7*sz), fill=light)
    else:
        # Wolf/Hound/Fox/Warg — pointy snout + triangle ears
        # Ears
        ear_h = int(12 * ear_f * sz)
        _draw_polygon(draw, [(head_x - int(6*sz), head_y - head_ry_val + int(2*sz)),
                              (head_x - int(10*sz), head_y - head_ry_val - ear_h),
                              (head_x - int(2*sz), head_y - head_ry_val + int(4*sz))], fill=secondary)
        _draw_polygon(draw, [(head_x + int(6*sz), head_y - head_ry_val + int(2*sz)),
                              (head_x + int(10*sz), head_y - head_ry_val - ear_h),
                              (head_x + int(2*sz), head_y - head_ry_val + int(4*sz))], fill=secondary)
        # Snout
        _draw_ellipse(draw, head_x - int(12*sz), head_y + int(2*sz), int(10*sz), int(7*sz), fill=light)
        _draw_ellipse(draw, head_x - int(17*sz), head_y + int(1*sz), int(3*sz), int(2.5*sz), fill=(30, 30, 30))

    # Eye
    eye_x = head_x - int(3*sz)
    eye_y = head_y - int(3*sz)
    _draw_ellipse(draw, eye_x, eye_y, int(3*sz), int(3*sz), fill=(255, 255, 255))
    eye_glow = (255, 200, 0) if noun == "Warg" else accent
    _draw_ellipse(draw, eye_x, eye_y, int(1.5*sz), int(1.5*sz), fill=eye_glow)

    # --- Noun-specific body extras ---
    if noun == "Warg":
        # Scars on body
        for _ in range(rng.randint(2, 4)):
            sx = 100 + body_lean + rng.randint(-int(body_rx*0.6), int(body_rx*0.6))
            sy = body_cy + rng.randint(-int(body_ry*0.5), int(body_ry*0.5))
            _draw_line(draw, sx, sy, sx + rng.randint(-6, 6), sy + rng.randint(3, 10),
                       fill=dark_dark + (120,), width=1.5)


def _draw_winged(draw, primary, secondary, dark, accent, light, rng, noun, sz):
    props = WINGED_PROPORTIONS.get(noun, (1.0, 1.0, 1.0, 1.0, 1.0, 1.0))
    body_s, wing_s, head_s_f, tail_f, neck_f, leg_f = props

    by = 185
    body_rx = int(35 * body_s * sz)
    body_ry = int(28 * body_s * sz)
    body_cy = by - int(35 * sz)
    wing_span = int(85 * wing_s * sz)
    head_r = int(14 * head_s_f * sz)
    neck_h = int(18 * neck_f * sz)
    leg_h = int(28 * leg_f * sz)

    # RNG variation
    wing_droop = rng.randint(-5, 5)

    # Body
    _draw_ellipse(draw, 100, body_cy, body_rx, body_ry, fill=primary)
    # Belly highlight
    _draw_ellipse(draw, 100, body_cy + int(10*sz), int(body_rx*0.6), int(body_ry*0.4), fill=light + (80,))

    # Wings — left
    wing_top = body_cy - int(25 * sz) + wing_droop
    _draw_polygon(draw, [(100 - int(body_rx*0.6), body_cy - int(body_ry*0.5)),
                          (100 - wing_span, wing_top - int(30*sz)),
                          (100 - int(wing_span*0.7), body_cy - int(5*sz)),
                          (100 - int(wing_span*0.4), wing_top),
                          (100 - int(body_rx*0.3), body_cy + int(body_ry*0.3))],
                  fill=secondary + (200,))
    # Wing membrane lines
    for i in range(3):
        frac = 0.3 + i * 0.25
        wx = 100 - int(body_rx*0.5) - int(wing_span * frac * 0.5)
        wy_top = wing_top - int(15*sz) + int(i * 10 * sz)
        _draw_line(draw, 100 - int(body_rx*0.5), body_cy - int(body_ry*0.3), wx, wy_top,
                   fill=primary + (100,), width=1)

    # Wings — right
    _draw_polygon(draw, [(100 + int(body_rx*0.6), body_cy - int(body_ry*0.5)),
                          (100 + wing_span, wing_top - int(30*sz)),
                          (100 + int(wing_span*0.7), body_cy - int(5*sz)),
                          (100 + int(wing_span*0.4), wing_top),
                          (100 + int(body_rx*0.3), body_cy + int(body_ry*0.3))],
                  fill=secondary + (200,))
    for i in range(3):
        frac = 0.3 + i * 0.25
        wx = 100 + int(body_rx*0.5) + int(wing_span * frac * 0.5)
        wy_top = wing_top - int(15*sz) + int(i * 10 * sz)
        _draw_line(draw, 100 + int(body_rx*0.5), body_cy - int(body_ry*0.3), wx, wy_top,
                   fill=primary + (100,), width=1)

    # Legs
    _draw_rect(draw, 100 - int(12*sz), body_cy + body_ry - int(5*sz), int(10*sz), leg_h, fill=dark)
    _draw_rect(draw, 100 + int(2*sz), body_cy + body_ry - int(5*sz), int(10*sz), leg_h, fill=dark)
    # Claws
    foot_y = body_cy + body_ry - int(5*sz) + leg_h
    for dx in [-12, -6, 2, 8]:
        cx_ = 100 + int(dx * sz)
        _draw_polygon(draw, [(cx_, foot_y), (cx_ + int(2*sz), foot_y + int(6*sz)),
                              (cx_ + int(4*sz), foot_y)], fill=accent)

    # Neck
    neck_w = max(8, int(14 * sz))
    neck_top = body_cy - body_ry + int(5*sz)
    _draw_rect(draw, 100 - neck_w//2, neck_top - neck_h, neck_w, neck_h, fill=primary)

    # Head
    head_y = neck_top - neck_h - head_r + int(5*sz)
    _draw_ellipse(draw, 100, head_y, int(16 * head_s_f * sz), head_r, fill=primary)

    # --- Noun-specific head/features ---
    if noun in ("Dragon", "Wyvern"):
        # Horns
        horn_h = int(20 * sz)
        _draw_polygon(draw, [(100 - int(8*sz), head_y - head_r + int(3*sz)),
                              (100 - int(14*sz), head_y - head_r - horn_h),
                              (100 - int(4*sz), head_y - head_r + int(6*sz))], fill=accent)
        _draw_polygon(draw, [(100 + int(8*sz), head_y - head_r + int(3*sz)),
                              (100 + int(14*sz), head_y - head_r - horn_h),
                              (100 + int(4*sz), head_y - head_r + int(6*sz))], fill=accent)
        # Snout
        _draw_ellipse(draw, 100, head_y + int(4*sz), int(10*sz), int(6*sz), fill=light)
        # Nostrils
        _draw_ellipse(draw, 100 - int(3*sz), head_y + int(5*sz), int(2*sz), int(1.5*sz), fill=dark)
        _draw_ellipse(draw, 100 + int(3*sz), head_y + int(5*sz), int(2*sz), int(1.5*sz), fill=dark)
    elif noun == "Phoenix":
        # Flame crest
        for dx in range(-10, 14, 4):
            fh = rng.randint(10, 25)
            _draw_polygon(draw, [(100 + int(dx*sz), head_y - head_r),
                                  (100 + int((dx-2)*sz), head_y - head_r - int(fh*sz)),
                                  (100 + int((dx+2)*sz), head_y - head_r - int(fh*sz*0.7))],
                          fill=(255, rng.randint(80, 200), 0, rng.randint(150, 230)))
        # Beak
        _draw_polygon(draw, [(100, head_y + int(3*sz)),
                              (100 - int(4*sz), head_y + int(8*sz)),
                              (100 + int(4*sz), head_y + int(8*sz))], fill=(255, 180, 0))
    elif noun == "Griffin":
        # Beak
        _draw_polygon(draw, [(100 - int(2*sz), head_y + int(2*sz)),
                              (100, head_y + int(10*sz)),
                              (100 + int(2*sz), head_y + int(2*sz))], fill=(220, 180, 50))
        # Ear tufts
        _draw_polygon(draw, [(100 - int(8*sz), head_y - head_r),
                              (100 - int(10*sz), head_y - head_r - int(10*sz)),
                              (100 - int(5*sz), head_y - head_r + int(2*sz))], fill=secondary)
        _draw_polygon(draw, [(100 + int(8*sz), head_y - head_r),
                              (100 + int(10*sz), head_y - head_r - int(10*sz)),
                              (100 + int(5*sz), head_y - head_r + int(2*sz))], fill=secondary)
    elif noun == "Harpy":
        # More humanoid face
        _draw_ellipse(draw, 100, head_y + int(3*sz), int(5*sz), int(3*sz), fill=light)
    elif noun in ("Eagle", "Raven"):
        # Hooked beak
        _draw_polygon(draw, [(100, head_y + head_r - int(2*sz)),
                              (100, head_y + head_r + int(6*sz)),
                              (100 - int(3*sz), head_y + head_r + int(2*sz))],
                      fill=(200, 170, 30) if noun == "Eagle" else (40, 40, 40))
    elif noun == "Bat":
        # Big ears
        _draw_polygon(draw, [(100 - int(8*sz), head_y - int(2*sz)),
                              (100 - int(14*sz), head_y - head_r - int(18*sz)),
                              (100 - int(2*sz), head_y - int(4*sz))], fill=primary)
        _draw_polygon(draw, [(100 + int(8*sz), head_y - int(2*sz)),
                              (100 + int(14*sz), head_y - head_r - int(18*sz)),
                              (100 + int(2*sz), head_y - int(4*sz))], fill=primary)

    # Eyes
    eye_sep = int(7 * head_s_f * sz)
    eye_r = max(2, int(3 * sz))
    eye_color = (255, 221, 0) if noun in ("Dragon", "Wyvern", "Phoenix") else (255, 255, 255)
    _draw_ellipse(draw, 100 - eye_sep, head_y - int(2*sz), eye_r, eye_r, fill=eye_color)
    _draw_ellipse(draw, 100 + eye_sep, head_y - int(2*sz), eye_r, eye_r, fill=eye_color)
    _draw_ellipse(draw, 100 - eye_sep, head_y - int(2*sz), int(eye_r*0.5), int(eye_r*0.5), fill=dark)
    _draw_ellipse(draw, 100 + eye_sep, head_y - int(2*sz), int(eye_r*0.5), int(eye_r*0.5), fill=dark)

    # Tail
    tail_len = int(40 * tail_f * sz)
    pts = _bezier_points((100, body_cy + body_ry - int(5*sz)),
                          (100 + int(20*sz), body_cy + body_ry + int(5*sz)),
                          (100 + int(tail_len*0.7), body_cy),
                          (100 + tail_len, body_cy + int(10*sz)))
    _draw_thick_curve(draw, pts, fill=primary, width=max(3, int(5 * sz)))
    if noun == "Phoenix":
        # Flaming tail tip
        for fi in range(3):
            tp = pts[-1 - fi * 2] if fi * 2 < len(pts) else pts[-1]
            _draw_polygon(draw, [(tp[0], tp[1]), (tp[0] - 4, tp[1] - rng.randint(8, 15)),
                                  (tp[0] + 4, tp[1] - rng.randint(6, 12))],
                          fill=(255, rng.randint(100, 200), 0, 180))


def _draw_serpentine(draw, primary, secondary, dark, accent, light, rng, noun, sz):
    # Vary the S-curve control points per creature
    cx_off = rng.randint(-10, 10)
    coil_w = rng.randint(30, 60)
    body_w = max(10, int(18 * sz))

    # Build a unique S-curve path
    p0 = (50 + cx_off, int(210 * sz + (1 - sz) * 50))
    p1 = (50 + cx_off - coil_w, int(160 * sz + (1 - sz) * 40))
    p2 = (100 + cx_off + coil_w, int(120 * sz + (1 - sz) * 30))
    p3 = (100 + cx_off, int(80 * sz + (1 - sz) * 20))
    pts = _bezier_points(p0, p1, p2, p3)

    p4 = (100 + cx_off - rng.randint(10, 30), int(50 * sz + (1 - sz) * 15))
    p5 = (100 + cx_off + rng.randint(15, 35), int(35 * sz + (1 - sz) * 10))
    p6 = (100 + cx_off + rng.randint(-5, 5), int(42 * sz + (1 - sz) * 10))
    pts2 = _bezier_points(p3, p4, p5, p6)
    all_pts = pts + pts2[1:]

    _draw_thick_curve(draw, all_pts, fill=primary, width=body_w)
    _draw_thick_curve(draw, all_pts, fill=light + (80,), width=max(4, body_w // 2))

    # Scale pattern
    num_scales = rng.randint(6, 14)
    for i in range(num_scales):
        idx = int(i * len(all_pts) / num_scales)
        if idx >= len(all_pts):
            break
        sx, sy = all_pts[idx]
        ds = rng.randint(2, 5)
        _draw_polygon(draw, [(sx, sy - ds), (sx + ds, sy), (sx, sy + ds), (sx - ds, sy)],
                      fill=secondary + (rng.randint(100, 200),))

    # Head
    hx, hy = all_pts[-1]
    head_rx = int(14 * sz)
    head_ry = int(11 * sz)
    _draw_ellipse(draw, hx, hy, head_rx, head_ry, fill=primary)

    if noun == "Hydra":
        # Extra heads branching from neck area
        for hi in range(rng.randint(2, 4)):
            branch_idx = max(0, len(all_pts) - rng.randint(5, 12))
            bx, by_ = all_pts[branch_idx]
            angle = rng.uniform(-1.2, 1.2)
            neck_len = rng.randint(15, 30)
            ehx = bx + int(math.cos(angle) * neck_len * sz)
            ehy = by_ - int(abs(math.sin(angle)) * neck_len * sz) - int(8 * sz)
            _draw_line(draw, bx, by_, ehx, ehy, fill=primary, width=max(4, int(8 * sz)))
            _draw_ellipse(draw, ehx, ehy, int(8*sz), int(6*sz), fill=primary)
            # Mini eyes
            _draw_ellipse(draw, ehx - int(2*sz), ehy - int(1*sz), int(2*sz), int(2*sz), fill=(255, 221, 0))
            _draw_ellipse(draw, ehx + int(2*sz), ehy - int(1*sz), int(2*sz), int(2*sz), fill=(255, 221, 0))

    # Main head eyes
    _draw_ellipse(draw, hx - int(5*sz), hy - int(3*sz), int(3*sz), int(3*sz), fill=(255, 221, 0))
    _draw_ellipse(draw, hx + int(5*sz), hy - int(3*sz), int(3*sz), int(3*sz), fill=(255, 221, 0))
    _draw_ellipse(draw, hx - int(5*sz), hy - int(3*sz), int(1.5*sz), int(1.5*sz), fill=dark)
    _draw_ellipse(draw, hx + int(5*sz), hy - int(3*sz), int(1.5*sz), int(1.5*sz), fill=dark)
    # Tongue
    _draw_line(draw, hx, hy + head_ry - int(2*sz), hx - int(2*sz), hy + head_ry + int(8*sz),
               fill=(204, 51, 51), width=1.5)
    _draw_line(draw, hx - int(2*sz), hy + head_ry + int(8*sz), hx - int(5*sz), hy + head_ry + int(12*sz),
               fill=(204, 51, 51), width=1)
    _draw_line(draw, hx - int(2*sz), hy + head_ry + int(8*sz), hx + int(1*sz), hy + head_ry + int(12*sz),
               fill=(204, 51, 51), width=1)

    # Tail taper
    tx, ty = all_pts[0]
    pts_tail = _quad_bezier_points((tx, ty), (tx - int(10*sz), ty + int(15*sz)), (tx - int(15*sz), ty + int(25*sz)))
    _draw_thick_curve(draw, pts_tail, fill=dark, width=max(4, int(8 * sz)))


def _draw_amorphous(draw, primary, secondary, dark, accent, light, rng, noun, sz):
    cx = 100 + rng.randint(-10, 10)
    cy = 140 + rng.randint(-10, 10)

    # Number and spread of blobs varies by noun
    if noun == "Slime":
        num_blobs = rng.randint(3, 5)
        spread = 15
        tendril_count = rng.randint(1, 2)
        num_eyes = 2
    elif noun in ("Wraith", "Shade", "Specter"):
        num_blobs = rng.randint(4, 7)
        spread = 25
        tendril_count = rng.randint(4, 8)
        num_eyes = rng.randint(1, 2)
    elif noun == "Elemental":
        num_blobs = rng.randint(6, 10)
        spread = 30
        tendril_count = rng.randint(5, 8)
        num_eyes = rng.randint(1, 3)
    elif noun == "Echo":
        num_blobs = rng.randint(3, 5)
        spread = 20
        tendril_count = rng.randint(2, 4)
        num_eyes = rng.randint(2, 4)
    else:  # Blight
        num_blobs = rng.randint(5, 8)
        spread = 25
        tendril_count = rng.randint(4, 7)
        num_eyes = rng.randint(1, 2)

    spread = int(spread * sz)

    # Blobs
    for i in range(num_blobs):
        rx = rng.randint(int(15*sz), int(50*sz))
        ry = rng.randint(int(12*sz), int(45*sz))
        dx = rng.randint(-spread, spread)
        dy = rng.randint(-spread, spread)
        alpha = rng.randint(40, 160)
        color = primary if rng.random() < 0.6 else secondary
        _draw_ellipse(draw, cx + dx, cy + dy, rx, ry, fill=color + (alpha,))

    # Core glow
    _draw_ellipse(draw, cx, cy, int(30*sz), int(25*sz), fill=light + (80,))

    # Noun-specific core shapes
    if noun == "Slime":
        # Visible objects trapped inside
        for _ in range(rng.randint(2, 4)):
            ox = cx + rng.randint(int(-15*sz), int(15*sz))
            oy = cy + rng.randint(int(-10*sz), int(15*sz))
            _draw_rect(draw, ox, oy, rng.randint(3, 8), rng.randint(3, 8), fill=dark + (100,))
    elif noun in ("Wraith", "Shade"):
        # Hooded shape silhouette
        hood_w = int(20 * sz)
        hood_h = int(25 * sz)
        _draw_polygon(draw, [(cx - hood_w, cy + int(10*sz)),
                              (cx, cy - hood_h),
                              (cx + hood_w, cy + int(10*sz))], fill=dark + (100,))
        # Spectral chains (Wraith)
        if noun == "Wraith":
            for _ in range(rng.randint(2, 3)):
                ch_x = cx + rng.randint(int(-20*sz), int(20*sz))
                ch_y = cy + rng.randint(0, int(20*sz))
                for ci in range(rng.randint(3, 6)):
                    _draw_ellipse(draw, ch_x + ci * int(4*sz), ch_y + ci * int(3*sz),
                                  int(3*sz), int(2*sz), fill=accent + (100,), outline=accent + (150,))
    elif noun == "Elemental":
        # Spinning core particles
        for _ in range(rng.randint(5, 10)):
            angle = rng.uniform(0, 2 * math.pi)
            dist = rng.randint(int(5*sz), int(20*sz))
            px = cx + int(math.cos(angle) * dist)
            py = cy + int(math.sin(angle) * dist)
            _draw_ellipse(draw, px, py, rng.randint(2, 5), rng.randint(2, 5), fill=accent + (180,))
    elif noun == "Blight":
        # Twisted vine/root shapes
        for _ in range(rng.randint(3, 5)):
            angle = rng.uniform(0, 2 * math.pi)
            length = rng.randint(int(20*sz), int(40*sz))
            ex = cx + int(math.cos(angle) * length)
            ey = cy + int(math.sin(angle) * length)
            mx = cx + int(math.cos(angle) * length * 0.4) + rng.randint(-8, 8)
            my = cy + int(math.sin(angle) * length * 0.4) + rng.randint(-8, 8)
            pts = _quad_bezier_points((cx, cy), (mx, my), (ex, ey))
            _draw_thick_curve(draw, pts, fill=(50, 80, 20, 180), width=rng.randint(2, 4))
            # Leaf/growth at tip
            _draw_ellipse(draw, ex, ey, rng.randint(3, 6), rng.randint(2, 4), fill=(80, 130, 30, 150))

    # Tendrils
    for _ in range(tendril_count):
        angle = rng.uniform(0, 2 * math.pi)
        length = rng.randint(int(25*sz), int(55*sz))
        ex = cx + int(math.cos(angle) * length)
        ey = cy + int(math.sin(angle) * length)
        mx = cx + int(math.cos(angle) * length * 0.5) + rng.randint(-10, 10)
        my = cy + int(math.sin(angle) * length * 0.5) + rng.randint(-10, 10)
        pts = _quad_bezier_points((cx, cy), (mx, my), (ex, ey))
        _draw_thick_curve(draw, pts, fill=primary + (rng.randint(80, 160),), width=rng.randint(2, 5))

    # Eyes
    for _ in range(num_eyes):
        ex = cx + rng.randint(int(-15*sz), int(15*sz))
        ey = cy + rng.randint(int(-20*sz), int(5*sz))
        er = max(3, int(5 * sz))
        _draw_ellipse(draw, ex, ey, er, er, fill=(255, 255, 255, 230))
        iris_color = accent if noun != "Shade" else (200, 0, 0)
        _draw_ellipse(draw, ex, ey, int(er*0.5), int(er*0.5), fill=iris_color)


def _draw_arachnid(draw, primary, secondary, dark, accent, light, rng, noun, sz):
    cx = 100 + rng.randint(-5, 5)
    cy = 140 + rng.randint(-5, 5)

    abd_rx = int(28 * sz)
    abd_ry = int(22 * sz)
    ceph_rx = int(18 * sz)
    ceph_ry = int(15 * sz)
    leg_reach = int(35 * sz)

    # Abdomen
    _draw_ellipse(draw, cx + int(5*sz), cy + int(15*sz), abd_rx, abd_ry, fill=primary)
    # Abdomen markings (unique pattern per creature)
    mark_type = rng.choice(["hourglass", "dots", "stripes", "diamond"])
    if mark_type == "hourglass":
        _draw_polygon(draw, [(cx + int(5*sz), cy + int(5*sz)),
                              (cx + int(5*sz) - int(6*sz), cy + int(15*sz)),
                              (cx + int(5*sz), cy + int(25*sz)),
                              (cx + int(5*sz) + int(6*sz), cy + int(15*sz))], fill=accent + (180,))
    elif mark_type == "dots":
        for _ in range(rng.randint(3, 7)):
            dx = rng.randint(-int(abd_rx*0.5), int(abd_rx*0.5))
            dy = rng.randint(-int(abd_ry*0.5), int(abd_ry*0.5))
            _draw_ellipse(draw, cx + int(5*sz) + dx, cy + int(15*sz) + dy,
                          rng.randint(2, 5), rng.randint(2, 5), fill=accent + (150,))
    elif mark_type == "stripes":
        for si in range(-2, 3):
            sy = cy + int(15*sz) + si * int(5*sz)
            _draw_line(draw, cx + int(5*sz) - int(abd_rx*0.6), sy,
                       cx + int(5*sz) + int(abd_rx*0.6), sy, fill=accent + (120,), width=2)
    else:  # diamond
        _draw_polygon(draw, [(cx + int(5*sz), cy + int(5*sz)),
                              (cx + int(5*sz) + int(8*sz), cy + int(15*sz)),
                              (cx + int(5*sz), cy + int(25*sz)),
                              (cx + int(5*sz) - int(8*sz), cy + int(15*sz))], fill=accent + (120,))

    # Cephalothorax
    _draw_ellipse(draw, cx - int(8*sz), cy - int(10*sz), ceph_rx, ceph_ry, fill=secondary)

    # Eyes (variable cluster)
    num_eyes = rng.randint(4, 8)
    for _ in range(num_eyes):
        edx = rng.randint(-int(8*sz), int(8*sz))
        edy = rng.randint(-int(8*sz), int(2*sz))
        er = max(1, rng.randint(1, int(3*sz)))
        _draw_ellipse(draw, cx - int(8*sz) + edx, cy - int(16*sz) + edy, er, er,
                      fill=(255, rng.randint(0, 50), 0, 230))

    # Fangs
    fang_len = int(12 * sz)
    _draw_line(draw, cx - int(12*sz), cy - int(4*sz), cx - int(14*sz), cy - int(4*sz) + fang_len,
               fill=accent, width=max(2, int(2.5*sz)))
    _draw_line(draw, cx - int(4*sz), cy - int(4*sz), cx - int(2*sz), cy - int(4*sz) + fang_len,
               fill=accent, width=max(2, int(2.5*sz)))

    # 8 legs with RNG-driven angles and lengths
    for sign in [-1, 1]:
        for i in range(4):
            angle = math.pi * (0.5 + sign * (0.15 + i * 0.18 + rng.uniform(-0.05, 0.05)))
            mid_x = cx + int(math.cos(angle) * leg_reach * (0.8 + rng.uniform(0, 0.4)))
            mid_y = cy - int(5*sz) + int(math.sin(angle) * int(20*sz)) - int(15*sz)
            end_x = mid_x + sign * rng.randint(int(5*sz), int(18*sz))
            end_y = cy + int(45*sz) + rng.randint(int(-8*sz), int(8*sz))
            start_x = cx - int(5*sz) * sign
            start_y = cy - int(8*sz) + i * int(5*sz)
            _draw_line(draw, start_x, start_y, mid_x, mid_y, fill=dark, width=max(2, int(2.5*sz)))
            _draw_line(draw, mid_x, mid_y, end_x, end_y, fill=dark, width=max(2, int(2.5*sz)))


def _draw_small(draw, primary, secondary, dark, accent, light, rng, noun, sz):
    cx = 100 + rng.randint(-5, 5)
    cy = 160

    body_rx = int(18 * sz)
    body_ry = int(22 * sz)
    head_r = int(14 * sz)

    # Body
    _draw_ellipse(draw, cx, cy, body_rx, body_ry, fill=primary)
    # Belly
    _draw_ellipse(draw, cx, cy + int(5*sz), int(12*sz), int(12*sz), fill=light + (75,))

    # Legs
    leg_w = max(4, int(8 * sz))
    leg_h = max(8, int(18 * sz))
    _draw_rect(draw, cx - int(10*sz), cy + body_ry - int(4*sz), leg_w, leg_h, fill=dark)
    _draw_rect(draw, cx + int(2*sz), cy + body_ry - int(4*sz), leg_w, leg_h, fill=dark)

    # Arms
    arm_w = max(4, int(8 * sz))
    arm_h = max(8, int(22 * sz))
    _draw_rect(draw, cx - body_rx - arm_w + int(2*sz), cy - int(8*sz), arm_w, arm_h, fill=primary)
    _draw_rect(draw, cx + body_rx - int(2*sz), cy - int(8*sz), arm_w, arm_h, fill=primary)

    # Head
    head_cy = cy - body_ry - head_r + int(8*sz)
    _draw_ellipse(draw, cx, head_cy, head_r, head_r, fill=primary)

    # --- Noun-specific features ---
    eye_sep = int(5 * sz)
    eye_r = max(2, int(4.5 * sz))

    if noun == "Imp":
        # Small horns
        horn_h = int(10 * sz)
        _draw_polygon(draw, [(cx - int(6*sz), head_cy - head_r + int(2*sz)),
                              (cx - int(10*sz), head_cy - head_r - horn_h),
                              (cx - int(3*sz), head_cy - head_r + int(4*sz))], fill=accent)
        _draw_polygon(draw, [(cx + int(6*sz), head_cy - head_r + int(2*sz)),
                              (cx + int(10*sz), head_cy - head_r - horn_h),
                              (cx + int(3*sz), head_cy - head_r + int(4*sz))], fill=accent)
        # Bat wings (small)
        _draw_polygon(draw, [(cx - body_rx, cy - int(10*sz)),
                              (cx - body_rx - int(18*sz), cy - int(25*sz)),
                              (cx - body_rx - int(10*sz), cy + int(2*sz))], fill=secondary + (180,))
        _draw_polygon(draw, [(cx + body_rx, cy - int(10*sz)),
                              (cx + body_rx + int(18*sz), cy - int(25*sz)),
                              (cx + body_rx + int(10*sz), cy + int(2*sz))], fill=secondary + (180,))
        # Barbed tail
        pts = _quad_bezier_points((cx + body_rx - int(3*sz), cy + int(5*sz)),
                                   (cx + body_rx + int(20*sz), cy + int(15*sz)),
                                   (cx + body_rx + int(15*sz), cy - int(5*sz)))
        _draw_thick_curve(draw, pts, fill=dark, width=2)
        _draw_polygon(draw, [(pts[-1][0], pts[-1][1]),
                              (pts[-1][0] + 3, pts[-1][1] - 6),
                              (pts[-1][0] + 6, pts[-1][1] + 1)], fill=accent)
        # Mischievous eyes
        _draw_ellipse(draw, cx - eye_sep, head_cy - int(2*sz), eye_r, int(eye_r*0.7), fill=(255, 200, 0))
        _draw_ellipse(draw, cx + eye_sep, head_cy - int(2*sz), eye_r, int(eye_r*0.7), fill=(255, 200, 0))
        _draw_ellipse(draw, cx - eye_sep + 1, head_cy - int(2*sz), int(eye_r*0.4), int(eye_r*0.4), fill=(20,20,20))
        _draw_ellipse(draw, cx + eye_sep + 1, head_cy - int(2*sz), int(eye_r*0.4), int(eye_r*0.4), fill=(20,20,20))
        # Grin
        pts_m = _quad_bezier_points((cx - int(5*sz), head_cy + int(4*sz)),
                                     (cx, head_cy + int(8*sz)),
                                     (cx + int(5*sz), head_cy + int(4*sz)))
        _draw_thick_curve(draw, pts_m, fill=dark, width=1.5)
    elif noun == "Sprite":
        # Delicate insect wings
        wing_h = int(22 * sz)
        _draw_ellipse(draw, cx - body_rx - int(5*sz), cy - int(15*sz),
                      int(12*sz), wing_h, fill=secondary + (100,), outline=accent + (150,))
        _draw_ellipse(draw, cx + body_rx + int(5*sz), cy - int(15*sz),
                      int(12*sz), wing_h, fill=secondary + (100,), outline=accent + (150,))
        # Luminous glow
        _draw_ellipse(draw, cx, cy, body_rx + int(8*sz), body_ry + int(8*sz), fill=light + (30,))
        # Big sparkly eyes
        _draw_ellipse(draw, cx - eye_sep, head_cy - int(2*sz), eye_r + 1, eye_r + 1, fill=(200, 230, 255))
        _draw_ellipse(draw, cx + eye_sep, head_cy - int(2*sz), eye_r + 1, eye_r + 1, fill=(200, 230, 255))
        _draw_ellipse(draw, cx - eye_sep, head_cy - int(3*sz), int(eye_r*0.3), int(eye_r*0.3), fill=(255,255,255))
        _draw_ellipse(draw, cx - eye_sep, head_cy - int(2*sz), int(eye_r*0.4), int(eye_r*0.4), fill=accent)
        _draw_ellipse(draw, cx + eye_sep, head_cy - int(2*sz), int(eye_r*0.4), int(eye_r*0.4), fill=accent)
    elif noun == "Rat":
        # Whiskers
        for side in [-1, 1]:
            for wi in range(3):
                wy = head_cy + int(2*sz) + wi * int(2*sz)
                _draw_line(draw, cx + side * int(3*sz), wy,
                           cx + side * int(18*sz), wy + rng.randint(-3, 5),
                           fill=dark + (150,), width=1)
        # Long tail
        pts = _bezier_points((cx + body_rx - int(3*sz), cy + body_ry - int(5*sz)),
                              (cx + body_rx + int(15*sz), cy + body_ry + int(10*sz)),
                              (cx + body_rx + int(25*sz), cy + body_ry - int(5*sz)),
                              (cx + body_rx + int(30*sz), cy + body_ry + int(5*sz)))
        _draw_thick_curve(draw, pts, fill=light, width=1.5)
        # Beady eyes
        _draw_ellipse(draw, cx - eye_sep, head_cy - int(2*sz), int(eye_r*0.6), int(eye_r*0.6), fill=(20, 20, 20))
        _draw_ellipse(draw, cx + eye_sep, head_cy - int(2*sz), int(eye_r*0.6), int(eye_r*0.6), fill=(20, 20, 20))
        # Big round ears
        _draw_ellipse(draw, cx - head_r + int(1*sz), head_cy - head_r + int(2*sz),
                      int(7*sz), int(7*sz), fill=secondary)
        _draw_ellipse(draw, cx + head_r - int(1*sz), head_cy - head_r + int(2*sz),
                      int(7*sz), int(7*sz), fill=secondary)
    elif noun == "Wasp":
        # Striped abdomen
        for si in range(3):
            sy = cy + int((-2 + si * 7) * sz)
            _draw_rect(draw, cx - body_rx + int(3*sz), sy, body_rx * 2 - int(6*sz), int(3*sz),
                       fill=(20, 20, 20, 150))
        # Translucent wings
        _draw_ellipse(draw, cx - body_rx - int(3*sz), cy - int(18*sz),
                      int(14*sz), int(20*sz), fill=secondary + (60,), outline=accent + (100,))
        _draw_ellipse(draw, cx + body_rx + int(3*sz), cy - int(18*sz),
                      int(14*sz), int(20*sz), fill=secondary + (60,), outline=accent + (100,))
        # Stinger
        _draw_polygon(draw, [(cx, cy + body_ry),
                              (cx - int(2*sz), cy + body_ry + int(10*sz)),
                              (cx + int(2*sz), cy + body_ry + int(10*sz))], fill=accent)
        # Compound eyes
        _draw_ellipse(draw, cx - eye_sep, head_cy - int(1*sz), eye_r, eye_r, fill=(150, 0, 0, 200))
        _draw_ellipse(draw, cx + eye_sep, head_cy - int(1*sz), eye_r, eye_r, fill=(150, 0, 0, 200))
        # Antennae
        _draw_line(draw, cx - int(3*sz), head_cy - head_r, cx - int(8*sz), head_cy - head_r - int(12*sz),
                   fill=dark, width=1)
        _draw_line(draw, cx + int(3*sz), head_cy - head_r, cx + int(8*sz), head_cy - head_r - int(12*sz),
                   fill=dark, width=1)
    elif noun == "Bug":
        # Segmented body (draw lines across)
        for si in range(3):
            sy = cy + int((-5 + si * 10) * sz)
            _draw_line(draw, cx - body_rx + int(2*sz), sy, cx + body_rx - int(2*sz), sy,
                       fill=dark + (100,), width=1.5)
        # Mandibles
        _draw_line(draw, cx - int(4*sz), head_cy + head_r, cx - int(10*sz), head_cy + head_r + int(6*sz),
                   fill=accent, width=2)
        _draw_line(draw, cx + int(4*sz), head_cy + head_r, cx + int(10*sz), head_cy + head_r + int(6*sz),
                   fill=accent, width=2)
        # Compound eyes
        _draw_ellipse(draw, cx - eye_sep - 1, head_cy - int(1*sz), eye_r + 1, eye_r + 1, fill=(80, 150, 80, 200))
        _draw_ellipse(draw, cx + eye_sep + 1, head_cy - int(1*sz), eye_r + 1, eye_r + 1, fill=(80, 150, 80, 200))
        # Antennae
        _draw_line(draw, cx - int(2*sz), head_cy - head_r, cx - int(12*sz), head_cy - head_r - int(15*sz),
                   fill=dark, width=1)
        _draw_line(draw, cx + int(2*sz), head_cy - head_r, cx + int(12*sz), head_cy - head_r - int(15*sz),
                   fill=dark, width=1)
        _draw_ellipse(draw, cx - int(12*sz), head_cy - head_r - int(15*sz), 2, 2, fill=accent)
        _draw_ellipse(draw, cx + int(12*sz), head_cy - head_r - int(15*sz), 2, 2, fill=accent)
    else:
        # Generic small creature (fallback)
        # Triangle ears
        _draw_polygon(draw, [(cx - int(8*sz), head_cy - head_r + int(3*sz)),
                              (cx - int(14*sz), head_cy - head_r - int(12*sz)),
                              (cx - int(2*sz), head_cy - head_r + int(5*sz))], fill=secondary)
        _draw_polygon(draw, [(cx + int(8*sz), head_cy - head_r + int(3*sz)),
                              (cx + int(14*sz), head_cy - head_r - int(12*sz)),
                              (cx + int(2*sz), head_cy - head_r + int(5*sz))], fill=secondary)
        # Eyes
        _draw_ellipse(draw, cx - eye_sep, head_cy - int(2*sz), eye_r, int(eye_r*0.8), fill=(255, 255, 255))
        _draw_ellipse(draw, cx + eye_sep, head_cy - int(2*sz), eye_r, int(eye_r*0.8), fill=(255, 255, 255))
        _draw_ellipse(draw, cx - eye_sep, head_cy - int(2*sz), int(eye_r*0.5), int(eye_r*0.5), fill=accent)
        _draw_ellipse(draw, cx + eye_sep, head_cy - int(2*sz), int(eye_r*0.5), int(eye_r*0.5), fill=accent)
        # Tail
        pts = _quad_bezier_points((cx + body_rx - int(3*sz), cy + int(5*sz)),
                                   (cx + body_rx + int(15*sz), cy + int(15*sz)),
                                   (cx + body_rx + int(12*sz), cy - int(3*sz)))
        _draw_thick_curve(draw, pts, fill=primary, width=2.5)


# --- Ability VFX ---

def _draw_ability_vfx(draw, ability, accent, rng):
    ab = ability.lower()

    if "fire" in ab or "flame" in ab or "lava" in ab:
        for _ in range(rng.randint(4, 8)):
            fx = rng.randint(30, 170)
            fy = rng.randint(20, 200)
            sz = rng.randint(5, 12)
            alpha = rng.randint(80, 180)
            _draw_polygon(draw, [(fx, fy), (fx - sz, fy + sz * 2), (fx + sz, fy + sz * 2)],
                          fill=(255, 68, 0, alpha))
            _draw_polygon(draw, [(fx, fy + 2), (fx - sz + 2, fy + sz * 2 - 2), (fx + sz - 2, fy + sz * 2 - 2)],
                          fill=(255, 215, 0, int(alpha * 0.7)))

    elif "ice" in ab or "frost" in ab or "frozen" in ab or "glacial" in ab:
        for _ in range(rng.randint(5, 10)):
            ix = rng.randint(25, 175)
            iy = rng.randint(20, 220)
            sz = rng.randint(4, 10)
            alpha = rng.randint(80, 180)
            _draw_polygon(draw, [(ix, iy - sz), (ix + sz // 2, iy), (ix, iy + sz), (ix - sz // 2, iy)],
                          fill=(170, 221, 255, alpha))

    elif "lightning" in ab or "thunder" in ab:
        for _ in range(rng.randint(2, 4)):
            lx = rng.randint(30, 170)
            ly = rng.randint(10, 80)
            pts = [(lx, ly)]
            for _ in range(rng.randint(3, 5)):
                lx += rng.randint(-8, 8)
                ly += rng.randint(15, 30)
                pts.append((lx, ly))
            _draw_thick_curve(draw, pts, fill=(255, 238, 68, 200), width=2)

    elif "poison" in ab or "plague" in ab or "venom" in ab or "acid" in ab:
        for _ in range(rng.randint(5, 10)):
            px = rng.randint(25, 175)
            py = rng.randint(40, 230)
            pr = rng.randint(3, 7)
            alpha = rng.randint(50, 130)
            _draw_ellipse(draw, px, py, pr, pr, fill=(68, 221, 68, alpha))
            _draw_ellipse(draw, px - 1, py - 1, pr // 2, pr // 2, fill=(136, 255, 136, int(alpha * 0.6)))

    elif "shadow" in ab or "void" in ab or "death" in ab or "soul" in ab:
        for _ in range(rng.randint(3, 6)):
            sx = rng.randint(30, 170)
            sy = rng.randint(30, 200)
            ex = sx + rng.randint(-30, 30)
            ey = sy + rng.randint(-30, 30)
            mx = (sx + ex) // 2 + rng.randint(-15, 15)
            my = (sy + ey) // 2 + rng.randint(-15, 15)
            pts = _quad_bezier_points((sx, sy), (mx, my), (ex, ey))
            _draw_thick_curve(draw, pts, fill=(51, 34, 68, 100), width=rng.randint(2, 5))

    elif "wind" in ab:
        for _ in range(rng.randint(3, 5)):
            wx = rng.randint(20, 160)
            wy = rng.randint(30, 200)
            pts = _quad_bezier_points((wx, wy), (wx + 15, wy - 5), (wx + 35, wy))
            _draw_thick_curve(draw, pts, fill=(170, 204, 221, 128), width=1.5)

    elif "crystal" in ab or "stone" in ab or "earth" in ab or "petrif" in ab:
        for _ in range(rng.randint(4, 7)):
            cx_ = rng.randint(25, 175)
            cy_ = rng.randint(30, 220)
            sz = rng.randint(4, 9)
            _draw_polygon(draw, [(cx_, cy_ - sz), (cx_ + sz, cy_ - sz // 3),
                                  (cx_ + sz // 2, cy_ + sz), (cx_ - sz // 2, cy_ + sz),
                                  (cx_ - sz, cy_ - sz // 3)], fill=accent + (128,))

    elif "blood" in ab:
        for _ in range(rng.randint(4, 8)):
            bx = rng.randint(30, 170)
            by = rng.randint(40, 220)
            bs = rng.randint(3, 6)
            _draw_ellipse(draw, bx, by, bs, bs, fill=(153, 0, 0, 128))

    elif "holy" in ab or "radiant" in ab:
        for _ in range(rng.randint(3, 6)):
            gx = rng.randint(30, 170)
            gy = rng.randint(20, 200)
            gr = rng.randint(5, 12)
            _draw_ellipse(draw, gx, gy, gr, gr, fill=(255, 238, 136, 75))
            _draw_ellipse(draw, gx, gy, gr // 2, gr // 2, fill=(255, 255, 255, 50))

    elif "thorn" in ab:
        for _ in range(rng.randint(4, 8)):
            tx = rng.randint(30, 170)
            ty = rng.randint(30, 220)
            ts = rng.randint(5, 10)
            angle = rng.uniform(0, math.pi * 2)
            ex = tx + int(math.cos(angle) * ts)
            ey = ty + int(math.sin(angle) * ts)
            _draw_line(draw, tx, ty, ex, ey, fill=(51, 102, 17, 150), width=2)

    elif "mind" in ab or "teleport" in ab:
        for i in range(3):
            r = 20 + i * 15
            alpha = max(20, 100 - i * 30)
            _draw_ellipse(draw, 100, 100, r, r, fill=None, outline=accent + (alpha,), width=2)

    elif "summon" in ab or "berserk" in ab or "regen" in ab:
        for _ in range(rng.randint(5, 10)):
            px = rng.randint(25, 175)
            py = rng.randint(25, 225)
            pr = rng.randint(2, 5)
            _draw_ellipse(draw, px, py, pr, pr, fill=accent + (65,))


# ---------------------------------------------------------------------------
# IMAGE -> PNG BYTES
# ---------------------------------------------------------------------------

def image_to_png_bytes(img):
    buf = io.BytesIO()
    img.save(buf, format="PNG")
    buf.seek(0)
    return buf


# ---------------------------------------------------------------------------
# ITEM ICON GENERATION (200x200 canvas)
# ---------------------------------------------------------------------------

ITEM_CANVAS = 200

def _draw_item_sword(draw, primary, secondary):
    """Vertical blade + crossguard + handle + pommel."""
    cx = 100
    # Blade
    draw.rectangle([cx - 8, 20, cx + 8, 130], fill=secondary)
    # Pointed tip
    draw.polygon([(cx - 8, 20), (cx, 5), (cx + 8, 20)], fill=secondary)
    # Edge highlight
    draw.rectangle([cx - 3, 25, cx - 1, 125], fill=_lighten_tuple(secondary, 0.4))
    # Crossguard
    draw.rectangle([cx - 30, 130, cx + 30, 140], fill=primary)
    # Handle
    draw.rectangle([cx - 5, 140, cx + 5, 175], fill=_darken_tuple(primary, 0.5))
    # Handle wrap lines
    for y in range(145, 175, 6):
        draw.line([(cx - 5, y), (cx + 5, y)], fill=primary, width=1)
    # Pommel
    draw.ellipse([cx - 10, 172, cx + 10, 192], fill=primary)


def _draw_item_shield(draw, primary, secondary):
    """Heraldic pentagon + inner border + center emblem."""
    cx, cy = 100, 95
    # Shield outline (pentagon)
    pts = [(cx, 15), (cx + 60, 40), (cx + 55, 130), (cx, 170), (cx - 55, 130), (cx - 60, 40)]
    draw.polygon(pts, fill=primary)
    # Inner border
    pts_inner = [(cx, 28), (cx + 48, 48), (cx + 44, 120), (cx, 155), (cx - 44, 120), (cx - 48, 48)]
    draw.polygon(pts_inner, fill=_darken_tuple(primary, 0.8))
    # Center emblem — cross
    draw.rectangle([cx - 4, 55, cx + 4, 135], fill=secondary)
    draw.rectangle([cx - 30, 85, cx + 30, 93], fill=secondary)


def _draw_item_potion(draw, primary, secondary):
    """Circle flask + neck + cork + liquid fill + bubbles."""
    cx = 100
    # Flask body
    draw.ellipse([cx - 40, 70, cx + 40, 170], fill=_lighten_tuple(primary, 0.3))
    # Liquid fill (lower portion)
    draw.pieslice([cx - 40, 70, cx + 40, 170], start=0, end=180, fill=primary)
    draw.rectangle([cx - 40, 120, cx + 40, 170], fill=primary)
    # Flask outline
    draw.ellipse([cx - 40, 70, cx + 40, 170], outline=_darken_tuple(primary, 0.6), width=2)
    # Neck
    draw.rectangle([cx - 12, 40, cx + 12, 75], fill=_lighten_tuple(primary, 0.3))
    draw.rectangle([cx - 12, 40, cx + 12, 75], outline=_darken_tuple(primary, 0.6), width=2)
    # Cork
    draw.rectangle([cx - 15, 25, cx + 15, 44], fill=(160, 120, 60))
    draw.rectangle([cx - 15, 25, cx + 15, 44], outline=(120, 90, 40), width=2)
    # Bubbles
    for bx, by, br in [(cx - 15, 110, 5), (cx + 10, 125, 4), (cx - 5, 140, 3), (cx + 18, 108, 3)]:
        draw.ellipse([bx - br, by - br, bx + br, by + br], fill=secondary + (120,))


def _draw_item_gem(draw, primary, secondary):
    """Faceted octagon + highlight + sparkle dots."""
    cx, cy = 100, 100
    r = 55
    # Octagon
    pts = []
    for i in range(8):
        angle = math.pi / 8 + i * math.pi / 4
        pts.append((cx + int(r * math.cos(angle)), cy + int(r * math.sin(angle))))
    draw.polygon(pts, fill=primary)
    # Inner facet
    r2 = 35
    pts2 = []
    for i in range(8):
        angle = math.pi / 8 + i * math.pi / 4
        pts2.append((cx + int(r2 * math.cos(angle)), cy + int(r2 * math.sin(angle))))
    draw.polygon(pts2, fill=_lighten_tuple(primary, 0.3))
    # Highlight triangle
    draw.polygon([(cx, cy - 30), (cx - 20, cy), (cx + 5, cy - 10)],
                 fill=_lighten_tuple(secondary, 0.5) + (100,))
    # Sparkle dots
    for sx, sy in [(cx - 25, cy - 30), (cx + 30, cy - 20), (cx + 20, cy + 25), (cx - 10, cy + 30)]:
        draw.ellipse([sx - 3, sy - 3, sx + 3, sy + 3], fill=(255, 255, 255, 180))


def _draw_item_skull(draw, primary, secondary):
    """Cranium ellipse + jaw + dark eye sockets + nose."""
    cx, cy = 100, 85
    # Cranium
    draw.ellipse([cx - 45, cy - 50, cx + 45, cy + 40], fill=primary)
    # Jaw
    draw.rectangle([cx - 30, cy + 20, cx + 30, cy + 60], fill=_darken_tuple(primary, 0.85))
    draw.arc([cx - 30, cy + 40, cx + 30, cy + 75], start=0, end=180, fill=_darken_tuple(primary, 0.7), width=3)
    # Eye sockets
    draw.ellipse([cx - 30, cy - 20, cx - 8, cy + 10], fill=secondary)
    draw.ellipse([cx + 8, cy - 20, cx + 30, cy + 10], fill=secondary)
    # Nose
    draw.polygon([(cx, cy + 8), (cx - 6, cy + 22), (cx + 6, cy + 22)], fill=secondary)
    # Teeth
    for tx in range(-20, 21, 10):
        draw.rectangle([cx + tx - 3, cy + 35, cx + tx + 3, cy + 50], fill=_lighten_tuple(primary, 0.3))
        draw.rectangle([cx + tx - 3, cy + 35, cx + tx + 3, cy + 50],
                       outline=_darken_tuple(primary, 0.6), width=1)


def _draw_item_scroll(draw, primary, secondary):
    """Horizontal rect + rolled ends + text lines."""
    cx, cy = 100, 100
    # Main parchment
    draw.rectangle([30, 55, 170, 145], fill=_lighten_tuple(primary, 0.5))
    # Rolled top
    draw.ellipse([25, 42, 175, 68], fill=primary)
    # Rolled bottom
    draw.ellipse([25, 132, 175, 158], fill=primary)
    # Text lines
    for y in range(72, 132, 10):
        lw = 80 if y < 120 else 50
        draw.line([(cx - lw // 2, y), (cx + lw // 2, y)], fill=secondary + (150,), width=2)
    # Seal/ribbon
    draw.ellipse([cx - 8, 140, cx + 8, 156], fill=secondary)


def _draw_item_ingot(draw, primary, secondary):
    """Trapezoid with highlight stripe + perspective shading."""
    # Top face (lighter)
    top_face = [(55, 60), (145, 60), (165, 90), (35, 90)]
    draw.polygon(top_face, fill=_lighten_tuple(primary, 0.3))
    # Front face
    front_face = [(35, 90), (165, 90), (155, 150), (45, 150)]
    draw.polygon(front_face, fill=primary)
    # Highlight stripe on top
    draw.polygon([(70, 65), (130, 65), (140, 85), (50, 85)], fill=_lighten_tuple(secondary, 0.4) + (100,))
    # Edge lines
    draw.line([(35, 90), (165, 90)], fill=_darken_tuple(primary, 0.6), width=2)
    draw.line([(35, 90), (45, 150)], fill=_darken_tuple(primary, 0.6), width=2)
    draw.line([(165, 90), (155, 150)], fill=_darken_tuple(primary, 0.6), width=2)
    draw.line([(45, 150), (155, 150)], fill=_darken_tuple(primary, 0.6), width=2)
    # Engraved mark
    draw.rectangle([90, 105, 110, 135], fill=secondary + (100,))


def _draw_item_orb(draw, primary, secondary):
    """Large circle + concentric glow + specular highlight."""
    cx, cy = 100, 100
    # Outer glow
    for i in range(3):
        r = 70 - i * 8
        alpha = 40 + i * 20
        draw.ellipse([cx - r, cy - r, cx + r, cy + r], fill=secondary + (alpha,))
    # Main sphere
    draw.ellipse([cx - 50, cy - 50, cx + 50, cy + 50], fill=primary)
    # Inner ring
    draw.ellipse([cx - 35, cy - 35, cx + 35, cy + 35], fill=_lighten_tuple(primary, 0.2), outline=secondary + (100,), width=2)
    # Specular highlight crescent
    draw.pieslice([cx - 40, cy - 45, cx + 10, cy + 5], start=200, end=340,
                  fill=_lighten_tuple(primary, 0.5) + (120,))
    # Core spark
    draw.ellipse([cx - 8, cy - 8, cx + 8, cy + 8], fill=(255, 255, 255, 160))


def _lighten_tuple(color, factor=0.3):
    """Lighten an RGB tuple."""
    if isinstance(color, str):
        color = _hex_to_rgb(color)
    return tuple(min(255, int(c + (255 - c) * factor)) for c in color[:3])


def _darken_tuple(color, factor=0.6):
    """Darken an RGB tuple."""
    if isinstance(color, str):
        color = _hex_to_rgb(color)
    return tuple(int(c * factor) for c in color[:3])


ITEM_SHAPE_DRAWERS = {
    "sword": _draw_item_sword,
    "shield": _draw_item_shield,
    "potion": _draw_item_potion,
    "gem": _draw_item_gem,
    "skull": _draw_item_skull,
    "scroll": _draw_item_scroll,
    "ingot": _draw_item_ingot,
    "orb": _draw_item_orb,
}


def generate_item_image(item_key, item_data):
    """Generate a 200x200 item icon with rarity border."""
    primary = _hex_to_rgb(item_data["icon_primary"])
    secondary = _hex_to_rgb(item_data["icon_secondary"])
    rarity = item_data["rarity"]
    shape = item_data["icon_shape"]
    border_hex = RARITY_COLORS.get(rarity, "#999999")
    border_color = _hex_to_rgb(border_hex)

    img = Image.new("RGBA", (ITEM_CANVAS, ITEM_CANVAS), (240, 240, 240, 255))
    draw = ImageDraw.Draw(img, "RGBA")

    # Background gradient feel
    bg_light = _lighten_tuple(primary, 0.85)
    draw.ellipse([30, 30, 170, 170], fill=bg_light + (80,))

    # Draw shape
    drawer = ITEM_SHAPE_DRAWERS.get(shape, _draw_item_gem)
    drawer(draw, primary, secondary)

    # Rarity border (4px)
    bw = 4
    draw.rectangle([0, 0, ITEM_CANVAS - 1, bw - 1], fill=border_color)
    draw.rectangle([0, ITEM_CANVAS - bw, ITEM_CANVAS - 1, ITEM_CANVAS - 1], fill=border_color)
    draw.rectangle([0, 0, bw - 1, ITEM_CANVAS - 1], fill=border_color)
    draw.rectangle([ITEM_CANVAS - bw, 0, ITEM_CANVAS - 1, ITEM_CANVAS - 1], fill=border_color)
    # Corner accents
    for cx, cy in [(0, 0), (ITEM_CANVAS - 12, 0), (0, ITEM_CANVAS - 12), (ITEM_CANVAS - 12, ITEM_CANVAS - 12)]:
        draw.rectangle([cx, cy, cx + 12, cy + 12], fill=border_color)

    # Convert to RGB
    final = Image.new("RGB", img.size, (255, 255, 255))
    final.paste(img, mask=img.split()[3])
    return final


# ---------------------------------------------------------------------------
# CROSS-REFERENCE SYSTEM
# ---------------------------------------------------------------------------

def build_cross_references(creatures, item_encyclopedia):
    """Build item_key -> list of (creature_name, drop_rate_percent)."""
    rarity_factor = {
        "Common": 1.0, "Uncommon": 0.75, "Rare": 0.5, "Epic": 0.3, "Legendary": 0.15,
    }
    xref = {k: [] for k in item_encyclopedia}
    for creature in creatures:
        for drop_key in creature["drops"]:
            if drop_key in item_encyclopedia:
                rng = _seeded_random(creature["name"] + drop_key)
                base_rate = rng.randint(5, 40)
                rarity = item_encyclopedia[drop_key]["rarity"]
                rate = round(base_rate * rarity_factor.get(rarity, 0.5), 1)
                xref[drop_key].append((creature["name"], rate))
    # Sort each list by drop rate descending
    for k in xref:
        xref[k].sort(key=lambda x: -x[1])
    return xref


# ---------------------------------------------------------------------------
# DOCX ASSEMBLY
# ---------------------------------------------------------------------------

def build_docx(creatures, cross_refs, output_path):
    doc = Document()

    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    # --- Title Page ---
    _add_title_page(doc)

    # --- About Voxels ---
    _add_about_page(doc)

    # --- Table of Contents ---
    doc.add_heading("Table of Contents", level=1)
    p = doc.add_paragraph(
        "This encyclopedia contains two sections:\n\n"
        "BESTIARY — 500 procedurally-generated creatures, each with vector-style art, "
        "visual descriptions, combat statistics, and lore.\n\n"
        "ARMORY & ARTIFACTS — 30 items of power with icons, effects, stat bonuses, "
        "and cross-references to the creatures that drop them."
    )
    p.paragraph_format.space_after = Pt(6)
    doc.add_page_break()

    # --- BESTIARY section divider ---
    _add_section_divider(doc, "BESTIARY", "500 Creatures of the Voxel World")

    # --- Creature entries ---
    total = len(creatures)
    for idx, creature in enumerate(creatures):
        print(f"  [{idx + 1}/{total}] {creature['name']}")

        h = doc.add_heading(creature["name"], level=2)
        h.runs[0].font.color.rgb = RGBColor(*_hex_to_rgb(creature["primary_color"]))

        try:
            img = generate_creature_image(creature)
            img_stream = image_to_png_bytes(img)
            p_img = doc.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p_img.add_run()
            run.add_picture(img_stream, width=Inches(2.0))
        except Exception as e:
            p_err = doc.add_paragraph(f"[Image generation failed: {e}]")
            p_err.runs[0].font.color.rgb = RGBColor(200, 0, 0)

        vd_heading = doc.add_paragraph()
        vd_run = vd_heading.add_run("Visual Description")
        vd_run.bold = True
        vd_run.font.size = Pt(10)
        vd_run.font.color.rgb = RGBColor(80, 80, 80)

        vd_para = doc.add_paragraph(creature["visual_description"])
        vd_para.paragraph_format.space_after = Pt(4)
        vd_para.runs[0].font.size = Pt(9)
        vd_para.runs[0].font.italic = True

        _add_stats_table(doc, creature)

        lore_heading = doc.add_paragraph()
        lore_run = lore_heading.add_run("Lore")
        lore_run.bold = True
        lore_run.font.size = Pt(10)
        lore_run.font.color.rgb = RGBColor(80, 80, 80)

        lore_para = doc.add_paragraph(creature["lore"])
        lore_para.paragraph_format.space_after = Pt(6)
        lore_para.runs[0].font.size = Pt(9)

        if (idx + 1) % 2 == 0 and idx < total - 1:
            doc.add_page_break()

    # --- ARMORY section divider ---
    doc.add_page_break()
    _add_section_divider(doc, "ARMORY & ARTIFACTS", "30 Items of Power")

    # --- Item entries ---
    item_keys = list(ITEM_ENCYCLOPEDIA.keys())
    total_items = len(item_keys)
    for idx, item_key in enumerate(item_keys):
        item_data = ITEM_ENCYCLOPEDIA[item_key]
        sources = cross_refs.get(item_key, [])
        print(f"  [Item {idx + 1}/{total_items}] {item_data['display_name']}")
        _add_item_entry(doc, item_key, item_data, sources)
        if (idx + 1) % 2 == 0 and idx < total_items - 1:
            doc.add_page_break()

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"\nSaved: {output_path}")
    print(f"Creatures: {total}, Items: {total_items}")


def _add_title_page(doc):
    for _ in range(4):
        doc.add_paragraph("")

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run("VOXELS")
    run.bold = True
    run.font.size = Pt(44)
    run.font.color.rgb = RGBColor(0, 170, 170)
    run.font.name = "Calibri"

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = subtitle_p.add_run("The Complete Encyclopedia")
    run2.font.size = Pt(18)
    run2.font.color.rgb = RGBColor(80, 80, 80)
    run2.font.name = "Calibri"

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_p.add_run(datetime.now().strftime("%B %d, %Y"))
    date_run.font.size = Pt(12)
    date_run.font.color.rgb = RGBColor(120, 120, 120)

    desc_p = doc.add_paragraph()
    desc_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    desc_p.paragraph_format.space_before = Pt(24)
    desc_run = desc_p.add_run(
        "A comprehensive guide to the voxel-built world.\n"
        "500 creatures in the Bestiary. 30 artifacts in the Armory.\n"
        "Stats, lore, art, and item effects \u2014 all procedurally generated."
    )
    desc_run.font.size = Pt(11)
    desc_run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_page_break()


def _add_about_page(doc):
    doc.add_heading("About Voxels", level=1)

    intro = doc.add_paragraph()
    intro.paragraph_format.space_after = Pt(12)
    r = intro.add_run(
        "Voxels is a world built from cubes, but nothing about it is simple. "
        "Every mountain, cavern, and floating island is assembled from voxel blocks "
        "that shift and evolve as the world ages. Creatures roam biomes shaped by "
        "elemental forces, and adventurers explore procedurally generated landscapes "
        "that never repeat."
    )
    r.font.size = Pt(11)

    mech = doc.add_paragraph()
    mech.paragraph_format.space_after = Pt(12)
    r2 = mech.add_run(
        "Items in Voxels are powerful but finite. Every artifact, potion, and weapon "
        "has a limited number of uses \u2014 the rarer the item, the fewer times it can be "
        "activated. This creates hard choices: burn a Legendary relic to survive a boss, "
        "or save it for the dungeon ahead? Resource management is survival."
    )
    r2.font.size = Pt(11)

    theme = doc.add_paragraph()
    theme.paragraph_format.space_after = Pt(12)
    r3 = theme.add_run(
        "This encyclopedia catalogs the creatures and artifacts discovered so far. "
        "The Bestiary documents 500 known species with their combat stats, habitats, "
        "and lore. The Armory details the 30 artifacts that can be looted from these "
        "creatures, including their effects, drawbacks, and drop rates. "
        "Everything here was procedurally generated \u2014 just like the world itself."
    )
    r3.font.size = Pt(11)

    doc.add_page_break()


def _add_section_divider(doc, title, subtitle):
    for _ in range(8):
        doc.add_paragraph("")

    tp = doc.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = tp.add_run(title)
    tr.bold = True
    tr.font.size = Pt(32)
    tr.font.color.rgb = RGBColor(0, 170, 170)
    tr.font.name = "Calibri"

    sp = doc.add_paragraph()
    sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sp.add_run(subtitle)
    sr.font.size = Pt(14)
    sr.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_page_break()


def _add_stats_table(doc, creature):
    # Build drops text with cross-references
    drops_parts = []
    for d in creature["drops"]:
        if d in ITEM_ENCYCLOPEDIA:
            drops_parts.append(f"{d} [see Armory]")
        else:
            drops_parts.append(d)
    drops_text = ", ".join(drops_parts)

    stats = [
        ("Size", f"{creature['size']} ({creature['height_m']}m)"),
        ("Level", str(creature["level"])),
        ("HP", str(creature["hp"])),
        ("Attack", str(creature["attack"])),
        ("Defense", str(creature["defense"])),
        ("Speed", str(creature["speed"])),
        ("Ability", creature["ability"]),
        ("Habitat", creature["habitat"]),
        ("Behavior", creature["behavior"]),
        ("Drops", drops_text),
    ]

    table = doc.add_table(rows=len(stats), cols=2)
    table.style = "Light Shading"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, (label, value) in enumerate(stats):
        cell_label = table.cell(i, 0)
        cell_value = table.cell(i, 1)
        cell_label.text = label
        cell_value.text = value

        for paragraph in cell_label.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)
        for paragraph in cell_value.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(9)
                # Italicize cross-ref hints
                if label == "Drops" and "[see Armory]" in value:
                    run.font.italic = True

    for row in table.rows:
        row.cells[0].width = Inches(1.2)
        row.cells[1].width = Inches(5.0)

    doc.add_paragraph("")


def _add_item_entry(doc, item_key, item_data, sources):
    """Add a single item encyclopedia entry to the DOCX."""
    rarity = item_data["rarity"]
    rarity_hex = RARITY_COLORS.get(rarity, "#999999")
    rarity_rgb = _hex_to_rgb(rarity_hex)

    # Heading (colored by rarity)
    h = doc.add_heading(item_data["display_name"], level=2)
    h.runs[0].font.color.rgb = RGBColor(*rarity_rgb)

    # Item icon
    try:
        img = generate_item_image(item_key, item_data)
        img_stream = image_to_png_bytes(img)
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_img.add_run()
        run.add_picture(img_stream, width=Inches(1.2))
    except Exception as e:
        p_err = doc.add_paragraph(f"[Icon generation failed: {e}]")
        p_err.runs[0].font.color.rgb = RGBColor(200, 0, 0)

    # Info line
    info_p = doc.add_paragraph()
    info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_r = info_p.add_run(
        f"Type: {item_data['item_type']}  |  Rarity: {rarity}  |  Uses: {item_data['uses']}"
    )
    info_r.font.size = Pt(10)
    info_r.bold = True
    info_r.font.color.rgb = RGBColor(*rarity_rgb)

    # Original drop name
    orig_p = doc.add_paragraph()
    orig_r = orig_p.add_run(f"Drop Name: {item_key}")
    orig_r.font.size = Pt(9)
    orig_r.font.color.rgb = RGBColor(120, 120, 120)
    orig_r.italic = True

    # Effect
    eff_label = doc.add_paragraph()
    eff_r = eff_label.add_run("Effect")
    eff_r.bold = True
    eff_r.font.size = Pt(10)
    eff_r.font.color.rgb = RGBColor(80, 80, 80)

    eff_body = doc.add_paragraph()
    eff_body.paragraph_format.space_after = Pt(6)
    eb_r = eff_body.add_run(item_data["effect"])
    eb_r.font.size = Pt(9)
    eb_r.italic = True

    # Stat bonuses table (only non-zero stats)
    bonuses = {k: v for k, v in item_data.get("stat_bonuses", {}).items() if v != 0}
    if bonuses:
        st_label = doc.add_paragraph()
        st_r = st_label.add_run("Stat Bonuses")
        st_r.bold = True
        st_r.font.size = Pt(10)
        st_r.font.color.rgb = RGBColor(80, 80, 80)

        table = doc.add_table(rows=len(bonuses), cols=2)
        table.style = "Light Shading"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, (stat, val) in enumerate(bonuses.items()):
            table.cell(i, 0).text = stat
            sign = "+" if val > 0 else ""
            table.cell(i, 1).text = f"{sign}{val}"
            for paragraph in table.cell(i, 0).paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = Pt(9)
            for paragraph in table.cell(i, 1).paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(0, 120, 0) if val > 0 else RGBColor(180, 0, 0)
        for row in table.rows:
            row.cells[0].width = Inches(1.0)
            row.cells[1].width = Inches(1.0)

    # How to Obtain
    if sources:
        obt_label = doc.add_paragraph()
        obt_r = obt_label.add_run("How to Obtain")
        obt_r.bold = True
        obt_r.font.size = Pt(10)
        obt_r.font.color.rgb = RGBColor(80, 80, 80)

        for creature_name, rate in sources[:10]:  # Cap at 10 creatures
            src_p = doc.add_paragraph(style="List Bullet")
            sr = src_p.add_run(f"{creature_name} \u2014 {rate}% drop rate")
            sr.font.size = Pt(9)
        if len(sources) > 10:
            more_p = doc.add_paragraph()
            more_r = more_p.add_run(f"  ...and {len(sources) - 10} more creatures")
            more_r.font.size = Pt(8)
            more_r.font.color.rgb = RGBColor(120, 120, 120)
            more_r.italic = True

    # Flavor text
    flav_p = doc.add_paragraph()
    flav_p.paragraph_format.space_before = Pt(8)
    flav_p.paragraph_format.space_after = Pt(6)
    flav_r = flav_p.add_run(item_data["flavor_text"])
    flav_r.font.size = Pt(9)
    flav_r.italic = True
    flav_r.font.color.rgb = RGBColor(100, 100, 100)


# ---------------------------------------------------------------------------
# MAIN
# ---------------------------------------------------------------------------

def main():
    output_path = r"C:\Users\Doug\Downloads\monster_manual.docx"
    print("=" * 60)
    print("  VOXELS ENCYCLOPEDIA GENERATOR")
    print("  Bestiary + Armory + Vector Art + DOCX")
    print("=" * 60)

    print("\n[1/4] Generating 500 creatures...")
    creatures = generate_creatures(500)
    print(f"  Generated {len(creatures)} unique creatures.")

    print("\n[2/4] Building cross-references (creatures <-> items)...")
    cross_refs = build_cross_references(creatures, ITEM_ENCYCLOPEDIA)
    items_with_sources = sum(1 for v in cross_refs.values() if v)
    print(f"  {items_with_sources}/{len(ITEM_ENCYCLOPEDIA)} items have creature sources.")

    print("\n[3/4] Building DOCX (title, about, bestiary, armory)...")
    build_docx(creatures, cross_refs, output_path)

    print("\n[4/4] Complete!")
    print(f"  Output: {output_path}")
    print(f"  Creatures: {len(creatures)}")
    print(f"  Items: {len(ITEM_ENCYCLOPEDIA)}")
    print("=" * 60)


if __name__ == "__main__":
    main()
