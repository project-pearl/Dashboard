"""Generate PIN Dashboard Audit Report as .docx"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os, datetime

doc = Document()

# -- Styles --
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(4)

for level in range(1, 4):
    h = doc.styles[f'Heading {level}']
    h.font.name = 'Calibri'
    h.font.color.rgb = RGBColor(0x1A, 0x9A, 0x8E)  # PIN teal

# -- Helper --
def add_grade_table(doc, grade, color_hex):
    """Add a large grade badge inline."""
    p = doc.add_paragraph()
    run = p.add_run(f"  {grade}  ")
    run.font.size = Pt(28)
    run.bold = True
    r, g, b = int(color_hex[0:2], 16), int(color_hex[2:4], 16), int(color_hex[4:6], 16)
    run.font.color.rgb = RGBColor(r, g, b)
    return p

def add_domain(doc, domain_name, grade, color, current, target, reasoning, how_to_improve):
    doc.add_heading(domain_name, level=2)
    add_grade_table(doc, grade, color)

    doc.add_heading("Current State", level=3)
    for item in current:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading("Where We Should Be", level=3)
    for item in target:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading("Why This Grade", level=3)
    for item in reasoning:
        doc.add_paragraph(item, style='List Bullet')

    if how_to_improve:
        doc.add_heading("How to Reach an A", level=3)
        for i, item in enumerate(how_to_improve, 1):
            doc.add_paragraph(f"{i}. {item}")

    doc.add_paragraph()  # spacer


# ============================================================
# TITLE PAGE
# ============================================================
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_heading('PIN Water Quality Dashboard', level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle = doc.add_heading('Comprehensive Technical Audit Report', level=2)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(f"Generated {datetime.date.today().strftime('%B %d, %Y')}")
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_page_break()

# ============================================================
# EXECUTIVE SUMMARY
# ============================================================
doc.add_heading('Executive Summary', level=1)
doc.add_paragraph(
    "This report evaluates the PIN Water Quality Dashboard across 12 engineering domains. "
    "Each domain receives a letter grade (A through F), an explanation of the current state, "
    "the target state for production/government readiness, and a concrete roadmap to close gaps. "
    "The dashboard demonstrates strong data architecture and design system maturity but has "
    "critical gaps in testing, security hardening, and observability that must be addressed "
    "before pursuing DARPA, FedRAMP, or ATO engagements."
)

# Summary table
table = doc.add_table(rows=1, cols=3)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table.rows[0].cells
hdr[0].text = 'Domain'
hdr[1].text = 'Grade'
hdr[2].text = 'Priority'
for cell in hdr:
    for p in cell.paragraphs:
        p.runs[0].bold = True

grades_summary = [
    ("Security", "D+", "CRITICAL"),
    ("Testing & QA", "F", "CRITICAL"),
    ("CI/CD & DevOps", "C+", "HIGH"),
    ("Code Quality & Maintainability", "C-", "HIGH"),
    ("Documentation", "D", "HIGH"),
    ("Accessibility (a11y)", "C+", "MEDIUM"),
    ("Performance Optimization", "B-", "MEDIUM"),
    ("SEO & Discoverability", "C-", "LOW"),
    ("Design System & UX", "A-", "LOW"),
    ("Data Architecture & Caching", "B+", "LOW"),
    ("Monitoring & Observability", "C-", "HIGH"),
    ("Compliance & Governance", "D", "CRITICAL"),
]
for domain, grade, priority in grades_summary:
    row = table.add_row().cells
    row[0].text = domain
    row[1].text = grade
    row[2].text = priority

doc.add_page_break()

# ============================================================
# DETAILED DOMAIN ASSESSMENTS
# ============================================================
doc.add_heading('Detailed Domain Assessments', level=1)

# 1. SECURITY
add_domain(doc,
    "1. Security",
    "D+", "CC3333",
    current=[
        "Supabase auth with JWT + RBAC (4-tier admin hierarchy) is properly implemented",
        "Security headers present: X-Frame-Options DENY, X-Content-Type-Options nosniff, Referrer-Policy, Permissions-Policy",
        "Rate limiting on AI endpoints (10 req/60s per IP) but in-memory only",
        "Cron endpoints protected by single CRON_SECRET bearer token",
        "No Content-Security-Policy (CSP) header",
        "No CSRF token validation on POST/DELETE endpoints",
        "No HSTS header configured",
        "Next.js 13.5.1 is 18+ months out of support with 15 known CVEs including critical SSRF and auth bypass",
        "dangerouslySetInnerHTML usage in chart.tsx and ShuckAndDestroy.tsx",
        "npm audit shows 16 vulnerabilities (1 critical, 7 high, 7 moderate)",
        "Public endpoints (/api/cache-status, /api/water-data) expose internal architecture without auth",
        "All 47 cron endpoints share a single secret (compromise of one = compromise of all)",
    ],
    target=[
        "Zero critical/high CVEs in dependencies",
        "CSP, HSTS, and all OWASP-recommended headers in place",
        "CSRF protection on all state-changing endpoints",
        "Per-endpoint or per-category rate limiting with Redis/Upstash backend",
        "Input validation (Zod schemas) on every API route",
        "Secrets rotated quarterly with automated rotation",
        "SAST + DAST scanning in CI with blocking on high/critical findings",
    ],
    reasoning=[
        "Auth/RBAC implementation is solid (+), but the Next.js version is critically outdated with known exploits",
        "Security headers are partially implemented but missing the most important one (CSP)",
        "Rate limiting exists but is ineffective in serverless (resets on cold start)",
        "No CSRF protection on any endpoint is a significant gap",
        "Public cache-status endpoint leaks internal architecture details",
        "Single shared cron secret creates a blast radius problem",
    ],
    how_to_improve=[
        "Upgrade Next.js to 15.x (closes 15 known CVEs including critical SSRF)",
        "Add Content-Security-Policy header with strict script-src and style-src directives",
        "Add HSTS header (Strict-Transport-Security: max-age=31536000; includeSubDomains)",
        "Implement CSRF tokens on all POST/PUT/DELETE routes (Next.js middleware approach)",
        "Add Zod schema validation to all 105 API route inputs",
        "Migrate rate limiting to Upstash Redis (@upstash/ratelimit) for serverless persistence",
        "Add auth check to /api/cache-status and /api/water-data endpoints",
        "Separate cron secrets by job category (data-rebuild, sentinel, insights)",
        "Replace dangerouslySetInnerHTML with sanitized rendering or CSS modules",
        "Run npm audit fix for auto-fixable vulnerabilities; track unfixable ones in an exceptions list",
    ]
)

# 2. TESTING & QA
add_domain(doc,
    "2. Testing & Quality Assurance",
    "F", "CC0000",
    current=[
        "Zero test files across 607 TypeScript/TSX source files",
        "No testing framework installed (no Jest, Vitest, Cypress, or Playwright)",
        "No test runner in CI/CD pipeline",
        "105 API routes with zero test coverage",
        "236 React components with zero unit or integration tests",
        "48 cache modules with zero tests despite complex logic (spatial indexing, build locks, delta detection)",
        "Sentinel validation utilities exist (lib/sentinel/__tests__/) but are scoring tools, not test suites",
        "No pre-commit hooks to enforce quality gates",
    ],
    target=[
        "80%+ code coverage on critical paths (cache modules, API routes, auth)",
        "Unit tests for all cache modules (grid indexing, build locks, persistence)",
        "Integration tests for API routes (auth flows, cron endpoints, data endpoints)",
        "Component tests for key UI (dashboard rendering, form validation, role-based views)",
        "E2E tests for critical user journeys (login, dashboard load, alert creation)",
        "Test runner in CI blocking merges on failure",
        "Coverage reports uploaded as PR comments",
    ],
    reasoning=[
        "This is the most critical gap in the entire project",
        "607 source files with zero tests means any change could silently break production",
        "Cache modules contain complex spatial indexing and concurrency logic that is impossible to verify without tests",
        "Auth/RBAC logic is untested despite being security-critical",
        "No regression safety net makes refactoring (like the upcoming role redesigns) extremely risky",
    ],
    how_to_improve=[
        "Install Vitest + @testing-library/react + msw (Mock Service Worker)",
        "Start with cache module unit tests (highest complexity, most critical): wqpCache, attainsCache, icisCache",
        "Add API route integration tests using Next.js test helpers",
        "Add component tests for FederalManagementCenter (most complex component)",
        "Add Playwright for E2E: login flow, dashboard rendering per role, alert creation",
        "Add Vitest to CI workflow (fail on test failure, upload coverage report)",
        "Set up Husky + lint-staged pre-commit hooks running tests on changed files",
        "Target: 40% coverage in 30 days, 70% in 90 days, 80% in 6 months",
    ]
)

# 3. CI/CD & DEVOPS
add_domain(doc,
    "3. CI/CD & DevOps",
    "C+", "DD8800",
    current=[
        "Two GitHub Actions workflows: PIN Pipeline (hourly data fetch) + Security Scan (SAST + audit)",
        "Semgrep SAST with SARIF upload to GitHub Security tab",
        "npm audit with artifact upload (advisory, non-blocking)",
        "PIN Pipeline: auto-commit fetched data with bot identity",
        "Vercel deployment via git push (no manual deploy step)",
        "Vercel preview deployments on PRs (inferred)",
        "No test execution in any CI workflow",
        "No build validation step (ESLint + TypeScript both disabled during builds)",
        "No branch protection rules or required reviewers",
        "No PR template or CODEOWNERS file",
    ],
    target=[
        "CI pipeline: lint + typecheck + test + build + SAST + audit (all blocking)",
        "Branch protection: require PR reviews, require status checks, no direct push to main",
        "CODEOWNERS for automatic reviewer assignment",
        "PR template with checklist (tests, security, docs)",
        "Staging environment with smoke tests before production",
        "Deployment notifications (Slack/Teams on failure)",
        "Automated dependency updates (Dependabot or Renovate)",
    ],
    reasoning=[
        "Security scanning workflow is well-designed and just started working",
        "PIN Pipeline is clever (hourly incremental fetch with smart commit messages)",
        "But no tests, no lint enforcement, and no build validation in CI defeats the purpose of CI",
        "Anyone can push directly to main without review",
        "No quality gates means broken code reaches production regularly",
    ],
    how_to_improve=[
        "Add a 'build-and-test' workflow: npm run lint && npm run typecheck && npm run test && npm run build",
        "Enable branch protection on main: require 1 reviewer, require CI checks to pass",
        "Create .github/CODEOWNERS mapping lib/ -> backend team, components/ -> frontend team",
        "Create .github/PULL_REQUEST_TEMPLATE.md with testing/security checklist",
        "Add Dependabot or Renovate for automated dependency PRs",
        "Add Slack notification on workflow failure (GitHub Actions -> Slack webhook)",
        "Consider adding a staging deployment step before production",
    ]
)

# 4. CODE QUALITY
add_domain(doc,
    "4. Code Quality & Maintainability",
    "C-", "DD8800",
    current=[
        "TypeScript strict mode enabled in tsconfig.json (noImplicitAny, strictNullChecks, etc.)",
        "ESLint installed but explicitly disabled during builds (ignoreDuringBuilds: true)",
        "TypeScript errors explicitly ignored during builds (ignoreBuildErrors: true)",
        "Prettier installed but no configuration file exists",
        "No pre-commit hooks (no Husky, no lint-staged)",
        "Significant code duplication: 48 cache modules each implement their own loadFromDisk/saveToDisk despite shared helpers existing in cacheUtils.ts",
        "ATTAINS cache has its own blob implementation instead of using shared blobPersistence.ts",
        "Build lock timeout logic (~12 min auto-clear) duplicated across multiple cache modules",
        "Error handling inconsistent: 17% of API routes have try-catch (18/105)",
        "Only 2 React Error Boundaries (Federal + PEARL dashboards); no global error boundary",
    ],
    target=[
        "ESLint and TypeScript enforced during builds (zero errors)",
        "Prettier configured and enforced via pre-commit hooks",
        "Shared cache helpers used by all modules (zero duplication of persistence logic)",
        "Try-catch on all API routes with structured error responses",
        "Global error boundary at app root + per-page boundaries for dashboards",
        "Consistent code style enforced automatically",
    ],
    reasoning=[
        "TypeScript strict mode is good, but disabling it during builds nullifies the benefit",
        "Having shared helpers (cacheUtils.ts, blobPersistence.ts) but not using them is worse than not having them",
        "83% of API routes can fail silently without error handling",
        "No formatting enforcement means inconsistent code style across 607 files",
    ],
    how_to_improve=[
        "Remove ignoreDuringBuilds and ignoreBuildErrors from next.config.js; fix all errors (this is a large effort but essential)",
        "Create .prettierrc with project-wide formatting rules; run prettier --write on entire codebase",
        "Set up Husky + lint-staged: run eslint + prettier on staged files before each commit",
        "Refactor cache modules to use shared loadCacheFromDisk/saveCacheToDisk from cacheUtils.ts",
        "Migrate ATTAINS blob implementation to shared blobPersistence.ts",
        "Add try-catch to all 87 unprotected API routes with consistent JSON error responses",
        "Add global error.tsx at app root level",
    ]
)

# 5. DOCUMENTATION
add_domain(doc,
    "5. Documentation",
    "D", "CC3333",
    current=[
        'README.md contains only the word "Test"',
        "Sentinel architecture documented in docs/sentinel-system-architecture.md (271 lines, excellent)",
        "Sentinel validation metrics in docs/METRICS.md",
        "No CONTRIBUTING.md, CHANGELOG.md, or API reference",
        "No root .env.example documenting required environment variables",
        "Minimal JSDoc/TSDoc comments in code",
        "Critical architectural knowledge exists only in Claude memory files, not in repo",
        "Scripts directory has .env.example for MS4 contacts scraper only",
    ],
    target=[
        "Comprehensive README.md: project overview, architecture diagram, setup instructions, deployment guide",
        "API reference documenting all 105 endpoints with request/response schemas",
        ".env.example at root listing all 15+ required environment variables with descriptions",
        "CONTRIBUTING.md with development setup, coding standards, PR process",
        "CHANGELOG.md tracking releases and breaking changes",
        "Architecture Decision Records (ADRs) for key design choices",
        "Inline JSDoc on all exported functions in lib/",
    ],
    reasoning=[
        "The Sentinel architecture doc is excellent and shows the capability exists",
        "But the README being just 'Test' means no new developer can onboard",
        "15+ environment variables with no documentation is a setup blocker",
        "Architecture knowledge in Claude memory files is not accessible to other team members",
    ],
    how_to_improve=[
        "Write README.md: project overview, tech stack, prerequisites, setup instructions, architecture overview, deployment",
        "Create .env.example at root with all required variables and descriptions",
        "Create CONTRIBUTING.md: dev setup, branch naming, PR checklist, coding standards",
        "Generate API reference from route files (can use typedoc or manual markdown)",
        "Move architectural knowledge from Claude memory into docs/ directory",
        "Add JSDoc to all exported functions in lib/ (start with cache modules and auth)",
    ]
)

# 6. ACCESSIBILITY
add_domain(doc,
    "6. Accessibility (a11y)",
    "C+", "DD8800",
    current=[
        "Radix UI primitives provide built-in keyboard navigation, focus management, and ARIA roles",
        "30+ Shadcn/ui components with accessibility baked in (dialog, accordion, alert, tabs, etc.)",
        "Semantic HTML: role='alert', aria-live='assertive', aria-expanded, aria-label on key elements",
        "Screen reader text (.sr-only) on carousel controls, dialog close buttons, pagination",
        "Reduced motion support: @media (prefers-reduced-motion) disables sentinel pulse animation",
        "Alt text on images via next/image component",
        "No accessibility testing tools installed (no axe, pa11y, jest-axe)",
        "Limited aria-describedby and aria-errormessage on form fields",
        "No documented WCAG compliance level",
        "No color contrast verification",
        "No focus trap verification beyond Radix defaults",
    ],
    target=[
        "WCAG 2.1 AA compliance verified and documented",
        "Automated a11y testing in CI (axe-core or pa11y)",
        "Manual audit with screen reader (NVDA/VoiceOver) on all dashboard views",
        "Color contrast AA verified across light and dark themes",
        "Comprehensive ARIA attributes on all interactive elements",
        "Keyboard-only navigation verified for all user flows",
    ],
    reasoning=[
        "Radix UI choice was excellent for baseline accessibility",
        "Semantic HTML and ARIA usage shows awareness of a11y needs",
        "But zero testing means compliance is assumed, not verified",
        "Government software (DARPA context) requires Section 508 compliance which maps to WCAG 2.1 AA",
    ],
    how_to_improve=[
        "Install jest-axe and add a11y assertions to component tests",
        "Add pa11y-ci to CI pipeline for automated WCAG checking",
        "Run Lighthouse accessibility audit and fix all flagged issues",
        "Verify color contrast ratios in both light and dark themes (tool: contrast-ratio.com)",
        "Add aria-describedby to all form inputs linking to help text/errors",
        "Document target WCAG compliance level in README",
        "Conduct manual screen reader testing (NVDA on Windows, VoiceOver on Mac)",
    ]
)

# 7. PERFORMANCE
add_domain(doc,
    "7. Performance Optimization",
    "B-", "228B22",
    current=[
        "Dynamic imports for heavy components (Mapbox, charts) with ssr: false",
        "Suspense boundaries with skeleton fallbacks in dashboard layouts",
        "next/image used throughout for automatic lazy loading",
        "next/font/google for zero-layout-shift font loading (Inter, JetBrains Mono)",
        "Console.log removed in production via compiler.removeConsole (except warn)",
        "3-tier caching (memory -> disk -> Vercel Blob) prevents cold-start data loss",
        "Image optimization DISABLED globally (images.unoptimized: true in next.config.js)",
        "No bundle analyzer installed or configured",
        "No Lighthouse CI or Core Web Vitals tracking",
        "No code splitting beyond dynamic imports",
    ],
    target=[
        "Image optimization enabled (Vercel's built-in or custom loader)",
        "Bundle analyzer integrated; bundle budget enforced in CI",
        "Lighthouse CI tracking Core Web Vitals (LCP < 2.5s, FID < 100ms, CLS < 0.1)",
        "Route-level code splitting verified and optimized",
        "Web Workers for heavy computation (cache processing, data transforms)",
    ],
    reasoning=[
        "Dynamic imports and Suspense show good performance awareness",
        "3-tier caching architecture is sophisticated and well-designed",
        "But disabling image optimization in production is leaving significant performance on the table",
        "No measurement means no baseline and no ability to track regressions",
    ],
    how_to_improve=[
        "Re-enable image optimization: remove images.unoptimized or set up custom loader for Vercel",
        "Install @next/bundle-analyzer; add npm run analyze script",
        "Set up Lighthouse CI in GitHub Actions (budget: LCP < 2.5s, CLS < 0.1)",
        "Add @vercel/speed-insights for production Real User Monitoring (RUM)",
        "Profile largest dashboard pages and lazy-load below-fold sections",
    ]
)

# 8. SEO
add_domain(doc,
    "8. SEO & Discoverability",
    "C-", "DD8800",
    current=[
        "Root metadata configured: title, description, Open Graph, Twitter card with image",
        "PWA manifest.json with app name, icons, theme color",
        "themeColor set for mobile browser chrome",
        "No robots.txt file",
        "No sitemap.xml (static or dynamic)",
        "No per-page dynamic metadata for state/jurisdiction/permit pages",
        "No structured data (JSON-LD) for Organization, BreadcrumbList, or FAQ",
        "No hreflang tags (consistent with no i18n)",
    ],
    target=[
        "robots.txt with appropriate crawl directives (allow public pages, block dashboard internals)",
        "Dynamic sitemap.xml generated from route structure",
        "Per-page metadata for all public-facing routes",
        "JSON-LD structured data for Organization and key pages",
        "Canonical URLs on all pages",
    ],
    reasoning=[
        "Root metadata is properly configured which covers the homepage",
        "But dashboard pages (which are the primary content) have no individual metadata",
        "Missing robots.txt and sitemap means search engines can't efficiently crawl the site",
        "For a government-facing tool, discoverability matters for adoption",
    ],
    how_to_improve=[
        "Add public/robots.txt: allow public pages, disallow /dashboard/, /api/",
        "Create app/sitemap.ts using Next.js 13+ generateSitemaps() for dynamic generation",
        "Add generateMetadata() to state report and public-facing pages",
        "Add JSON-LD Organization schema to root layout",
        "Add canonical <link> tags via Next.js metadata API",
    ]
)

# 9. DESIGN SYSTEM
add_domain(doc,
    "9. Design System & UX",
    "A-", "228B22",
    current=[
        "Tailwind CSS 3.3 + Radix UI + Shadcn/ui: 30+ composable UI components",
        "Comprehensive design token system: 80+ CSS custom properties (HSL-based)",
        "Dark mode with class-based toggle + 80 CSS overrides for seamless transition",
        "4 accent color presets (Sapphire, Indigo, Emerald, Copper) with full theme support",
        "Domain-specific semantic colors (surface water teal, groundwater amber, etc.)",
        "Custom utility classes: .pin-card-glass, .pin-stat-hero, .pin-card-hover, .pin-card-tinted",
        "Typography system: Inter (body) + JetBrains Mono (data/stats) loaded via next/font",
        "Animations respect prefers-reduced-motion",
        "Theme persistence via next-themes + localStorage",
        "Pre-flight script prevents flash of unstyled content on theme load",
    ],
    target=[
        "Component documentation (Storybook or similar) for design system reference",
        "Visual regression testing to catch unintended style changes",
        "Design tokens exported for use outside the app (design handoff, docs)",
        "Consistent spacing/sizing scale documented",
    ],
    reasoning=[
        "This is the strongest domain in the project",
        "The design token architecture is production-grade with semantic naming and theme variants",
        "Radix UI provides the right balance of accessibility and customization",
        "Reduced-motion support and theme persistence show attention to user experience",
        "Only gap is lack of component documentation and visual regression testing",
    ],
    how_to_improve=[
        "Add Storybook for component documentation and visual testing",
        "Set up Chromatic or Percy for visual regression testing in CI",
        "Document the design token naming convention and usage guidelines",
        "Create a component playground page for internal reference",
    ]
)

# 10. DATA ARCHITECTURE
add_domain(doc,
    "10. Data Architecture & Caching",
    "B+", "228B22",
    current=[
        "12+ cache modules with grid-based spatial indexing (0.1-degree resolution)",
        "3-tier persistence: in-memory -> disk (.cache/) -> Vercel Blob (REST API)",
        "Cold-start survival: ensureWarmed() cascades through all three tiers",
        "Build lock pattern with 12-minute auto-clear timeout prevents stuck locks",
        "47 cron jobs staggered across 4-11 AM UTC to prevent overload",
        "ATTAINS self-chaining: 20-hop cascade with time-budgeted chunks and deferred queue",
        "Delta detection via signalsHash for efficient rebuild skipping",
        "Semaphore-based concurrency control (4 parallel states) for AI insights",
        "React Context API for auth, admin state, jurisdiction (no Redux/Zustand)",
        "Raw fetch() throughout (no SWR or React Query for caching/deduplication)",
        "Custom polling hooks for real-time data (sentinel: 60s, flood: 2min)",
    ],
    target=[
        "SWR or React Query for client-side data fetching (deduplication, cache invalidation, optimistic updates)",
        "Shared cache module base class eliminating persistence code duplication",
        "Redis/Upstash for distributed cache state (not in-memory per instance)",
        "Cache warming metrics and alerting on cold-start frequency",
        "Formal data flow documentation with sequence diagrams",
    ],
    reasoning=[
        "The caching architecture is genuinely sophisticated and well-thought-out",
        "Self-chaining ATTAINS pattern is innovative and solves a real serverless constraint",
        "Spatial indexing with grid keys is appropriate for the domain",
        "But client-side fetching with raw fetch() and useEffect means no request deduplication",
        "Cache module code duplication (48x persistence pattern) is a maintainability risk",
    ],
    how_to_improve=[
        "Adopt SWR or React Query for client-side data fetching (replaces manual useEffect + fetch patterns)",
        "Create a base CacheModule class/factory that handles loadFromDisk/saveToDisk/ensureWarmed",
        "Migrate ATTAINS to shared blobPersistence.ts helper",
        "Add cache hit/miss metrics to /api/cache-status endpoint",
        "Consider Upstash Redis for shared cache state across Vercel instances",
        "Document data flow with Mermaid diagrams in docs/",
    ]
)

# 11. MONITORING
add_domain(doc,
    "11. Monitoring & Observability",
    "C-", "DD8800",
    current=[
        "Custom health check endpoint (/api/source-health) monitoring 35 external data sources",
        "Unified cache status endpoint (/api/cache-status) tracking 50+ cache modules with staleness detection",
        "Console.warn preserved in production for runtime warnings",
        "Cache staleness flagging (> 48 hours = stale)",
        "Delta tracking (before/after record counts) on cache rebuilds",
        "No error tracking service (no Sentry, DataDog, Rollbar, or New Relic)",
        "No structured logging (console.log/warn/error only)",
        "No APM or distributed tracing",
        "No alerting on cron failures beyond GitHub Actions status",
        "No uptime monitoring",
    ],
    target=[
        "Sentry (or equivalent) for error tracking with source maps",
        "Structured logging (JSON format) for log aggregation",
        "APM with transaction tracing across API routes and cron jobs",
        "Uptime monitoring with alerting (PagerDuty, Opsgenie, or Slack)",
        "Dashboard for cron job success rates and execution times",
        "Real User Monitoring (RUM) for Core Web Vitals",
    ],
    reasoning=[
        "Custom health endpoints are well-designed and comprehensive",
        "But without an error tracking service, production errors are invisible",
        "Console logging in serverless is ephemeral and hard to search",
        "47 cron jobs running daily with no failure alerting is a reliability gap",
    ],
    how_to_improve=[
        "Install @sentry/nextjs for error tracking + performance monitoring",
        "Add Sentry to all API routes and cron jobs (automatic instrumentation)",
        "Add @vercel/speed-insights for Real User Monitoring",
        "Set up Vercel log drains to a log aggregation service (Axiom, Datadog, or Better Stack)",
        "Add Slack webhook notifications on cron job failures",
        "Create an operational dashboard (Grafana or Vercel Analytics) tracking cron success rates",
    ]
)

# 12. COMPLIANCE
add_domain(doc,
    "12. Compliance & Governance",
    "D", "CC3333",
    current=[
        "Supabase RBAC with 4-tier admin hierarchy (good foundation)",
        "CRON_SECRET auth on cron endpoints",
        "Security scanning workflow just added (Semgrep SAST + npm audit)",
        "No SBOM (Software Bill of Materials) generation",
        "No branch protection rules on main",
        "No CODEOWNERS file for required reviewers",
        "No PR template with security/compliance checklist",
        "No audit logging of admin actions",
        "No data retention or privacy policy implementation",
        "No FedRAMP, CMMC, or NIST 800-53 control mapping",
        "No secrets rotation policy or automation",
        "No incident response documentation",
    ],
    target=[
        "SBOM generated in CI (CycloneDX or SPDX format) per EO 14028",
        "Branch protection with required reviews and status checks",
        "Audit logging of all admin actions (role changes, user management, config changes)",
        "NIST 800-53 control mapping document for relevant controls",
        "Secrets rotation policy (quarterly) with automated rotation where possible",
        "Incident response plan documented",
        "Data handling classification and retention policies",
        "Supply chain security (npm provenance, lock file integrity verification)",
    ],
    reasoning=[
        "Auth/RBAC is solid but compliance requires more than authentication",
        "No SBOM is a blocker for EO 14028 compliance (required for federal software)",
        "No audit logging means admin actions are untrackable",
        "No branch protection means anyone can push to production without review",
        "Government contracts (DARPA included) require documented security controls",
    ],
    how_to_improve=[
        "Add CycloneDX SBOM generation to CI: npx @cyclonedx/cyclonedx-npm --output-file sbom.json",
        "Enable GitHub branch protection: require 1 reviewer, require status checks, no force push",
        "Create .github/CODEOWNERS mapping critical paths to required reviewers",
        "Add audit logging to admin API routes (log to Supabase table: who, what, when)",
        "Create a NIST 800-53 control mapping document for relevant controls (AC, AU, CM, IA, SC)",
        "Document secrets rotation policy and set calendar reminders",
        "Create incident response plan in docs/INCIDENT_RESPONSE.md",
        "Enable npm provenance in CI: npm publish --provenance (when applicable)",
    ]
)

# ============================================================
# OVERALL ASSESSMENT
# ============================================================
doc.add_page_break()
doc.add_heading('Overall Assessment', level=1)

doc.add_heading('Composite Grade: C-', level=2)

doc.add_paragraph(
    "The PIN Dashboard is architecturally ambitious and delivers genuine value through its "
    "data integration layer (47 cron jobs, 12 cache modules, 35+ external sources). The design system "
    "is production-grade. However, the complete absence of testing (0 tests across 607 files) and "
    "critical security gaps (outdated Next.js, no CSP, no CSRF) represent existential risks to "
    "production reliability and government adoption."
)

doc.add_heading('DARPA Positioning Impact', level=2)
doc.add_paragraph(
    "For DARPA specifically, the technical merit of the data architecture and real-time monitoring "
    "system is strong. However, DARPA evaluators will flag:"
)
items = [
    "No test suite = no confidence in correctness claims",
    "No SBOM = non-compliant with EO 14028 (Executive Order on Improving the Nation's Cybersecurity)",
    "Out-of-support Next.js version = known exploitable vulnerabilities",
    "No audit logging = no accountability trail for admin actions",
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('Priority Roadmap', level=2)
priorities = [
    ("Week 1-2 (Critical)", [
        "Upgrade Next.js to 15.x",
        "Install Vitest + write first 20 cache module tests",
        "Add CSP and HSTS headers",
        "Enable branch protection on main",
        "Add SBOM generation to CI",
    ]),
    ("Week 3-4 (High)", [
        "Add Sentry error tracking",
        "Remove ignoreDuringBuilds / ignoreBuildErrors; fix errors",
        "Add build-and-test CI workflow",
        "Write README.md and .env.example",
        "Add CSRF protection to state-changing endpoints",
    ]),
    ("Month 2 (Medium)", [
        "Reach 40% test coverage on critical paths",
        "Add Lighthouse CI for performance tracking",
        "Add accessibility testing (jest-axe + pa11y)",
        "Refactor cache modules to use shared helpers",
        "Create NIST 800-53 control mapping",
    ]),
    ("Month 3 (Polish)", [
        "Reach 70% test coverage",
        "Add Storybook for component documentation",
        "Adopt SWR/React Query for client-side data fetching",
        "Add structured logging with Axiom or Better Stack",
        "Document incident response plan",
    ]),
]
for phase, items in priorities:
    doc.add_heading(phase, level=3)
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

# ============================================================
# SAVE
# ============================================================
output_path = os.path.expanduser("~/Downloads/PIN_Dashboard_Audit_Report.docx")
doc.save(output_path)
print(f"Report saved to: {output_path}")
