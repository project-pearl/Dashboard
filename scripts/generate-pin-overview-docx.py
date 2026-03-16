"""Generate PIN Platform Overview .docx — drops into ~/Downloads"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os, datetime

doc = Document()

# -- Page margins --
for section in doc.sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# -- Base style --
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(10.5)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.space_before = Pt(0)

NAVY = RGBColor(0x0a, 0x1a, 0x3a)
DARK = RGBColor(0x1a, 0x1a, 0x2e)
GRAY = RGBColor(0x55, 0x55, 0x55)
ACCENT = RGBColor(0x00, 0x6d, 0x77)

def heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = NAVY if level == 1 else DARK
    return h

def para(text, bold=False, italic=False, size=None, color=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    if bold: r.bold = True
    if italic: r.italic = True
    if size: r.font.size = Pt(size)
    if color: r.font.color.rgb = color
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    if level > 0:
        p.paragraph_format.left_indent = Cm(1.2 * level)
    return p

def add_table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Light Grid Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
    for row_data in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row_data):
            cells[i].text = str(val)
            for p in cells[i].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
    return t

def add_contrast(before_text, after_text):
    """Add a Before/With PIN contrast block"""
    p = doc.add_paragraph()
    r = p.add_run('Before PIN: ')
    r.bold = True
    r.font.color.rgb = RGBColor(0x99, 0x33, 0x33)
    r.font.size = Pt(10)
    r2 = p.add_run(before_text)
    r2.font.size = Pt(10)
    p2 = doc.add_paragraph()
    r3 = p2.add_run('With PIN: ')
    r3.bold = True
    r3.font.color.rgb = RGBColor(0x00, 0x6d, 0x33)
    r3.font.size = Pt(10)
    r4 = p2.add_run(after_text)
    r4.font.size = Pt(10)
    doc.add_paragraph()  # spacer


# ================================================================
# TITLE PAGE
# ================================================================
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = t.add_run('PIN')
run.font.size = Pt(48)
run.bold = True
run.font.color.rgb = NAVY

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run('PEARL Intelligence Network')
r.font.size = Pt(24)
r.font.color.rgb = DARK

doc.add_paragraph()
sub2 = doc.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub2.add_run('Comprehensive Platform Overview')
r.font.size = Pt(16)
r.font.color.rgb = GRAY

doc.add_paragraph()
doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = meta.add_run(f'March 2026  |  Version 1.0\nUNCLASSIFIED \u2014 All data sources are publicly available federal APIs')
r.font.size = Pt(10)
r.font.color.rgb = GRAY

doc.add_page_break()

# ================================================================
# TABLE OF CONTENTS (manual)
# ================================================================
heading('Table of Contents')
toc_items = [
    'Executive Summary',
    'Platform Architecture Overview',
    'Complete Feature Inventory',
    '    1. Water Quality Monitoring',
    '    2. Regulatory Compliance & Enforcement',
    '    3. Infrastructure & Facilities',
    '    4. Health & Public Safety',
    '    5. Climate, Weather & Environment',
    '    6. Military & National Security',
    '    7. Environmental Justice & Demographics',
    '    8. Satellite, Ocean & Remote Sensing',
    '    9. PFAS & Emerging Contaminants',
    '    10. Funding, Grants & Governance',
    '    11. Agriculture & Land Use',
    '    12. AI-Powered Intelligence',
    '    13. Sentinel Anomaly Detection System',
    '    14. User Roles & Lens System',
    '    15. Composite Water Quality Scoring',
    '    16. Reporting & Export',
    'What PIN Makes Available That Was Never Before Accessible',
    'Correlations Never Before Possible',
    'Appendix: Federal Data Source Inventory',
]
for item in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(item)
    r.font.size = Pt(10)
    if not item.startswith('    '):
        r.bold = True

doc.add_page_break()

# ================================================================
# EXECUTIVE SUMMARY
# ================================================================
heading('Executive Summary')

para(
    'The PEARL Intelligence Network (PIN) is the first unified water quality intelligence platform '
    'that fuses data from 80+ federal agencies and programs into a single, continuously updated '
    'operational picture. Before PIN, the nation\'s water quality data existed in dozens of incompatible '
    'databases \u2014 EPA\'s SDWIS for drinking water, USGS NWIS for streamflow, CDC\'s surveillance systems '
    'for health outcomes, NOAA for climate, DoD for military installations \u2014 each requiring separate '
    'credentials, APIs, query languages, and domain expertise. No single analyst, agency, or tool could '
    'see across all of them simultaneously.'
)

para(
    'PIN eliminates that fragmentation. It operates 120+ automated cache modules fed by 90 staggered '
    'cron jobs that continuously ingest, normalize, spatially index, and cross-reference data from every '
    'major federal water, environmental, health, and infrastructure program. The result is a living '
    'intelligence layer that enables correlations, predictions, and threat detections that were literally '
    'impossible when these datasets existed in isolation.'
)

doc.add_paragraph()
heading('Key Metrics', level=2)
add_table(
    ['Metric', 'Value'],
    [
        ('Federal data sources integrated', '80+'),
        ('Automated cache modules', '120+'),
        ('Scheduled cron pipelines', '90'),
        ('User roles with tailored AI behavior', '17'),
        ('Dashboard lens configurations', '198+'),
        ('Unique dashboard section types', '127+'),
        ('Sentinel anomaly adapters', '19'),
        ('Compound threat detection patterns', '13'),
        ('Composite water quality index layers', '9'),
        ('Cross-agency correlation discoveries', '6 automated + 7 derived'),
        ('Spatial resolution', '0.1\u00b0 grid (~11 km cells)'),
        ('Real-time data refresh', 'Every 5 minutes (streamflow, sentinel)'),
    ]
)

doc.add_paragraph()
heading('Core Innovation', level=2)
para(
    'PIN\'s breakthrough is real-time cross-agency data fusion with AI-powered analysis. It doesn\'t merely '
    'aggregate data \u2014 it spatially indexes every record onto a uniform 0.1\u00b0 grid, enabling any data point '
    'from any agency to be correlated with any other data point in the same geographic area. When a USGS gauge '
    'detects anomalous chemistry, PIN can instantly cross-reference it with upstream EPA discharge permits, '
    'ATTAINS impairment records, CDC health surveillance data, NWS weather forecasts, FEMA flood zones, and '
    'EJScreen demographic overlays \u2014 a synthesis that previously would have taken an analyst weeks of manual '
    'cross-referencing across a dozen separate systems.'
)

doc.add_page_break()

# ================================================================
# ARCHITECTURE
# ================================================================
heading('Platform Architecture Overview')

heading('Technology Stack', level=2)
para(
    'PIN is built on Next.js deployed to Vercel\'s serverless infrastructure, providing auto-scaling, '
    'edge caching, and zero-downtime deployments. The architecture is designed for government-grade '
    'reliability with consumer-grade responsiveness.'
)

heading('Grid-Based Spatial Indexing', level=2)
para(
    'Every data record ingested by PIN is assigned to a 0.1\u00b0 latitude/longitude grid cell (~11 km '
    'resolution at mid-latitudes). This uniform spatial index is the foundation of all cross-agency '
    'correlation \u2014 a hospital, a USGS gauge, an NPDES permit, and an EJScreen block group in the same '
    'grid cell are automatically linkable without requiring explicit geographic joins. Adjacent-cell queries '
    'enable proximity analysis up to configurable radii (typically 15\u201340 km depending on the analysis type).'
)

heading('Dual-Persistence Caching', level=2)
para('Every cache module implements a two-tier persistence strategy for cold-start survival:')
bullet('Disk cache (.cache/ directory) \u2014 fast local reads during normal operation')
bullet('Vercel Blob \u2014 durable cloud storage that survives serverless function recycling')
para(
    'On startup, each module tries disk first, then falls back to Blob if empty. This ensures the platform '
    'recovers to full operational state within seconds of a cold start, even after a deployment or '
    'infrastructure restart.'
)

heading('Automated Data Pipeline', level=2)
para(
    'PIN\'s 90 cron jobs are staggered across a 3:00 AM \u2013 11:00 PM UTC daily window to avoid API rate '
    'limits and distribute compute load:'
)
bullet('Every 5 minutes: USGS real-time streamflow (NWIS IV), Sentinel anomaly polling, Sentinel scoring, alert dispatch')
bullet('Every 10 minutes: NWS weather alerts, military installation threat assessment')
bullet('Every 15 minutes: Embassy air quality, burn pit assessment, real-time weather')
bullet('Every 30 minutes: ATTAINS assessments, NWPS flood data, air quality, seismic activity')
bullet('Every 6 hours: AI insight generation, CO-OPS tidal data, NWM river forecasts, USBR reservoirs, HEFS forecasts, NWS 7-day forecasts')
bullet('Daily (staggered 3 AM\u201311 PM): 60+ full cache rebuilds covering all major federal APIs')
bullet('Weekly (Sunday): Heavy-compute caches \u2014 RCRA, SEMS, EJScreen, climate normals, NARS, demographics, infrastructure, cyber risk')

para(
    'All cron jobs include build locks (auto-clearing after 12 minutes), empty-data guards (skip writes when '
    '0 records to preserve last-known-good data), and retry logic with exponential backoff for failed states.'
)

heading('AI Layer', level=2)
bullet('GPT-4o \u2014 powers interactive "Ask PIN" conversational Q&A across all 17 user roles, briefing Q&A, and resolution plan generation')
bullet('GPT-4o-mini \u2014 handles high-volume batch insight generation (408 state/role combinations per cycle), operating at temperature 0.3 to minimize hallucination')
para(
    'The AI layer receives live, grounded context from 14 domain retrievers that query 44+ cache modules '
    'in real-time, ensuring every AI response reflects the latest data rather than stale training knowledge.'
)

doc.add_page_break()

# ================================================================
# COMPLETE FEATURE INVENTORY
# ================================================================
heading('Complete Feature Inventory')

# --- 1. WATER QUALITY ---
heading('1. Water Quality Monitoring', level=2)
add_table(
    ['Capability', 'Data Source', 'Cache Module', 'Frequency'],
    [
        ('Discrete water quality samples (nutrients, metals, bacteria)', 'EPA Water Quality Portal', 'wqpCache.ts', 'Daily'),
        ('Real-time streamflow & water quality (discharge, DO, pH, turbidity)', 'USGS NWIS Instantaneous Values', 'nwisIvCache.ts', '5 min'),
        ('Daily statistical streamflow values', 'USGS NWIS Daily Values', 'usgsDvCache.ts', 'Daily'),
        ('Groundwater level monitoring', 'USGS NWIS Groundwater', 'nwisGwCache.ts', 'Daily'),
        ('Waterbody impairment assessments & TMDLs', 'EPA ATTAINS', 'attainsCache.ts', '30 min'),
        ('California-specific water quality', 'CEDEN', 'cedenCache.ts', 'Daily'),
        ('Beach water quality advisories', 'EPA BEACON', 'Sentinel adapter', '5 min'),
        ('Harmful algal bloom observations', 'NOAA HAB-SOS', 'habsosCache.ts', 'Daily'),
        ('HAB forecasts', 'NOAA HAB Forecast', 'habForecastCache.ts', 'Daily'),
        ('Hypoxia (dead zone) monitoring', 'Various', 'hypoxiaCache.ts', 'Daily'),
        ('WQX modern results', 'EPA WQX', 'wqxModernCache.ts', 'Daily'),
        ('USGS WQP cross-reference', 'USGS/WQP', 'usgsWqpCache.ts', 'Daily'),
        ('OGC API monitoring locations', 'USGS OGC', 'usgsOgcCache.ts', 'Daily'),
        ('Real-time threshold alerts', 'USGS', 'nwisAlertCache.ts', 'Daily'),
    ]
)
doc.add_paragraph()
para(
    'Key capability: PIN\'s discharge signature analysis detects multi-parameter anomalies in real-time \u2014 a '
    'simultaneous E. coli spike + dissolved oxygen crash + turbidity surge is automatically flagged as a probable '
    'sewage discharge event, a correlation invisible when monitoring parameters individually.',
    italic=True, size=9.5
)

# --- 2. COMPLIANCE ---
heading('2. Regulatory Compliance & Enforcement', level=2)
add_table(
    ['Capability', 'Data Source', 'Cache Module', 'Frequency'],
    [
        ('Drinking water violations & enforcement', 'EPA SDWIS', 'sdwisCache.ts', 'Daily'),
        ('NPDES discharge permits & compliance', 'EPA ECHO', 'echoCache.ts', 'Daily'),
        ('Detailed enforcement actions', 'EPA ICIS', 'icisCache.ts', 'Daily'),
        ('Air emission compliance', 'EPA ICIS-Air', 'icisAirCache.ts', 'Daily'),
        ('Toxic chemical releases', 'EPA TRI', 'triCache.ts', 'Daily'),
        ('Hazardous waste facilities & violations', 'EPA RCRA', 'rcraCache.ts', 'Weekly'),
        ('Discharge monitoring report violations', 'EPA ECHO DMR', 'echoDmrViolationsCache.ts', 'Daily'),
        ('Effluent monitoring data', 'EPA ECHO Effluent', 'echoEffluentCache.ts', 'Daily'),
        ('MS4 stormwater permits', 'EPA ECHO', 'ms4PermitCache.ts', 'Daily'),
        ('Sewer overflows (SSO/CSO)', 'EPA', 'ssoCsoCache.ts', 'Daily'),
        ('Pesticide registrations', 'EPA OPP', 'epaOppPesticideCache.ts', 'Daily'),
        ('OSHA water-related violations', 'OSHA', 'oshaWaterCache.ts', 'Daily'),
    ]
)

# --- 3. INFRASTRUCTURE ---
heading('3. Infrastructure & Facilities', level=2)
add_table(
    ['Capability', 'Data Source', 'Cache Module', 'Frequency'],
    [
        ('Dam inventory & hazard classification', 'USACE NID', 'damCache.ts', 'Daily'),
        ('Army Corps projects & operations', 'USACE', 'usaceCache.ts', '6-hourly'),
        ('Bureau of Reclamation reservoirs', 'USBR', 'usbrCache.ts', '6-hourly'),
        ('Facility Registry Service (all EPA-regulated)', 'EPA FRS', 'frsCache.ts', 'Daily'),
        ('Water infrastructure data', 'BWB', 'bwbCache.ts', 'Daily'),
        ('Superfund NPL sites', 'EPA Superfund', 'superfundCache.ts', 'Daily'),
        ('SEMS contaminated sites', 'EPA SEMS', 'semsCache.ts', 'Weekly'),
        ('Basin characteristics', 'USGS StreamStats', 'streamStatsCache.ts', 'Daily'),
        ('Pipeline incidents', 'PHMSA', 'phmsaPipelineCache.ts', 'Daily'),
    ]
)

# --- 4. HEALTH ---
heading('4. Health & Public Safety', level=2)
add_table(
    ['Capability', 'Data Source', 'Cache Module', 'Frequency'],
    [
        ('Hospital facilities (capacity, quality ratings)', 'HealthData.gov CMS', 'hospitalCache.ts', 'Daily'),
        ('Healthcare shortage areas', 'HRSA HPSA', 'hrsaHpsaCache.ts', 'Daily'),
        ('Waterborne illness surveillance', 'CDC', 'waterborneIllnessCache.ts', 'Daily'),
        ('Mortality/morbidity correlation', 'CDC WONDER', 'cdcWonderCache.ts', 'Daily'),
        ('Environmental health tracking', 'CDC Tracking Network', 'environmentalTrackingCache.ts', 'Daily'),
        ('Hospital capacity & outbreak surveillance', 'HealthData.gov Socrata', 'healthDataGovCache.ts', 'Daily'),
        ('Drug/device enforcement, recalls', 'FDA Open FDA', 'openFDACache.ts', 'Daily'),
        ('Chronic disease prevalence', 'CDC PLACES', 'cdcPlacesCache.ts', 'Daily'),
        ('Broader CDC datasets', 'data.cdc.gov', 'dataCdcGovCache.ts', 'Daily'),
        ('Wastewater pathogen surveillance', 'CDC NWSS', 'cdcNwssCache.ts', 'Daily'),
        ('Health literacy & wellness', 'HHS MyHealthfinder', 'myhealthfinderCache.ts', 'Daily'),
        ('Toxic substance profiles', 'ATSDR', 'atsdrToxicologyCache.ts', 'Daily'),
    ]
)

# --- 5. CLIMATE ---
heading('5. Climate, Weather & Environment', level=2)
add_table(
    ['Capability', 'Data Source', 'Cache Module', 'Frequency'],
    [
        ('Real-time weather alerts', 'NWS', 'nwsAlertCache.ts', '10 min'),
        ('7-day weather forecasts', 'NWS', 'nwsForecastCache.ts', '6 hours'),
        ('Climate normals', 'NOAA CDO', 'climateNormalsCache.ts', 'Weekly'),
        ('Drought monitoring', 'USDM', 'usdmCache.ts', 'Daily'),
        ('FEMA disaster declarations & risk', 'FEMA', 'femaCache.ts', 'Daily'),
        ('NFIP flood insurance claims', 'FEMA NFIP', 'nfipClaimsCache.ts', 'Daily'),
        ('Hazard mitigation grants', 'FEMA', 'hazMitCache.ts', 'Weekly'),
        ('Flood stage forecasts', 'NOAA NWPS', 'nwpsCache.ts', '30 min'),
        ('National Water Model forecasts', 'NOAA NWM', 'nwmCache.ts', '6 hours'),
        ('Ensemble flood forecasts', 'NOAA HEFS', 'hefsCache.ts', '6 hours'),
        ('Flood event high-water marks', 'USGS STN', 'stnFloodCache.ts', 'Daily'),
        ('NEXRAD precipitation estimates', 'NOAA', 'nexradQpeCache.ts', 'Daily'),
        ('Historical climate data', 'NOAA NCEI', 'nceiCache.ts', 'Daily'),
        ('Snowpack monitoring', 'NRCS SNOTEL', 'snotelCache.ts', 'Daily'),
        ('Severe weather events', 'NOAA SPC', 'swdiCache.ts', 'Daily'),
        ('Wildfire detection', 'NASA FIRMS', 'firmsCache.ts', '4 hours'),
        ('Seismic activity', 'USGS Earthquake', 'seismicCache.ts', '30 min'),
        ('Volcanic activity', 'USGS Volcano', 'volcanoCache.ts', 'Daily'),
        ('ERA5 climate reanalysis', 'Copernicus CDS', 'copernicusCdsCache.ts', 'Weekly'),
        ('Global freshwater quality', 'GEMStat (UNEP)', 'gemstatCache.ts', 'Weekly'),
        ('Power sector emissions', 'EPA CAMPD', 'campdCache.ts', 'Daily'),
        ('Air quality monitoring', 'Multiple', 'airQualityCache.ts', '30 min'),
        ('Embassy air quality', 'State Dept', 'embassyAqiCache.ts', '15 min'),
        ('Flood impact (derived)', 'NWPS + NWM + FRS', 'floodImpactCache.ts', 'Daily'),
    ]
)

# --- 6. MILITARY ---
heading('6. Military & National Security', level=2)
add_table(
    ['Capability', 'Data Source', 'Cache Module', 'Frequency'],
    [
        ('DoD PFAS installation assessments', 'Curated DoD data', 'dodPfasCache.ts', 'Daily'),
        ('DoD PFAS investigation sites', 'Curated DoD data', 'dodPfasSitesCache.ts', 'Daily'),
        ('Military installation proximity analysis', 'Static + derived', 'Cross-cache', 'Continuous'),
        ('Water utility cyber risk scoring', 'Derived (ECHO/SDWIS/FRS)', 'cyberRiskCache.ts', 'Weekly'),
        ('Installation threat assessment', 'Real-time derived', 'Cron', '10 min'),
        ('Burn pit proximity assessment', 'Derived', 'Cron', '15 min'),
    ]
)

# --- 7. EJ ---
heading('7. Environmental Justice & Demographics', level=2)
add_table(
    ['Capability', 'Data Source', 'Cache Module', 'Frequency'],
    [
        ('Block-group EJ demographics', 'EPA EJScreen (Harvard DataVerse)', 'ejscreenCache.ts', 'Weekly'),
        ('Environmental health risk scoring', 'EPA EJScreen + CDC Tracking', 'environmentalHealthCache.ts', 'Daily'),
        ('Healthcare shortage correlation', 'HRSA HPSA', 'hrsaHpsaCache.ts', 'Daily'),
        ('Community vulnerability scoring', 'Derived (multi-source)', 'Cross-cache', 'Continuous'),
        ('Environmental Justice Index', 'CDC EJI', 'National summary', 'Continuous'),
        ('Social Vulnerability Index', 'CDC SVI', 'National summary', 'Continuous'),
    ]
)

# --- 8. SATELLITE ---
heading('8. Satellite, Ocean & Remote Sensing', level=2)
add_table(
    ['Capability', 'Data Source', 'Cache Module', 'Frequency'],
    [
        ('Ocean buoy observations', 'NOAA NDBC', 'ndbcCache.ts', 'Daily'),
        ('Tidal observations', 'NOAA CO-OPS', 'coopsCache.ts', '6 hours'),
        ('Derived tidal analytics', 'NOAA CO-OPS', 'coopsDerivedCache.ts', 'Weekly'),
        ('Great Lakes conditions', 'NOAA GLERL', 'glerlCache.ts', 'Daily'),
        ('Satellite ocean color & SST', 'NOAA CoastWatch', 'coastwatchCache.ts', 'Daily'),
        ('ERDDAP satellite data', 'NOAA ERDDAP', 'erddapSatCache.ts', 'Daily'),
        ('NASA Earth observation', 'NASA CMR', 'nasaCmrCache.ts', 'Daily'),
        ('NASA streamflow data', 'NASA', 'nasaStreamCache.ts', 'Weekly'),
        ('Species habitat', 'USFWS IPAC', 'ipacCache.ts', 'Weekly'),
    ]
)

# --- 9. PFAS ---
heading('9. PFAS & Emerging Contaminants', level=2)
add_table(
    ['Capability', 'Data Source', 'Cache Module', 'Frequency'],
    [
        ('UCMR5 PFAS detections', 'EPA UCMR', 'pfasCache.ts', 'Daily'),
        ('ECHO PFAS facility analytics', 'EPA ECHO', 'epaPfasAnalyticsCache.ts', 'Daily'),
        ('DoD PFAS assessments (50 installations)', 'Curated DoD', 'dodPfasCache.ts', 'Daily'),
        ('DoD investigation sites', 'Curated DoD', 'dodPfasSitesCache.ts', 'Daily'),
        ('Atmospheric PFAS deposition', 'NADP', 'nadpPfasCache.ts', 'Daily'),
    ]
)

# --- 10. FUNDING ---
heading('10. Funding, Grants & Governance', level=2)
add_table(
    ['Capability', 'Data Source', 'Cache Module', 'Frequency'],
    [
        ('Federal spending on water programs', 'USASpending.gov', 'usaSpendingCache.ts', 'Weekly'),
        ('Grant opportunities', 'Grants.gov', 'grantsGovCache.ts', 'Daily'),
        ('Federal contractor data', 'SAM.gov', 'samGovCache.ts', 'Weekly'),
        ('Open government datasets', 'Data.gov', 'dataGovCache.ts', 'Weekly'),
        ('Congressional legislation tracking', 'Congress.gov', 'congressCache.ts', 'Daily'),
        ('Advocacy & stakeholder tracking', 'Internal', 'advocacyCache.ts', 'Daily'),
    ]
)

# --- 11. AGRICULTURE ---
heading('11. Agriculture & Land Use', level=2)
add_table(
    ['Capability', 'Data Source', 'Cache Module', 'Frequency'],
    [
        ('Livestock operations', 'USDA NASS', 'nassLivestockCache.ts', 'Weekly'),
        ('Crop data', 'USDA NASS', 'nassCropsCache.ts', 'Weekly'),
        ('Soil characteristics', 'USDA SSURGO', 'ssurgoCache.ts', 'Daily'),
        ('Land cover classification', 'USGS NLCD', 'nlcdCache.ts', 'Daily'),
        ('National aquatic resource surveys', 'EPA NARS', 'narsCache.ts', 'Weekly'),
        ('HUC-8 water budgets + drought', 'USGS/USDM', 'usgsWaterAvailCache.ts', 'Weekly'),
        ('Multi-agency groundwater', 'USGS NGWMN/CIDA', 'ngwmnCache.ts', 'Daily'),
    ]
)

doc.add_page_break()

# --- 12. AI ---
heading('12. AI-Powered Intelligence', level=2)

heading('Ask PIN (Universal Q&A)', level=3)
bullet('Available across all 17 user roles with role-tailored tone and priorities')
bullet('14 domain retrievers query 44+ cache modules to build grounded context for every question')
bullet('Domains: PFAS, Groundwater, Compliance, Health, Military, Climate, Realtime, Superfund, Infrastructure, Environmental Justice, Water Quality, Stormwater, Correlations, baseline state reports')
bullet('Each role receives a customized system prompt (e.g., K-12 gets educational framing; Federal gets CWA implications; Military gets base security threats)')

heading('AI Insights (Batch Generation)', level=3)
bullet('Pre-generated for every state across all applicable roles on 6-hour cycles')
bullet('Delta detection via signalsHash \u2014 only regenerates when underlying data changes')
bullet('Semaphore-based concurrency (4 states in parallel) with exponential backoff on rate limits')
bullet('Each insight classified by type (predictive, anomaly, comparison, recommendation, summary) and severity')

heading('Briefing Q&A', level=3)
bullet('Specialized conversational AI for executive briefing scenarios')
bullet('Tailored prompts for Federal, Federal+Military, State, Local, and MS4 roles')
bullet('Grounded in live cache data, not static reports')

heading('Resolution Plans & Correlation Discovery', level=3)
bullet('AI-generated remediation and compliance action plans based on specific violations and conditions')
bullet('6 automated cross-agency correlation algorithms (detailed in Correlations section)')

doc.add_page_break()

# --- 13. SENTINEL ---
heading('13. Sentinel Anomaly Detection System', level=2)
para(
    'The Sentinel system is PIN\'s real-time threat detection engine, operating continuously with 5-minute polling cycles.'
)

heading('19 Real-Time Adapters', level=3)
add_table(
    ['Adapter', 'Source', 'Monitors'],
    [
        ('NWS', 'NWS Alerts', 'Severe weather affecting water infrastructure'),
        ('Air Quality', 'Multiple', 'AQI exceedances affecting watersheds'),
        ('NWPS Flood', 'NOAA NWPS', 'Active flood warnings at stream gauges'),
        ('NWPS Forecast', 'NOAA NWPS', 'Predicted flood conditions'),
        ('USGS IV', 'USGS NWIS', 'Real-time streamflow/chemistry anomalies'),
        ('SSO/CSO', 'EPA', 'Sewer overflow events'),
        ('NPDES DMR', 'EPA ICIS', 'Discharge monitoring exceedances'),
        ('QPE Rainfall', 'NOAA NEXRAD', 'Extreme precipitation events'),
        ('ATTAINS', 'EPA', 'New impairment designations'),
        ('State Discharge', 'State agencies', 'State-reported discharge anomalies'),
        ('FEMA', 'FEMA', 'Disaster declarations'),
        ('ECHO Enforcement', 'EPA ECHO', 'New enforcement actions'),
        ('CDC NWSS', 'CDC', 'Wastewater pathogen detections'),
        ('HAB SOS', 'NOAA', 'Harmful algal bloom events'),
        ('EPA BEACON', 'EPA', 'Beach pathogen advisories'),
        ('TRI', 'EPA TRI', 'Toxic chemical releases'),
        ('RCRA', 'EPA RCRA', 'Hazardous waste violations'),
        ('SEMS', 'EPA SEMS', 'Superfund site activity'),
        ('CAMPD', 'EPA CAMPD', 'Power sector emission events'),
    ]
)

doc.add_paragraph()
heading('13 Compound Pattern Detectors', level=3)
add_table(
    ['Pattern', 'Sources Crossed', 'Multiplier', 'Description'],
    [
        ('Potomac Multi-Hazard', 'NWS + NWPS + SSO/NPDES', '2.5x', 'Multi-hazard flooding with sewage impact'),
        ('Infrastructure Stress', 'SSO + QPE/NWPS', '2.0x', 'Sewer overflows triggered by rainfall'),
        ('Spreading Contamination', 'NPDES + SSO + ECHO', '3.0x', 'Contamination migrating across watersheds'),
        ('Regulatory Escalation', 'NPDES + ECHO', '1.8x', 'Enforcement on discharge violations'),
        ('Enforcement Cascade', 'ECHO + NPDES', '2.2x', 'EPA enforcement cascade across regions'),
        ('Bio-Threat Correlation', 'CDC NWSS + USGS/SSO/NPDES', '3.5x', 'Wastewater pathogen + WQ changes'),
        ('Flood Prediction Cascade', 'NWPS Forecast + QPE/NWPS', '3.0x', 'Predicted floods confirmed by observations'),
        ('Airborne Public Health', 'Air Quality + NWS', '2.0x', 'Air quality + severe weather compound risk'),
        ('Predicted Infra Failure', 'NWPS Forecast + SSO/NPDES', '2.5x', 'Infrastructure stress during forecast flood'),
        ('HAB-WQ Correlation', 'HABSOS + USGS IV', '2.5x', 'Algal blooms + water quality changes'),
        ('Beach Pathogen-WQ', 'EPA BEACON + USGS/CDC', '2.0x', 'Beach pathogens + upstream WQ signals'),
        ('Toxic Release Cascade', 'TRI + NPDES', '2.5x', 'Toxic release followed by discharge changes'),
        ('Contamination Cluster', 'SEMS/RCRA + ECHO', '3.0x', 'Superfund/RCRA + enforcement clustering'),
    ]
)

doc.add_paragraph()
heading('Coordination Detection (Multi-Site Attack Identification)', level=3)
bullet('Three-layer spatial-temporal clustering at HUC-6 watershed level')
bullet('6-hour coordination window for identifying simultaneous events')
bullet('Attack indicator parameter pairs (e.g., conductivity + pH = industrial discharge signature)')
bullet('Scoring: cluster size (35%) + parameter breadth (35%) + temporal tightness (30%)')
bullet('Threshold: coordination score > 0.6 triggers CoordinatedEvent alert')
bullet('CBRN classification: chemical, biological, radiological, nuclear categorization')

heading('Tier-2 Scoring Engine', level=3)
bullet('5 severity tiers: NOMINAL (<30) \u2192 ADVISORY (30\u201374) \u2192 WATCH (75\u2013149) \u2192 CRITICAL (150\u2013299) \u2192 ANOMALY (300+)')
bullet('Time-decay scoring with 48-hour window (floor: 0.1 \u2014 events never fully zero out)')
bullet('Adjacent HUC bonus: 1.5x multiplier for geographically nearby watershed events')
bullet('Per-HUC-8 persistent scoring with compound pattern multipliers')

doc.add_page_break()

# --- 14. ROLES ---
heading('14. User Roles & Lens System', level=2)
para('PIN serves 17 distinct user roles, each with tailored dashboard configurations, AI behavior, and data prioritization:')

add_table(
    ['Role', 'Primary Audience', 'AI Tone Focus'],
    [
        ('Federal', 'EPA leadership, federal analysts', 'Cross-state patterns, national trends, CWA implications'),
        ('Federal+Military', 'Base commanders, DoD water security', 'PFAS proximity, CISA advisories, force readiness'),
        ('State', 'State environmental program managers', 'Statewide compliance, TMDL progress, resource allocation'),
        ('Local', 'Municipal/county coordinators', 'Jurisdiction deadlines, council briefing points'),
        ('MS4', 'Stormwater permit holders', 'MCM deliverables, BMP inspection, audit readiness'),
        ('K-12', 'Students and teachers', 'Fun facts, wildlife impacts, educational framing'),
        ('College', 'University researchers', 'Research anomalies, methodology rigor, publication readiness'),
        ('Researcher', 'Water quality scientists', 'Statistical patterns, data quality, peer-comparable findings'),
        ('Corporate', 'ESG/sustainability teams', 'Portfolio risk, ESG disclosure, supply chain water risk'),
        ('NGO', 'Nonprofit advocates', 'Environmental justice, community health, advocacy'),
        ('Utility', 'Water/wastewater operators', 'Source water, treatment challenges, distribution compliance'),
        ('Biotech', 'Pharma/biotech operations', 'GMP water purity, USP standards, process water'),
        ('Investor', 'Financial analysts', 'Infrastructure investment, regulatory risk, utility financials'),
        ('Agriculture', 'Agricultural operators', 'Irrigation quality, nutrient runoff, ag BMPs'),
        ('Lab', 'Laboratory operations', 'QA/QC protocols, method compliance, data validation'),
        ('PEARL', 'Platform operations', 'System health, cache status, pipeline integrity'),
        ('Temp', 'General public', 'Clear, helpful water quality answers'),
    ]
)

doc.add_paragraph()
para(
    'The dashboard lens system provides 198+ configurable views across these roles, with each lens controlling '
    'which of 127+ section types are visible. The Federal Management Center alone offers 18 lens views (Overview, '
    'Briefing, Emergency Response, Policy, Compliance, Water Quality, Public Health, Habitat, Agriculture, '
    'Infrastructure, Monitoring, Disaster, TMDL, Scorecard, Reports, Permits, Funding, and more). Each section '
    'is individually draggable and configurable via the LayoutEditor.'
)

# --- 15. COMPOSITE SCORING ---
heading('15. Composite Water Quality Scoring', level=2)
para('PIN computes a 9-layer weighted composite score (0\u2013100) for every HUC-8 watershed in the United States:')

add_table(
    ['Layer', 'Weight', 'Measures', 'Key Sources'],
    [
        ('PEARL Load Velocity', '15%', 'Nutrient loading trends & severity', 'WQP, ICIS DMR'),
        ('Infrastructure Failure Probability', '14%', 'Drinking water system compliance risk', 'SDWIS, ICIS'),
        ('Per Capita Load Contribution', '13%', 'Population-normalized discharge burden', 'SDWIS, WQP, ICIS DMR'),
        ('Permit Risk Exposure', '12%', 'NPDES permit compliance risk', 'ICIS, ECHO DMR'),
        ('Ecological Health', '12%', 'T&E species risk & recovery pressure', 'USFWS ECOS, ATTAINS'),
        ('Watershed Recovery Rate', '10%', 'Impairment status & TMDL progress', 'ATTAINS'),
        ('Waterfront Value Exposure', '8%', 'Economic risk from WQ degradation', 'Census, NOAA, WQP, ATTAINS'),
        ('EJ Vulnerability', '8%', 'Environmental justice burden', 'CDC EJI/SVI, Census, SDWIS'),
        ('Governance Response', '8%', 'Regulatory oversight gaps', 'ATTAINS, ICIS, SDWIS'),
    ]
)

doc.add_paragraph()
para(
    'Each layer includes confidence scoring (data density 30%, recency 45%, source diversity 25%) and '
    'trend detection (improving/stable/declining). Low-confidence scores are automatically regressed toward '
    'neutral to prevent misleading results. The composite enables national comparative analytics, state '
    'rankings, and 7-day/30-day score projections.'
)

# --- 16. REPORTING ---
heading('16. Reporting & Export', level=2)
bullet('State Reports \u2014 Aggregated per-state compliance, violation, and impairment summaries')
bullet('Federal Briefing Metrics \u2014 National summary with cross-state patterns and worst-state identification')
bullet('Compliance Scorecards \u2014 Multi-category grades (regulatory, operational, infrastructure, financial)')
bullet('PDF/DOCX Generation \u2014 Exportable formatted reports for offline distribution')
bullet('Data Export Hub \u2014 Configurable data extraction across all cached sources')
bullet('Delta Changelog \u2014 Track what changed since last review')

doc.add_page_break()

# ================================================================
# SECTION 1: WHAT'S NEW
# ================================================================
heading('What PIN Makes Available That Was Never Before Accessible')
para(
    'This section documents capabilities that did not exist in any form before PIN \u2014 not improvements to '
    'existing tools, but entirely new intelligence that was impossible when federal water data existed in silos.',
    italic=True, size=10
)

# 1.1
heading('1. Unified Real-Time Federal Data Access', level=2)
add_contrast(
    'An analyst investigating water quality in a single watershed needed to separately query EPA ECHO '
    '(permit compliance), EPA SDWIS (drinking water violations), USGS NWIS (streamflow), EPA ATTAINS '
    '(impairment status), EPA WQP (lab results), USACE NID (dam conditions), EPA FRS (facility locations), '
    'CDC (health data), NOAA (weather), and FEMA (flood risk). Each system has different APIs, query languages, '
    'date formats, coordinate systems, and access methods. A comprehensive assessment of a single watershed '
    'could take days of manual data gathering.',
    'A single dashboard query returns all data from all sources for any geographic area, pre-correlated on '
    'a uniform spatial grid, updated continuously. What took days now takes seconds.'
)

# 1.2
heading('2. Multi-Parameter Discharge Signature Detection', level=2)
add_contrast(
    'USGS real-time data was available site-by-site through the NWIS web interface. Detecting that a '
    'simultaneous E. coli spike + dissolved oxygen crash + turbidity surge at the same gauge station '
    'indicates a probable sewage discharge event required a human expert monitoring multiple parameters '
    'simultaneously \u2014 something that effectively never happened at scale.',
    'The system automatically evaluates multi-parameter combinations every 5 minutes across all monitored '
    'gauges. Discharge signatures, chemical dump indicators (conductivity + pH), and sediment load events '
    '(DO + turbidity) are detected within minutes, triggering Sentinel alerts with severity scoring. '
    'PIN monitors 4 specific attack indicator parameter pairs with configurable weights.'
)

# 1.3
heading('3. Cross-Agency Health-Environment Correlation', level=2)
add_contrast(
    'CDC health surveillance data (WONDER mortality, NWSS wastewater pathogens, PLACES chronic disease) '
    'and EPA environmental data (SDWIS violations, EJScreen demographics, ECHO compliance) existed in '
    'completely separate institutional universes. No tool, database, or analyst could simultaneously '
    'query both and perform spatial correlation.',
    'Health and environmental data are co-located on the same spatial grid, enabling automatic correlation: '
    'CDC WONDER mortality rates overlaid on SDWIS violation density; hospital proximity to active violations; '
    'CDC PLACES chronic disease prevalence correlated with environmental health risk scores; waterborne '
    'illness outbreak data cross-referenced with upstream discharge permits.'
)

# 1.4
heading('4. Military Installation Water Security Intelligence', level=2)
add_contrast(
    'DoD water infrastructure risk assessment required manual cross-referencing of PFAS investigation data '
    '(scattered across installation-level reports), civilian drinking water system data (EPA SDWIS), '
    'CISA cyber advisories (DHS), and geographic proximity analysis. No integrated threat picture existed.',
    'Automated, continuous threat monitoring: PFAS proximity alerts for 50 high-profile installations, '
    'installation threat assessment every 10 minutes, burn pit assessment every 15 minutes, weekly cyber '
    'risk scoring, and AI-generated briefings tailored for base commanders with security-relevant framing.'
)

# 1.5
heading('5. Role-Specific Intelligence at Scale', level=2)
add_contrast(
    'Federal water data was one-size-fits-all. A K-12 teacher, a military base commander, a Wall Street '
    'ESG analyst, and an EPA enforcement officer all saw the same raw data \u2014 if they could find it at all.',
    '17 role-specific AI profiles deliver fundamentally different intelligence from the same underlying data. '
    'K-12 receives educational fun facts; Federal+Military receives base security threats and PFAS proximity; '
    'Corporate receives ESG disclosure metrics; NGO receives environmental justice and advocacy opportunities. '
    'This is not cosmetic rebranding \u2014 each role\'s AI restructures how data is interpreted and what actions '
    'are recommended.'
)

# 1.6
heading('6. Automated Anomaly Detection with Compound Pattern Recognition', level=2)
add_contrast(
    'Environmental anomalies were discovered after the fact through periodic manual review of individual '
    'agency databases. A sewer overflow during a flood event upstream of a drinking water intake would '
    'appear as three separate data points in three separate systems.',
    'The Sentinel system operates 19 real-time adapters on 5-minute polling cycles, evaluating every new '
    'data point against 13 compound pattern detectors. Bio-Threat Correlation (3.5x multiplier) detects '
    'CDC wastewater pathogens + water quality changes. The coordination detection engine identifies potential '
    'multi-site deliberate contamination by clustering simultaneous anomalies within 6-hour windows.'
)

# 1.7
heading('7. Environmental Justice as a Native Data Layer', level=2)
add_contrast(
    'EPA\'s EJScreen was a standalone mapping tool. Understanding how EJ burden relates to water quality '
    'violations, infrastructure conditions, health outcomes, and compliance enforcement required manual '
    'queries of EJScreen plus SDWIS, ECHO, ATTAINS, CDC, and HRSA, followed by manual spatial overlay.',
    'EJScreen block-group demographics are a native data layer that automatically overlays every metric. '
    'Every violation is visible in its EJ context. Every Sentinel alert includes the affected population\'s '
    'EJ vulnerability. The composite score includes a dedicated EJ Vulnerability Index (8% weight) combining '
    'CDC EJI percentiles, poverty, minority percentage, uninsured rates, and linguistic isolation.'
)

# 1.8
heading('8. Predictive Flood-Infrastructure Risk Fusion', level=2)
add_contrast(
    'NOAA flood forecasts and EPA facility data existed in separate systems. An approaching flood\'s threat '
    'to specific drinking water treatment plants, hazardous waste sites, and dam structures required manual '
    'geographic overlay during the crisis \u2014 often too late for preventive action.',
    'The Flood Impact Cache automatically cross-references NWPS flood stage forecasts with drinking water '
    'facilities, NWM river model predictions with dam conditions, HEFS ensemble forecasts with hazmat '
    'proximity, and FEMA flood zones with critical infrastructure. Preemptive risk assessment happens '
    'before the flood arrives.'
)

# 1.9
heading('9. National Comparative Water Quality Analytics', level=2)
add_contrast(
    'No system could compare water quality across states using a consistent methodology. Each state reports '
    'to different federal systems using different metrics, timeframes, and standards. Ranking states required '
    'months of manual data harmonization.',
    'The 9-layer composite scoring algorithm applies a uniform, weighted methodology to every HUC-8 '
    'watershed in the nation, enabling state-by-state rankings, national trend detection, worst-state '
    'identification, and 7-day/30-day projections based on 90-day historical trends.'
)

# 1.10
heading('10. Continuous AI Insight Generation for Every State and Role', level=2)
add_contrast(
    'Environmental intelligence was produced manually by analysts reviewing data and writing reports. '
    'The volume of data across 50 states and dozens of programs made comprehensive coverage impossible.',
    'AI insights are automatically generated for every state across all applicable roles on 6-hour cycles '
    'using delta detection. The system produces 408 state/role combinations per cycle, each classified by '
    'type and severity, creating a living intelligence feed that ensures no significant change goes unnoticed.'
)

doc.add_page_break()

# ================================================================
# SECTION 2: CORRELATIONS
# ================================================================
heading('Correlations Never Before Possible')
para(
    'PIN\'s most transformative capability is its Correlation Discovery Engine \u2014 six automated algorithms that '
    'perform cross-agency spatial joins to reveal relationships invisible when federal data existed in silos. '
    'These are not statistical approximations or literature-based associations \u2014 they are live, data-driven '
    'discoveries computed from current federal data, updated with every cache rebuild.',
    italic=True, size=10
)
doc.add_paragraph()

# 2.1
heading('1. PFAS Contamination \u00d7 Healthcare Deserts \u00d7 Environmental Justice', level=2)
para('Function: discoverPfasHealthDeserts()', bold=True, size=9.5, color=ACCENT)
para('Datasets Crossed:', bold=True)
bullet('DoD PFAS investigation sites (dodPfasSitesCache) \u2014 military PFAS contamination locations')
bullet('HRSA Health Professional Shortage Areas (hrsaHpsaCache) \u2014 healthcare-underserved communities')
bullet('EPA EJScreen (ejscreenCache) \u2014 block-group level demographic and environmental justice data')
para(
    'Reveals: Communities where three crises intersect \u2014 PFAS contamination from military installations, '
    'shortage of healthcare providers to treat PFAS-related health effects, and high environmental justice '
    'burden indicating vulnerable populations least equipped to advocate for remediation.'
)
para(
    'Algorithm: Filters to confirmed PFAS detections, finds HPSA shortage areas and high-EJ block groups '
    '(\u226580th percentile) within 30 km. Severity is CRITICAL if HPSA score \u226518 AND EJ index \u226590.'
)
para(
    'Why impossible before: DoD PFAS data, HRSA healthcare data, and EPA demographic data are maintained by '
    'three different agencies with no shared identifier or coordinate system. PIN\'s spatial grid makes this a '
    'trivial join.',
    italic=True, size=9.5
)

# 2.2
heading('2. Flood Damage \u00d7 Drinking Water Contamination \u00d7 Environmental Justice', level=2)
para('Function: discoverFloodWaterContamination()', bold=True, size=9.5, color=ACCENT)
para('Datasets Crossed:', bold=True)
bullet('FEMA NFIP flood insurance claims (nfipClaimsCache) \u2014 documented flood damage with dollar amounts')
bullet('EPA SDWIS drinking water violations (sdwisCache) \u2014 water quality failures')
bullet('EPA EJScreen (ejscreenCache) \u2014 demographic vulnerability')
para(
    'Reveals: Zones where documented flood damage co-occurs with drinking water quality failures in '
    'environmental justice communities \u2014 compound disaster zones where flood-damaged infrastructure '
    'produces contaminated drinking water for vulnerable populations.'
)
para(
    'Algorithm: Groups NFIP claims and SDWIS violations by state, finds violations within 15 km of flood '
    'centroid. Triggers if \u22653 violations AND flood damage >$100K. CRITICAL if \u22655 health-based violations.'
)

# 2.3
heading('3. Upstream Discharge Violations \u00d7 Downstream Impairment', level=2)
para('Function: discoverDischargeImpairmentLinks()', bold=True, size=9.5, color=ACCENT)
para('Datasets Crossed:', bold=True)
bullet('EPA ECHO Significant Non-Compliance facilities (echoCache) \u2014 facilities violating discharge permits')
bullet('EPA ATTAINS impaired waters (attainsCache) \u2014 waterbodies failing water quality standards')
para(
    'Reveals: Facilities with active discharge violations positioned upstream of impaired waterbodies, '
    'suggesting causal links between illegal discharges and water quality degradation. Identifies regulatory '
    'enforcement priority targets where action would have highest impact.'
)
para(
    'Algorithm: For each SNC facility, finds impaired waterbodies within 0.5\u201325 km. CRITICAL if \u22655 nearby; '
    'HIGH if \u22653.'
)

# 2.4
heading('4. Dam Failure Risk \u00d7 Downstream Hazmat \u00d7 Drinking Water Cascade', level=2)
para('Function: discoverDamCascadeRisk()', bold=True, size=9.5, color=ACCENT)
para('Datasets Crossed:', bold=True)
bullet('USACE high-hazard dams (damCache) \u2014 dams whose failure would cause loss of life')
bullet('EPA RCRA/Superfund hazmat sites (rcraCache, superfundCache) \u2014 hazardous waste locations')
bullet('EPA SDWIS drinking water systems (sdwisCache) \u2014 public water supply intakes')
para(
    'Reveals: High-hazard dams positioned such that failure would flood hazardous waste sites before reaching '
    'drinking water intakes \u2014 a three-step catastrophic cascade: dam breach \u2192 hazmat release \u2192 drinking '
    'water contamination.'
)
para(
    'Algorithm: Filters to high-hazard dams, finds hazmat sites within 40 km downstream, then drinking water '
    'systems in same radius. CRITICAL if \u22653 hazmat sites in flood path OR population >50K.'
)

# 2.5
heading('5. Drought Severity \u00d7 Reservoir Depletion \u00d7 Violation Clustering', level=2)
para('Function: discoverDroughtWaterStress()', bold=True, size=9.5, color=ACCENT)
para('Datasets Crossed:', bold=True)
bullet('USDA Drought Monitor (usdmCache) \u2014 drought intensity classification (D0\u2013D4)')
bullet('Interior Bureau of Reclamation reservoirs (usbrCache) \u2014 reservoir storage levels')
bullet('EPA SDWIS violations (sdwisCache) \u2014 drinking water compliance failures')
para(
    'Reveals: States experiencing severe drought where declining reservoir levels correlate with rising '
    'drinking water violations \u2014 a compound water stress signal indicating drought is not just a supply '
    'issue but a quality and compliance crisis.'
)
para(
    'Algorithm: Calculates severe drought percentage (D2+D3+D4), filters to states with >20% severe drought, '
    'identifies reservoirs below 40% capacity. CRITICAL if >5% exceptional drought (D4).'
)

# 2.6
heading('6. Wastewater Pathogen Detection \u00d7 Drinking Water Arrival Time', level=2)
para('Function: discoverWastewaterPathogenRisk()', bold=True, size=9.5, color=ACCENT)
para('Datasets Crossed:', bold=True)
bullet('CDC NWSS wastewater pathogen surveillance (cdcNwssCache) \u2014 pathogen detections in wastewater')
bullet('USGS stream flow routing (HUC-8 flow estimation) \u2014 travel time calculations')
bullet('EPA SDWIS drinking water intakes (sdwisCache) \u2014 surface water intake locations')
para(
    'Reveals: Pathogens detected in upstream wastewater with quantified arrival time to downstream drinking '
    'water intakes. This is predictive epidemiology \u2014 giving utilities hours to days of advance warning '
    'before a pathogen reaches their intake.'
)
para(
    'Algorithm: Groups pathogen detections by HUC-8, filters intakes to surface water, calculates flow '
    'travel time for upstream detections \u226472 hours. CRITICAL if estimated arrival <24 hours.'
)
para(
    'Why impossible before: CDC wastewater surveillance and EPA drinking water locations are in completely '
    'separate systems. The pathogen travel time calculation requires HUC-8 flow routing \u2014 a hydrological '
    'computation crossing agency boundaries that no single agency could build alone.',
    italic=True, size=9.5
)

doc.add_paragraph()
heading('Additional Derived Correlations', level=2)
add_table(
    ['Correlation', 'Sources Crossed', 'What It Reveals'],
    [
        ('Mortality \u00d7 Violation Density', 'CDC WONDER + SDWIS', 'Counties with elevated environmental death rates in high-violation areas'),
        ('Climate Normals \u00d7 WQ Trends', 'NOAA CDO + WQP', 'Regions where climate shifts correlate with water quality degradation'),
        ('Hospital Capacity \u00d7 Env Health Risk', 'HealthData.gov + Env Health', 'High-risk areas with limited hospital capacity'),
        ('DoD PFAS \u00d7 Civilian Drinking Water', 'dodPfasCache + SDWIS', 'Military PFAS plumes migrating toward civilian intakes'),
        ('Wildfire \u00d7 Watershed Impairment', 'NASA FIRMS + ATTAINS', 'Post-fire water quality degradation prediction'),
        ('Power Emissions \u00d7 Downwind WQ', 'EPA CAMPD + WQP', 'Atmospheric deposition impact on water quality'),
        ('Groundwater \u00d7 Surface Water Quality', 'NGWMN + NWIS IV/WQP', 'Aquifer-surface interaction anomalies'),
    ]
)

doc.add_page_break()

# ================================================================
# APPENDIX: DATA SOURCE INVENTORY
# ================================================================
heading('Appendix: Federal Data Source Inventory')
para('Complete inventory of all data sources integrated into PIN, organized by managing agency.')

doc.add_paragraph()
heading('EPA (Environmental Protection Agency)', level=2)
add_table(
    ['#', 'Source', 'Data Type', 'Cache Module', 'Frequency'],
    [
        ('1', 'Water Quality Portal (WQP)', 'Discrete WQ samples', 'wqpCache.ts', 'Daily'),
        ('2', 'ATTAINS', 'Impairment assessments & TMDLs', 'attainsCache.ts', '30 min'),
        ('3', 'ECHO Compliance', 'NPDES permits & compliance', 'echoCache.ts', 'Daily'),
        ('4', 'ECHO DMR Violations', 'Discharge monitoring violations', 'echoDmrViolationsCache.ts', 'Daily'),
        ('5', 'ECHO Effluent', 'Effluent monitoring data', 'echoEffluentCache.ts', 'Daily'),
        ('6', 'SDWIS', 'Drinking water violations', 'sdwisCache.ts', 'Daily'),
        ('7', 'ICIS Enforcement', 'NPDES enforcement actions', 'icisCache.ts', 'Daily'),
        ('8', 'ICIS Air', 'Air emission compliance', 'icisAirCache.ts', 'Daily'),
        ('9', 'FRS', 'Facility Registry Service', 'frsCache.ts', 'Daily'),
        ('10', 'TRI', 'Toxic Release Inventory', 'triCache.ts', 'Daily'),
        ('11', 'RCRA', 'Hazardous waste facilities', 'rcraCache.ts', 'Weekly'),
        ('12', 'SEMS', 'Superfund/contaminated sites', 'semsCache.ts', 'Weekly'),
        ('13', 'Superfund NPL', 'National Priorities List', 'superfundCache.ts', 'Daily'),
        ('14', 'EJScreen', 'Block-group EJ demographics', 'ejscreenCache.ts', 'Weekly'),
        ('15', 'UCMR5 PFAS', 'PFAS detections', 'pfasCache.ts', 'Daily'),
        ('16', 'PFAS Analytics', 'PFAS facility analytics', 'epaPfasAnalyticsCache.ts', 'Daily'),
        ('17', 'OPP Pesticides', 'Pesticide registrations', 'epaOppPesticideCache.ts', 'Daily'),
        ('18', 'BEACON Beach', 'Beach pathogen advisories', 'Sentinel adapter', '5 min'),
        ('19', 'SSO/CSO', 'Sewer overflow events', 'ssoCsoCache.ts', 'Daily'),
        ('20', 'MS4 Permits', 'Stormwater permits', 'ms4PermitCache.ts', 'Daily'),
        ('21', 'WQX Modern', 'Water quality exchange', 'wqxModernCache.ts', 'Daily'),
        ('22', 'NARS', 'Aquatic resource surveys', 'narsCache.ts', 'Weekly'),
        ('23', 'CAMPD', 'Power sector emissions', 'campdCache.ts', 'Daily'),
        ('24', 'BWB', 'Water infrastructure', 'bwbCache.ts', 'Daily'),
    ]
)

doc.add_paragraph()
heading('USGS (U.S. Geological Survey)', level=2)
add_table(
    ['#', 'Source', 'Data Type', 'Cache Module', 'Frequency'],
    [
        ('25', 'NWIS Instantaneous Values', 'Real-time streamflow & chemistry', 'nwisIvCache.ts', '5 min'),
        ('26', 'NWIS Daily Values', 'Statistical daily flow', 'usgsDvCache.ts', 'Daily'),
        ('27', 'NWIS Groundwater', 'Groundwater levels', 'nwisGwCache.ts', 'Daily'),
        ('28', 'NWIS Alerts', 'Threshold exceedances', 'nwisAlertCache.ts', 'Daily'),
        ('29', 'OGC API Stations', 'Monitoring locations', 'usgsOgcCache.ts', 'Daily'),
        ('30', 'USGS WQP Cross-Ref', 'WQ cross-reference', 'usgsWqpCache.ts', 'Daily'),
        ('31', 'StreamStats', 'Basin characteristics', 'streamStatsCache.ts', 'Daily'),
        ('32', 'Water Availability', 'HUC-8 water budgets', 'usgsWaterAvailCache.ts', 'Weekly'),
        ('33', 'NGWMN/CIDA', 'Multi-agency groundwater', 'ngwmnCache.ts', 'Daily'),
        ('34', 'STN Flood Events', 'High-water marks', 'stnFloodCache.ts', 'Daily'),
        ('35', 'Earthquake/Seismic', 'Seismic activity', 'seismicCache.ts', '30 min'),
        ('36', 'Volcano', 'Volcanic activity', 'volcanoCache.ts', 'Daily'),
        ('37', 'NLCD', 'Land cover classification', 'nlcdCache.ts', 'Daily'),
    ]
)

doc.add_paragraph()
heading('NOAA (National Oceanic & Atmospheric Administration)', level=2)
add_table(
    ['#', 'Source', 'Data Type', 'Cache Module', 'Frequency'],
    [
        ('38', 'NWS Alerts', 'Severe weather alerts', 'nwsAlertCache.ts', '10 min'),
        ('39', 'NWS Forecasts', '7-day weather forecasts', 'nwsForecastCache.ts', '6 hours'),
        ('40', 'NWPS Flood', 'Flood stage observations', 'nwpsCache.ts', '30 min'),
        ('41', 'NWM', 'National Water Model', 'nwmCache.ts', '6 hours'),
        ('42', 'HEFS', 'Ensemble flood forecasts', 'hefsCache.ts', '6 hours'),
        ('43', 'NEXRAD QPE', 'Quantitative precipitation', 'nexradQpeCache.ts', 'Daily'),
        ('44', 'Climate Normals (CDO)', 'County-level normals', 'climateNormalsCache.ts', 'Weekly'),
        ('45', 'NCEI', 'Historical climate data', 'nceiCache.ts', 'Daily'),
        ('46', 'NDBC', 'Ocean buoy observations', 'ndbcCache.ts', 'Daily'),
        ('47', 'CO-OPS Tidal', 'Tidal observations', 'coopsCache.ts', '6 hours'),
        ('48', 'CO-OPS Derived', 'Tidal analytics', 'coopsDerivedCache.ts', 'Weekly'),
        ('49', 'CoastWatch', 'Satellite ocean data', 'coastwatchCache.ts', 'Daily'),
        ('50', 'GLERL', 'Great Lakes conditions', 'glerlCache.ts', 'Daily'),
        ('51', 'HAB-SOS', 'Harmful algal blooms', 'habsosCache.ts', 'Daily'),
        ('52', 'HAB Forecast', 'HAB predictions', 'habForecastCache.ts', 'Daily'),
        ('53', 'ERDDAP', 'Satellite environmental', 'erddapSatCache.ts', 'Daily'),
        ('54', 'SWDI', 'Severe weather events', 'swdiCache.ts', 'Daily'),
    ]
)

doc.add_paragraph()
heading('HHS / CDC / FDA', level=2)
add_table(
    ['#', 'Source', 'Agency', 'Cache Module', 'Frequency'],
    [
        ('55', 'Env. Tracking Network', 'CDC', 'environmentalTrackingCache.ts', 'Daily'),
        ('56', 'CDC WONDER', 'CDC', 'cdcWonderCache.ts', 'Daily'),
        ('57', 'CDC NWSS', 'CDC', 'cdcNwssCache.ts', 'Daily'),
        ('58', 'CDC PLACES', 'CDC', 'cdcPlacesCache.ts', 'Daily'),
        ('59', 'data.cdc.gov', 'CDC', 'dataCdcGovCache.ts', 'Daily'),
        ('60', 'ATSDR Toxicology', 'HHS/ATSDR', 'atsdrToxicologyCache.ts', 'Daily'),
        ('61', 'HealthData.gov', 'HHS', 'healthDataGovCache.ts', 'Daily'),
        ('62', 'Hospital Facilities', 'HHS/CMS', 'hospitalCache.ts', 'Daily'),
        ('63', 'Waterborne Illness', 'CDC', 'waterborneIllnessCache.ts', 'Daily'),
        ('64', 'Environmental Health', 'EPA/CDC', 'environmentalHealthCache.ts', 'Daily'),
        ('65', 'Open FDA', 'FDA', 'openFDACache.ts', 'Daily'),
        ('66', 'MyHealthfinder', 'HHS', 'myhealthfinderCache.ts', 'Daily'),
        ('67', 'HRSA HPSA', 'HRSA', 'hrsaHpsaCache.ts', 'Daily'),
    ]
)

doc.add_paragraph()
heading('DoD / FEMA / USACE / Interior / USDA / Other', level=2)
add_table(
    ['#', 'Source', 'Agency', 'Cache Module', 'Frequency'],
    [
        ('68', 'NID Dams', 'USACE', 'damCache.ts', 'Daily'),
        ('69', 'USACE Projects', 'USACE', 'usaceCache.ts', '6 hours'),
        ('70', 'USBR Reservoirs', 'Interior', 'usbrCache.ts', '6 hours'),
        ('71', 'FEMA Disasters', 'FEMA', 'femaCache.ts', 'Daily'),
        ('72', 'NFIP Claims', 'FEMA', 'nfipClaimsCache.ts', 'Daily'),
        ('73', 'Hazard Mitigation', 'FEMA', 'hazMitCache.ts', 'Weekly'),
        ('74', 'DoD PFAS Assessments', 'DoD', 'dodPfasCache.ts', 'Daily'),
        ('75', 'DoD PFAS Sites', 'DoD', 'dodPfasSitesCache.ts', 'Daily'),
        ('76', 'NADP PFAS', 'NADP', 'nadpPfasCache.ts', 'Daily'),
        ('77', 'SNOTEL', 'NRCS/USDA', 'snotelCache.ts', 'Daily'),
        ('78', 'USDM Drought', 'USDA', 'usdmCache.ts', 'Daily'),
        ('79', 'NASS Livestock', 'USDA', 'nassLivestockCache.ts', 'Weekly'),
        ('80', 'NASS Crops', 'USDA', 'nassCropsCache.ts', 'Weekly'),
        ('81', 'SSURGO Soils', 'USDA/NRCS', 'ssurgoCache.ts', 'Daily'),
        ('82', 'NASA CMR', 'NASA', 'nasaCmrCache.ts', 'Daily'),
        ('83', 'NASA Streamflow', 'NASA', 'nasaStreamCache.ts', 'Weekly'),
        ('84', 'FIRMS Active Fires', 'NASA', 'firmsCache.ts', '4 hours'),
        ('85', 'IPAC Species', 'USFWS', 'ipacCache.ts', 'Weekly'),
        ('86', 'USASpending', 'Treasury', 'usaSpendingCache.ts', 'Weekly'),
        ('87', 'Grants.gov', 'GSA', 'grantsGovCache.ts', 'Daily'),
        ('88', 'SAM.gov', 'GSA', 'samGovCache.ts', 'Weekly'),
        ('89', 'Data.gov', 'GSA', 'dataGovCache.ts', 'Weekly'),
        ('90', 'Congress.gov', 'LOC', 'congressCache.ts', 'Daily'),
        ('91', 'CEDEN (California)', 'CA SWRCB', 'cedenCache.ts', 'Daily'),
        ('92', 'OSHA Water', 'DOL/OSHA', 'oshaWaterCache.ts', 'Daily'),
        ('93', 'Embassy AQI', 'State Dept', 'embassyAqiCache.ts', '15 min'),
        ('94', 'PHMSA Pipelines', 'DOT', 'phmsaPipelineCache.ts', 'Daily'),
        ('95', 'Hypoxia', 'Various', 'hypoxiaCache.ts', 'Daily'),
        ('96', 'Copernicus CDS', 'ECMWF (EU)', 'copernicusCdsCache.ts', 'Weekly'),
        ('97', 'GEMStat', 'UNEP', 'gemstatCache.ts', 'Weekly'),
    ]
)

doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run(
    'PIN is built and maintained by the PEARL team. All data sources referenced in this document are '
    'publicly available through federal government APIs. No classified, proprietary, or access-restricted '
    'data is used.'
)
r.font.size = Pt(9)
r.font.color.rgb = GRAY
r.italic = True

# ================================================================
# SAVE
# ================================================================
output_path = os.path.expanduser('~/Downloads/PIN-Platform-Overview.docx')
os.makedirs(os.path.dirname(output_path), exist_ok=True)
doc.save(output_path)
print(f'Saved to {output_path}')
