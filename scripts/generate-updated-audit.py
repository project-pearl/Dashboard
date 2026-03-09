"""Generate updated PIN Dashboard Audit Report as .docx"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

doc = Document()

# -- Styles --
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# -- Title --
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('PIN Water Quality Dashboard')
run.bold = True
run.font.size = Pt(22)
run.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Comprehensive Technical Audit Report')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Updated March 08, 2026 (Rev 2)')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

doc.add_paragraph()

# -- Executive Summary --
doc.add_heading('Executive Summary', level=1)
doc.add_paragraph(
    'This report evaluates the PIN Water Quality Dashboard across 12 engineering domains. '
    'Each domain receives a letter grade (A through F), an explanation of the current state, '
    'the target state for production/government readiness, and a concrete roadmap to close gaps. '
    'The dashboard demonstrates strong data architecture, a production-grade design system, '
    'and has recently implemented significant security hardening, observability, and compliance infrastructure. '
    'This revision reflects newly added SBOM generation, README, .env.example, PR template, and Prettier config. '
    'Remaining gaps are primarily in test coverage depth and build-time enforcement.'
)

# -- Grade Summary Table --
doc.add_heading('Grade Summary', level=1)

grades = [
    ('Security', 'B+', 'HIGH', 'CSP, CSRF, HSTS, Zod validation, Upstash rate limiting all implemented. Next.js 15.3.0 current.'),
    ('Testing & QA', 'D+', 'CRITICAL', 'Framework installed (Vitest + Playwright + MSW). 15 test files exist but coverage is ~15%.'),
    ('CI/CD & DevOps', 'B+', 'MEDIUM', '3 workflows (SAST + SBOM, test, pipeline). CODEOWNERS, Husky, PR template configured.'),
    ('Code Quality & Maintainability', 'B-', 'HIGH', 'Global error boundary, Sentry, .prettierrc. But ignoreDuringBuilds/ignoreBuildErrors still enabled.'),
    ('Documentation', 'B-', 'MEDIUM', 'README, .env.example, incident response, NIST mapping, secrets rotation docs all in place.'),
    ('Accessibility (a11y)', 'C+', 'MEDIUM', 'Radix UI primitives, 38+ ARIA attributes, prefers-reduced-motion. No automated a11y testing.'),
    ('Performance Optimization', 'B', 'MEDIUM', 'Speed Insights RUM, dynamic imports, 3-tier caching. Image optimization disabled.'),
    ('SEO & Discoverability', 'B-', 'LOW', 'robots.txt, sitemap.ts, JSON-LD Organization schema, OpenGraph/Twitter metadata.'),
    ('Design System & UX', 'A-', 'LOW', '196 CSS custom properties, Radix UI + Shadcn/ui, dark mode, 4 accent presets.'),
    ('Data Architecture & Caching', 'B+', 'LOW', '51 cache modules, 3-tier persistence, grid-based spatial indexing, self-chaining ATTAINS.'),
    ('Monitoring & Observability', 'B+', 'HIGH', 'Sentry on all 52 crons, Slack failure alerts, cronHealth tracking, ops dashboard, Speed Insights.'),
    ('Compliance & Governance', 'B', 'MEDIUM', 'Audit logging, CODEOWNERS, NIST mapping, incident response, secrets rotation. CycloneDX SBOM in CI.'),
]

table = doc.add_table(rows=1, cols=4)
table.style = 'Light Shading Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table.rows[0].cells
hdr[0].text = 'Domain'
hdr[1].text = 'Grade'
hdr[2].text = 'Priority'
hdr[3].text = 'Summary'
for h in hdr:
    for p in h.paragraphs:
        p.runs[0].bold = True

for domain, grade, priority, summary in grades:
    row = table.add_row().cells
    row[0].text = domain
    row[1].text = grade
    row[2].text = priority
    row[3].text = summary

doc.add_paragraph()
doc.add_paragraph('Composite Grade: B-', style='Intense Quote')

# -- Detailed Domain Assessments --
doc.add_heading('Detailed Domain Assessments', level=1)

# ---- 1. Security ----
doc.add_heading('1. Security', level=2)
doc.add_paragraph('B+', style='Intense Quote')

doc.add_heading('Current State', level=3)
items = [
    'Next.js 15.3.0 (current stable — closes all 15 known CVEs from the 13.x era)',
    'Content-Security-Policy header with per-request nonce generation (middleware.ts)',
    'HSTS: max-age=63072000; includeSubDomains; preload',
    'CSRF double-submit cookie protection on all POST/PUT/DELETE/PATCH (middleware.ts)',
    'Client-side CSRF helper (lib/csrf.ts) for fetch requests',
    'Zod schema validation on all API route inputs — 24 named schemas (lib/schemas.ts)',
    'Request body parser with typed Zod validation (lib/validateRequest.ts)',
    'Upstash Redis rate limiting: 10 req/60s per IP with graceful fallback (lib/rateLimit.ts)',
    'Security headers: X-Content-Type-Options, Referrer-Policy, Permissions-Policy, X-DNS-Prefetch-Control',
    'Supabase Auth with JWT + RBAC (4-tier admin hierarchy)',
    '/api/cache-status and /api/water-data are auth-protected',
    'dangerouslySetInnerHTML: 2 instances, both safe (game CSS, chart config)',
    'Semgrep SAST + npm audit in CI (security-scan.yml)',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('Remaining Gaps', level=3)
gaps = [
    'All 52 cron endpoints share a single CRON_SECRET (compromise of one = compromise of all)',
    'npm audit may still show vulnerabilities from transitive dependencies',
]
for g in gaps:
    doc.add_paragraph(g, style='List Bullet')

doc.add_heading('How to Reach an A', level=3)
steps = [
    'Separate cron secrets by job category (data-rebuild, sentinel, insights)',
    'Add npm audit --audit-level=critical as blocking step in CI',
    'Replace dangerouslySetInnerHTML with CSS modules where possible',
    'Add Subresource Integrity (SRI) for external scripts if any are added',
]
for s in steps:
    doc.add_paragraph(s, style='List Number')

# ---- 2. Testing ----
doc.add_heading('2. Testing & Quality Assurance', level=2)
doc.add_paragraph('D+', style='Intense Quote')

doc.add_heading('Current State', level=3)
items = [
    'Vitest ^4.0.18 installed with vitest.config.ts (coverage thresholds: 15% statements, 5% branches)',
    '@testing-library/react ^16.3.2 + @testing-library/jest-dom + @testing-library/user-event',
    'Playwright ^1.58.2 with playwright.config.ts (Chromium, sequential, screenshots on failure)',
    'MSW (Mock Service Worker) with handlers.ts and fixtures for ATTAINS, ICIS, WQP data',
    '15 test files total:',
    '  - 7 unit tests (apiAuth, attainsCache, authTypes, blobPersistence, cacheUtils, icisCache, wqpCache)',
    '  - 3 integration tests (alerts-history, cache-status, cron-rebuild-wqp)',
    '  - 2 component tests (FederalManagementCenter, login-page)',
    '  - 3 E2E tests (login, dashboard-federal, dashboard-role-based)',
    'Test scripts: test, test:watch, test:ui, test:coverage, test:e2e',
    'Husky pre-commit hooks run vitest on staged files via lint-staged',
    'CI test workflow runs unit/integration tests + uploads coverage artifact',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('Remaining Gaps', level=3)
gaps = [
    '15 test files across 607 source files — coverage is approximately 15%',
    'Most cache modules (44 of 51) have no dedicated tests',
    'Most API routes (90+ of 105) have no integration tests',
    'Most components (230+ of 236) have no component tests',
    'No mutation/property-based testing',
]
for g in gaps:
    doc.add_paragraph(g, style='List Bullet')

doc.add_heading('How to Reach an A', level=3)
steps = [
    'Add unit tests for remaining cache modules (start with sdwisCache, echoCache, pfasCache)',
    'Add integration tests for all auth-protected API routes',
    'Add component tests for PEARLManagementCenter, AlertsManagementPanel',
    'Add E2E tests for alert creation, CSV upload, cron trigger flows',
    'Target: 40% coverage in 30 days, 70% in 90 days, 80% in 6 months',
    'Add coverage reports as PR comments via CI',
]
for s in steps:
    doc.add_paragraph(s, style='List Number')

# ---- 3. CI/CD ----
doc.add_heading('3. CI/CD & DevOps', level=2)
doc.add_paragraph('B+', style='Intense Quote')

doc.add_heading('Current State', level=3)
items = [
    '3 GitHub Actions workflows:',
    '  - security-scan.yml: Semgrep SAST (SARIF upload) + npm audit (artifact) + CycloneDX SBOM generation',
    '  - test.yml: Vitest unit/integration + Playwright E2E (coverage artifact)',
    '  - pin-pipeline.yml: Hourly incremental data fetch with auto-commit',
    '.github/CODEOWNERS: lib/, app/api/admin/, app/api/cron/, lib/sentinel/, .github/, vercel.json',
    '.github/PULL_REQUEST_TEMPLATE.md: Security/testing checklist for all PRs',
    'Husky pre-commit hooks with lint-staged running vitest on changed files',
    'Branch protection rules configured (PR required + 2 status checks)',
    'Vercel deployment via git push with preview deployments on PRs',
    'npm run sbom: Local SBOM generation script (CycloneDX JSON)',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('Remaining Gaps', level=3)
gaps = [
    'No Dependabot or Renovate for automated dependency updates',
    'No staging environment with smoke tests before production',
]
for g in gaps:
    doc.add_paragraph(g, style='List Bullet')

doc.add_heading('How to Reach an A', level=3)
steps = [
    'Add Dependabot or Renovate for automated dependency updates',
    'Add Slack notification on workflow failure',
    'Consider staging deployment step before production',
]
for s in steps:
    doc.add_paragraph(s, style='List Number')

# ---- 4. Code Quality ----
doc.add_heading('4. Code Quality & Maintainability', level=2)
doc.add_paragraph('B-', style='Intense Quote')

doc.add_heading('Current State', level=3)
items = [
    'TypeScript strict mode enabled in tsconfig.json',
    'Global error boundary (app/global-error.tsx) with Sentry integration',
    'Route-level error boundary (app/dashboard/federal/error.tsx)',
    'console.error and console.warn preserved in production (removeConsole excludes both)',
    'Husky + lint-staged pre-commit hooks enforcing tests on changed files',
    '.prettierrc configured (semi, singleQuote, trailingComma, tabWidth 2, printWidth 100)',
    'Shared cache helpers in cacheUtils.ts (gridKey, neighborKeys, loadCacheFromDisk, saveCacheToDisk)',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('Remaining Gaps', level=3)
gaps = [
    'ignoreDuringBuilds: true and ignoreBuildErrors: true still suppress all ESLint/TS errors at build time',
    'Cache module persistence code still duplicated across 51 modules (no base class/factory)',
    'ATTAINS cache still uses its own blob implementation instead of shared blobPersistence.ts',
]
for g in gaps:
    doc.add_paragraph(g, style='List Bullet')

doc.add_heading('How to Reach an A', level=3)
steps = [
    'Remove ignoreDuringBuilds and ignoreBuildErrors; fix all underlying errors',
    'Enforce Prettier formatting via pre-commit hooks (add to lint-staged)',
    'Create a CacheModule base class/factory to eliminate persistence duplication',
    'Migrate ATTAINS blob implementation to shared blobPersistence.ts',
]
for s in steps:
    doc.add_paragraph(s, style='List Number')

# ---- 5. Documentation ----
doc.add_heading('5. Documentation', level=2)
doc.add_paragraph('B-', style='Intense Quote')

doc.add_heading('Current State', level=3)
items = [
    'README.md: Project overview, architecture, data pipeline, roles, setup, scripts, project structure, security, monitoring, deployment',
    '.env.example: 49 environment variables documented with categories and descriptions',
    'docs/INCIDENT_RESPONSE.md (98 lines): Severity levels P1-P4, roles, detection, triage, containment, recovery, comms templates',
    'docs/NIST-800-53-CONTROLS.md (24 lines): 11 controls mapped (AC-2, AC-6, AU-2/3/6, CM-2/8, IA-2/5, SC-8/13)',
    'docs/SECRETS_ROTATION.md (73 lines): 8 secrets inventoried, quarterly rotation, emergency procedures',
    'docs/sentinel-system-architecture.md (271 lines): Comprehensive Sentinel architecture',
    'docs/METRICS.md: Sentinel validation metrics',
    'docs/migrations/001_admin_audit_log.sql: Audit log table schema',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('Remaining Gaps', level=3)
gaps = [
    'No CONTRIBUTING.md with development workflow and PR process',
    'No API reference documenting the 105 endpoints',
    'Architecture knowledge in Claude memory files is not accessible to other developers',
]
for g in gaps:
    doc.add_paragraph(g, style='List Bullet')

doc.add_heading('How to Reach an A', level=3)
steps = [
    'Create CONTRIBUTING.md with dev setup, branch naming, PR checklist',
    'Generate API reference from route files',
    'Add JSDoc to all exported functions in lib/',
]
for s in steps:
    doc.add_paragraph(s, style='List Number')

# ---- 6. Accessibility ----
doc.add_heading('6. Accessibility (a11y)', level=2)
doc.add_paragraph('C+', style='Intense Quote')

doc.add_heading('Current State', level=3)
items = [
    'Radix UI primitives provide built-in keyboard navigation, focus management, and ARIA roles',
    '30+ Shadcn/ui components with accessibility baked in',
    '38+ ARIA attributes across components (aria-label, aria-live, aria-expanded, etc.)',
    'Screen reader text (.sr-only) on carousel controls, dialog close buttons',
    'Reduced motion support: @media (prefers-reduced-motion) in globals.css',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('Remaining Gaps', level=3)
gaps = [
    'No automated a11y testing tools installed (no jest-axe, pa11y, or axe-core)',
    'No documented WCAG compliance level',
    'No color contrast verification across light and dark themes',
]
for g in gaps:
    doc.add_paragraph(g, style='List Bullet')

# ---- 7. Performance ----
doc.add_heading('7. Performance Optimization', level=2)
doc.add_paragraph('B', style='Intense Quote')

doc.add_heading('Current State', level=3)
items = [
    '@vercel/speed-insights for Real User Monitoring (RUM) — installed and used in layout.tsx',
    'Dynamic imports for heavy components (Mapbox, charts) with ssr: false',
    'Suspense boundaries with skeleton fallbacks',
    'next/image for automatic lazy loading; next/font/google for zero-layout-shift fonts',
    '3-tier caching (memory -> disk -> Vercel Blob) prevents cold-start data loss',
    'console.log stripped in production; warn/error preserved',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('Remaining Gaps', level=3)
gaps = [
    'images.unoptimized: true — Next.js image optimization disabled globally',
    'No @next/bundle-analyzer installed',
    'No Lighthouse CI or Core Web Vitals tracking in CI',
]
for g in gaps:
    doc.add_paragraph(g, style='List Bullet')

# ---- 8. SEO ----
doc.add_heading('8. SEO & Discoverability', level=2)
doc.add_paragraph('B-', style='Intense Quote')

doc.add_heading('Current State', level=3)
items = [
    'public/robots.txt: Allows public pages, blocks /dashboard/, /api/, /_next/',
    'app/sitemap.ts: 9 public routes with priority weighting',
    'Root metadata: title template, description, OpenGraph, Twitter card with images',
    'JSON-LD Organization schema in layout.tsx',
    'Canonical URL support via alternates.canonical',
    'PWA manifest.json with app name, icons, theme color',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('Remaining Gaps', level=3)
gaps = [
    'No per-page generateMetadata for state report or jurisdiction pages',
    'No structured data beyond Organization (no BreadcrumbList, FAQ, etc.)',
]
for g in gaps:
    doc.add_paragraph(g, style='List Bullet')

# ---- 9. Design System ----
doc.add_heading('9. Design System & UX', level=2)
doc.add_paragraph('A-', style='Intense Quote')

doc.add_heading('Current State', level=3)
items = [
    'Tailwind CSS 3.3 + Radix UI + Shadcn/ui: 30+ composable UI components',
    '196 CSS custom properties (HSL-based design tokens)',
    'Dark mode with class-based toggle + full CSS overrides',
    '4 accent color presets (Sapphire, Indigo, Emerald, Copper)',
    'Domain-specific semantic colors (surface water teal, groundwater amber)',
    'Typography: Inter (body) + JetBrains Mono (data/stats)',
    'Animations respect prefers-reduced-motion',
    'Theme persistence via next-themes + localStorage',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('Remaining Gaps', level=3)
gaps = [
    'No Storybook for component documentation',
    'No visual regression testing (Chromatic, Percy)',
]
for g in gaps:
    doc.add_paragraph(g, style='List Bullet')

# ---- 10. Data Architecture ----
doc.add_heading('10. Data Architecture & Caching', level=2)
doc.add_paragraph('B+', style='Intense Quote')

doc.add_heading('Current State', level=3)
items = [
    '51 cache modules with grid-based spatial indexing (0.1-degree resolution)',
    '3-tier persistence: in-memory -> disk (.cache/) -> Vercel Blob (REST API)',
    'Cold-start survival: ensureWarmed() cascades through all three tiers',
    'Build lock pattern with 12-minute auto-clear timeout',
    '52 cron jobs staggered across 4-11 AM UTC',
    'ATTAINS self-chaining: 20-hop cascade with time-budgeted chunks',
    'Delta detection via signalsHash for efficient rebuild skipping',
    'Semaphore-based concurrency control (4 parallel states) for AI insights',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('Remaining Gaps', level=3)
gaps = [
    'No SWR or React Query for client-side data fetching (raw fetch + useEffect)',
    'No shared CacheModule base class — persistence pattern duplicated across 51 modules',
    'No Redis/Upstash for distributed cache state',
]
for g in gaps:
    doc.add_paragraph(g, style='List Bullet')

# ---- 11. Monitoring ----
doc.add_heading('11. Monitoring & Observability', level=2)
doc.add_paragraph('B+', style='Intense Quote')

doc.add_heading('Current State', level=3)
items = [
    '@sentry/nextjs ^10.42.0: error tracking + performance monitoring',
    'Sentry config files: sentry.client.config.ts, sentry.server.config.ts, sentry.edge.config.ts',
    'instrumentation.ts: Next.js instrumentation hook for runtime-aware Sentry loading',
    'next.config.js wrapped with withSentryConfig (hideSourceMaps, automaticVercelMonitors)',
    'Global error boundary (app/global-error.tsx) with Sentry.captureException',
    'All 52 cron routes instrumented with: Sentry.captureException, notifySlackCronFailure, recordCronRun',
    'lib/slackNotify.ts: Fire-and-forget Slack webhook on cron failures (Block Kit, 5s timeout)',
    'lib/cronHealth.ts: In-memory ring buffer (50 runs/cron) with disk + blob persistence, debounced saves',
    'app/api/admin/cron-health/route.ts: Auth-gated API returning summary + history',
    'components/CronHealthDashboard.tsx: Ops dashboard with summary strip, success rate chart, cache grid, failure table',
    '@vercel/speed-insights for Real User Monitoring (RUM)',
    'Custom health check endpoint (/api/source-health) monitoring 35 external data sources',
    'Unified cache status endpoint (/api/cache-status) tracking 50+ caches with staleness detection',
    'Delta tracking (before/after record counts) on cache rebuilds',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('Remaining Gaps', level=3)
gaps = [
    'No structured logging (JSON format) for log aggregation — still console.log/warn/error',
    'No Vercel log drains configured (Axiom, Datadog, Better Stack)',
    'No uptime monitoring service (PagerDuty, Opsgenie)',
]
for g in gaps:
    doc.add_paragraph(g, style='List Bullet')

doc.add_heading('How to Reach an A', level=3)
steps = [
    'Set up Vercel log drains to Axiom or Better Stack',
    'Add structured JSON logging wrapper',
    'Add uptime monitoring with alerting',
]
for s in steps:
    doc.add_paragraph(s, style='List Number')

# ---- 12. Compliance ----
doc.add_heading('12. Compliance & Governance', level=2)
doc.add_paragraph('B', style='Intense Quote')

doc.add_heading('Current State', level=3)
items = [
    'Supabase RBAC with 4-tier admin hierarchy',
    'lib/auditLog.ts: Fire-and-forget admin audit logging to admin_audit_log Supabase table',
    '.github/CODEOWNERS for required review on critical paths',
    'Branch protection rules: require PR, require 2 status checks',
    'Semgrep SAST + npm audit in CI workflow',
    'CycloneDX SBOM generated on every CI run (security-scan.yml sbom job) — satisfies EO 14028 CM-8',
    'npm run sbom: Local SBOM generation script',
    'docs/INCIDENT_RESPONSE.md: Full severity levels, roles, containment, recovery procedures',
    'docs/NIST-800-53-CONTROLS.md: 11 controls mapped (AC, AU, CM, IA, SC families)',
    'docs/SECRETS_ROTATION.md: 8 secrets inventoried with quarterly rotation schedule',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('Remaining Gaps', level=3)
gaps = [
    'No data retention or privacy policy implementation',
    'No secrets rotation automation (manual quarterly process)',
    'NIST mapping covers 11 of 20+ applicable controls — could expand to AC-4, SC-7, SI-4, etc.',
]
for g in gaps:
    doc.add_paragraph(g, style='List Bullet')

doc.add_heading('How to Reach an A', level=3)
steps = [
    'Expand NIST 800-53 mapping to cover AC-4, SC-7, SI-4, AU-10',
    'Implement automated secrets rotation where possible',
    'Add data classification and retention policies',
]
for s in steps:
    doc.add_paragraph(s, style='List Number')

# -- Overall Assessment --
doc.add_heading('Overall Assessment', level=1)
doc.add_paragraph('Composite Grade: B-', style='Intense Quote')

doc.add_paragraph(
    'The PIN Dashboard has made substantial progress across all 12 audit domains since the initial assessment. '
    'Security has been hardened with CSP, CSRF, Zod validation, and Upstash rate limiting. '
    'The observability stack (Sentry, Slack alerts, cron health tracking) provides production-grade monitoring. '
    'Compliance infrastructure is comprehensive: incident response docs, NIST 800-53 mapping, secrets rotation, '
    'CycloneDX SBOM generation in CI (EO 14028), and admin audit logging. '
    'Developer documentation now includes a full README, .env.example with 49 variables, and PR templates. '
    'The primary remaining gaps are test coverage depth (15 test files across 607 source files) '
    'and build-time enforcement (ignoreDuringBuilds still enabled).'
)

# -- DARPA Positioning --
doc.add_heading('DARPA Positioning Impact', level=2)
doc.add_paragraph(
    'The security and compliance posture has improved significantly. Key DARPA-relevant strengths:'
)
darpa_items = [
    'NIST 800-53 control mapping demonstrates security awareness (11 controls mapped)',
    'CycloneDX SBOM generated on every CI run — satisfies EO 14028 software supply chain requirement',
    'Incident response documentation shows operational maturity',
    'Sentry + Slack + cronHealth monitoring provides reliability evidence',
    'SAST scanning in CI shows secure development lifecycle',
    'Comprehensive README and .env.example demonstrate developer onboarding readiness',
]
for item in darpa_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph('Remaining DARPA concerns:')
darpa_gaps = [
    'Test coverage at ~15% — insufficient for correctness claims',
    'ignoreBuildErrors: true means type safety is not enforced',
]
for g in darpa_gaps:
    doc.add_paragraph(g, style='List Bullet')

# -- Priority Roadmap --
doc.add_heading('Priority Roadmap', level=1)

doc.add_heading('Week 1-2 (Critical)', level=3)
items = [
    'DONE: CycloneDX SBOM generation added to CI (closes EO 14028 gap)',
    'DONE: Comprehensive README.md written (architecture, setup, scripts, security)',
    'DONE: .env.example created with 49 documented environment variables',
    'DONE: .github/PULL_REQUEST_TEMPLATE.md created with security/testing checklist',
    'DONE: .prettierrc configured for consistent formatting',
    'Add 10 more cache module unit tests (target: 30% coverage)',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('Week 3-4 (High)', level=3)
items = [
    'Remove ignoreDuringBuilds / ignoreBuildErrors; fix all build errors',
    'Add integration tests for remaining auth-protected API routes',
    'Add Dependabot for automated dependency updates',
    'Set up Vercel log drains (Axiom)',
    'Enforce Prettier formatting via pre-commit hooks',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('Month 2 (Medium)', level=3)
items = [
    'Reach 40% test coverage on critical paths',
    'Add accessibility testing (jest-axe + pa11y)',
    'Re-enable image optimization',
    'Separate cron secrets by category',
    'Expand NIST 800-53 mapping',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('Month 3 (Polish)', level=3)
items = [
    'Reach 70% test coverage',
    'Add Storybook for component documentation',
    'Adopt SWR/React Query for client-side data fetching',
    'Create CacheModule base class to eliminate duplication',
    'Add Lighthouse CI for performance regression tracking',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

# -- Final Grade Table --
doc.add_heading('Final Grade Comparison', level=1)

table = doc.add_table(rows=1, cols=4)
table.style = 'Light Shading Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table.rows[0].cells
hdr[0].text = 'Domain'
hdr[1].text = 'Initial Grade'
hdr[2].text = 'Current Grade'
hdr[3].text = 'Change'
for h in hdr:
    for p in h.paragraphs:
        p.runs[0].bold = True

comparisons = [
    ('Security', 'D+', 'B+', '+4 levels'),
    ('Testing & QA', 'F', 'D+', '+2 levels'),
    ('CI/CD & DevOps', 'C+', 'B+', '+2 levels'),
    ('Code Quality', 'C-', 'B-', '+2 levels'),
    ('Documentation', 'D', 'B-', '+2.5 levels'),
    ('Accessibility', 'C+', 'C+', 'No change'),
    ('Performance', 'B-', 'B', '+0.5 level'),
    ('SEO', 'C-', 'B-', '+1.5 levels'),
    ('Design System', 'A-', 'A-', 'No change'),
    ('Data Architecture', 'B+', 'B+', 'No change'),
    ('Monitoring', 'C-', 'B+', '+3 levels'),
    ('Compliance', 'D', 'B', '+3 levels'),
]
for domain, initial, current, change in comparisons:
    row = table.add_row().cells
    row[0].text = domain
    row[1].text = initial
    row[2].text = current
    row[3].text = change

doc.add_paragraph()
doc.add_paragraph('Initial Composite: C-  |  Current Composite: B-', style='Intense Quote')

# -- Save --
output = r'C:\Users\Doug\Downloads\PIN_Dashboard_Audit_Report.docx'
doc.save(output)
print(f'Report saved to: {output}')
